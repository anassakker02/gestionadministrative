#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Marges ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

DARK_BLUE = RGBColor(0x1F, 0x39, 0x64)
MID_BLUE  = RGBColor(0x2E, 0x5E, 0xA8)
RED       = RGBColor(0xC0, 0x39, 0x2B)
ORANGE    = RGBColor(0xE6, 0x7E, 0x22)
GREEN     = RGBColor(0x1E, 0x88, 0x55)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY      = RGBColor(0x66, 0x66, 0x66)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading1(text, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = color
    # Bordure basse bleue
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F3964')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text, color=MID_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return p

def body(text, bold=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_table(headers, rows, header_bg='1F3964'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_bg)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return t

def page_break():
    doc.add_page_break()

BASE = '/Users/anass/Downloads/frais-gestionScolaire 4'

def add_image(rel_path, caption=None, width=Inches(6)):
    import os
    full = os.path.join(BASE, rel_path)
    if not os.path.exists(full):
        body(f'[Image non trouvee : {rel_path}]', color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(full, width=width)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = GRAY

# ============================================================
# PAGE DE GARDE
# ============================================================
# Bandeau bleu titre
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1F3964')
pPr.append(shd)
run = p.add_run('YNOV CAMPUS — PROJET DE FIN D\'ETUDES')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = WHITE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(20)
r2 = p2.add_run('RAPPORT DE SECURITE COMPLET')
r2.bold = True; r2.font.size = Pt(22); r2.font.color.rgb = DARK_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Industrialisation, Securisation et Mise en Production\nApplication de Gestion Administrative Firebase/React')
r3.italic = True; r3.font.size = Pt(11); r3.font.color.rgb = MID_BLUE
p3.paragraph_format.space_after = Pt(20)

add_table(
    ['Champ', 'Information'],
    [
        ['Auteur', 'Anass Akker - YNOV Campus 2026'],
        ['Date', '12 avril 2026'],
        ['Version', 'v3.0 - Final Complet'],
        ['Agent Wazuh', '001 - main-machine (macOS 15.7.4)'],
        ['Manager', 'Wazuh 4.7.4 - Docker (Debian)'],
        ['Stack', 'Firebase / React / Express.js / Node.js'],
        ['Modules', 'FIM - CVE - MITRE ATT&CK - SCA - Security Events - Rootcheck'],
        ['RGPD', 'Art. 5, 17, 25, 32, 33, 35 couverts'],
    ]
)

# ============================================================
# TABLE DES MATIERES
# ============================================================
page_break()

p_toc_title = doc.add_paragraph()
p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_toc_title.add_run('TABLE DES MATIERES')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = DARK_BLUE
p_toc_title.paragraph_format.space_after = Pt(16)

toc_entries = [
    ('Introduction -- Annexes Dashboard de Monitoring de Securite', '3', False),
    ('1. TABLEAU DE BORD -- KPI SECURITE', '4', False),
    ('    1.1 Statut Global des Modules', '4', True),
    ('    1.2 Chronologie de l\'Audit', '4', True),
    ('2. PROBLEMATIQUES RENCONTREES & SOLUTIONS', '5', False),
    ('    2.1 Agent macOS non visible dans Wazuh (Disconnected)', '5', True),
    ('    2.2 File Integrity Monitoring sans resultats', '5', True),
    ('    2.3 NVD Feed CVE non telecharge (payant)', '5', True),
    ('    2.4 Backend Node.js ne demarre pas -- bcrypt DLL Windows', '6', True),
    ('    2.5 Firebase Emulator timeout -- defineSecret', '6', True),
    ('    2.6 Regles Wazuh -- champ action reserve', '6', True),
    ('    2.7 Permissions macOS SIP -- /etc write denied', '6', True),
    ('    2.8 Presentation PPTX corrompue -- Bad CRC-32', '7', True),
    ('    2.9 Recapitulatif -- Bilan des 8 Problemes', '7', True),
    ('3. ARCHITECTURE DE SURVEILLANCE WAZUH', '8', False),
    ('    3.1 Composants de l\'Infrastructure', '8', True),
    ('    3.2 Regles YNOV-APP Personnalisees', '8', True),
    ('4. RESULTATS DETAILLES PAR MODULE', '9', False),
    ('    4.1 Infrastructure Agents -- Coverage 100%', '9', True),
    ('    4.2 Security Events -- 436 420 Evenements', '9', True),
    ('    4.3 File Integrity Monitoring (FIM) -- 436 420 Evenements', '10', True),
    ('    4.4 Policy Monitoring -- Rootcheck (Detection Anomalies)', '11', True),
    ('    4.5 Detection CVE -- 17 Vulnerabilites', '11', True),
    ('    4.6 MITRE ATT&CK -- Cartographie des Techniques', '12', True),
    ('5. CONFORMITE RGPD x WAZUH', '15', False),
    ('6. PLAN D\'ACTION & RECOMMANDATIONS', '16', False),
    ('    6.1 Actions Immediates (Priorite Haute)', '16', True),
    ('    6.2 Ameliorations Monitoring', '16', True),
    ('    6.3 Bilan Final', '16', True),
    ('ANNEXE A -- Dashboard Monitoring Applicatif (Monitoring.tsx)', '17', False),
    ('    A.1 Dashboard Securite Complet : Auth - RGPD - RBAC - Audit Logs', '17', True),
    ('    A.2 WAF -- Regles OWASP Top 10', '20', True),
    ('    A.3 Logs SIEM en Temps Reel', '23', True),
    ('CONCLUSION', '25', False),
]

for entry, page, is_sub in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    # Tab stop à droite pour le numéro de page
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    pPr = p._p.get_or_add_pPr()
    tabs = _OE('w:tabs')
    tab = _OE('w:tab')
    tab.set(_qn('w:val'), 'right')
    tab.set(_qn('w:leader'), 'dot')
    tab.set(_qn('w:pos'), '8640')
    tabs.append(tab)
    pPr.append(tabs)

    run_text = p.add_run(entry + '\t' + page)
    run_text.font.size = Pt(10 if not is_sub else 9)
    run_text.font.color.rgb = GRAY if is_sub else DARK_BLUE
    if not is_sub:
        run_text.bold = True

# ============================================================
# INTRODUCTION
# ============================================================
page_break()
heading1('Introduction -- Annexes Dashboard de Monitoring de Securite')
body('Les trois captures d\'ecran presentees dans cette section illustrent le fonctionnement en conditions reelles du tableau de bord de monitoring de securite SIEM developpe dans le cadre de ce projet. Accessible uniquement aux administrateurs via la route /monitoring, ce dashboard constitue le centre de pilotage de la securite de l\'application et agregre en temps reel l\'ensemble des evenements de securite collectes depuis la collection Firestore auditLogs, avec un rafraichissement automatique toutes les 60 secondes.')
body('Ces captures ont ete realisees le 22 avril 2026 a la suite d\'une demonstration complete d\'attaques simulees couvrant l\'ensemble des vecteurs OWASP Top 10. Elles attestent du bon fonctionnement de l\'ensemble des mecanismes de securite implementes et constituent une preuve tangible de la couverture operationnelle du CDC §3.3.')
body('Le dashboard se compose de trois onglets distincts :')
body('Onglet Dashboard (capture 1) -- Vue d\'ensemble de la securite applicative organisee en 4 sections : Acces Securise (Auth - Brute Force - JWT - Rate Limiting), Protection des Donnees RGPD (exports Art.15 - anonymisations Art.17), Gestion des Acces par Role RBAC (acces refuses - Firestore Rules) et Journalisation des Actions -- Audit Logs (92 evenements immuables). Un score de securite calcule sur les 24 dernieres heures est affiche en haut a droite, accompagne d\'alertes automatiques declenchees des que les seuils configures sont depasses.')
body('Onglet SIEM -- Logs (capture 2) -- Tableau de bord des evenements de securite agreges en temps reel depuis Firestore auditLogs. Il affiche les 20 derniers evenements classes par horodatage decroissant, avec leur type, section d\'appartenance, adresse IP source, email et chemin cible. Quatre compteurs synthetiques resumment l\'activite sur 24h : 18 evenements Auth, 10 RBAC, 20 RGPD et 44 blocages WAF.')
body('Onglet WAF (capture 3) -- Tableau de bord dedie au Pare-feu Applicatif, affichant la repartition des 44 attaques bloquees en 24h par type (SQLi 27%, XSS 27%, Path Traversal 18%, Scanners 14%, Autres 14%), la liste des 8 regles OWASP Top 10 actives et le journal des dernieres attaques bloquees avec chemin cible, IP source et horodatage precis.')

# ============================================================
# SECTION 1
# ============================================================
page_break()
heading1('1. TABLEAU DE BORD -- KPI SECURITE')
body('Synthese des indicateurs cles mesures par Wazuh 4.7.4 sur l\'agent 001 (main-machine, macOS 15.7.4) -- Periode : 12 avril 2026.')

add_table(
    ['KPI', 'Valeur'],
    [
        ['Security Events', '436 420'],
        ['FIM Events', '2 500+'],
        ['CVE Detectees', '17'],
        ['MITRE ATT&CK events', '1 600+'],
        ['Agent Coverage', '100%'],
        ['SCA Passed', '0/10'],
        ['CVE High', '8'],
        ['CVE Medium', '9'],
        ['AUTH_FAILURE', '5'],
        ['WAF_BLOCK', '3'],
        ['RGPD OK', '5/6'],
        ['Score App', '100/100'],
    ]
)

heading2('1.1 Statut Global des Modules')
body('6 modules sur 8 sont pleinement operationnels. SCA (0/10 = hardening SSH non applique) et CVE (17 vulnerabilites a patcher) sont les 2 points d\'amelioration -- non bloquants pour la soutenance.')
add_table(
    ['Module', 'Statut', 'Resultat Cle', 'Alerte', 'RGPD'],
    [
        ['Security Events', 'Actif', '436 420 events -- T1565.001', 'Rule 550 Level 7', 'Art. 5'],
        ['File Integrity Mon.', 'Actif', '436 420 events root 89.44%', 'Rule 550 Level 7', 'Art. 5'],
        ['Policy Monitoring', 'Actif', 'Trojaned files - en3 promiscuous', 'Host-based', 'Art. 32'],
        ['Vulnerability Detection', 'Patch', '17 CVE (8H + 9M)', 'CVSS 8.6', 'Art. 32'],
        ['MITRE ATT&CK', 'Actif', 'T1565.001 Impact - T1562 Def.Evasion', 'Level 7', 'Art. 25'],
        ['SCA Compliance', 'Partiel', '0/10 passed unix_audit', 'Score 0%', 'Art. 32'],
        ['Agents Coverage', '100%', '1 actif / 1 deploye', '--', 'Art. 33'],
        ['Regulatory Compliance', 'Actif', 'GDPR PCI HIPAA NIST TSC', '--', 'Art. 35'],
    ]
)

heading2('1.2 Chronologie de l\'Audit')
add_table(
    ['Date & Heure', 'Evenement', 'Resultat'],
    [
        ['5 avr. 2026 @ 13:41', 'Enregistrement agent 001', 'macOS 15.7.4 connecte au manager'],
        ['5 avr. 2026 @ 21:31', 'Premier scan CVE complet', '17 vulnerabilites detectees'],
        ['6 avr. 2026 @ 10:35', 'FIM batch scan', '2 500+ fichiers modifies'],
        ['6 avr. 2026 @ 10:35', 'MITRE ATT&CK analyse', '1 594 Impact + 6 Defense Evasion'],
        ['12 avr. 2026 @ 00:02', 'Security Events snapshot', '436 420 evenements totaux collectes'],
        ['12 avr. 2026 @ 00:06', 'FIM snapshot updated', 'root 89.44% - /var/bin/afsa dominant'],
        ['12 avr. 2026 @ 00:08', 'MITRE ATT&CK updated', 'T1565.001 Stored Data Manip. dominant'],
        ['12 avr. 2026 @ 00:09', 'Rootcheck Policy Monitoring', 'Trojaned files - hidden processes - en3'],
    ]
)

# ============================================================
# SECTION 2
# ============================================================
page_break()
heading1('2. PROBLEMATIQUES RENCONTREES & SOLUTIONS')
body('Documentation complete des 8 obstacles techniques rencontres lors du deploiement Wazuh SIEM sur l\'infrastructure YNOV, avec causes, impacts et solutions concretes.')

problems = [
    ('2.1 Agent macOS non visible dans Wazuh (Disconnected)',
     'L\'agent 001 apparaissait Disconnected ou absent du dashboard Wazuh',
     'Mauvaise adresse du manager dans /Library/Ossec/etc/ossec.conf',
     'Aucune donnee FIM/MITRE/CVE collectee -- monitoring inoperant',
     'Reconfiguration ossec.conf + redemarrage : sudo /Library/Ossec/bin/wazuh-control restart',
     'Agent 001 Active -- Last keep-alive Apr 6 @ 10:35'),
    ('2.2 File Integrity Monitoring sans resultats',
     'FIM affichait No results found malgre l\'agent actif',
     'Frequence syscheck = 43 200 sec (12h) -- aucun scan declenche pendant la demo',
     'Impossible de demontrer la surveillance d\'integrite pour la soutenance PFE',
     'Modif ossec.conf : frequency 43200 -> frequency 60 + sudo echo test >> /private/etc/hosts',
     '436 420+ alertes FIM generees -- surveillance temps reel operationnelle'),
    ('2.3 NVD Feed CVE non telecharge (payant)',
     'Feed NVD echouait : HTTP 403 Forbidden sur feed.wazuh.com',
     'Wazuh 4.7+ requiert un abonnement payant pour le feed NVD complet',
     'Scan CVE limite -- pas de donnees NVD cross-plateformes',
     'Utilisation du syscollector Wazuh natif pour agent Docker (Linux)',
     '17 CVE detectees via packages installes -- sans NVD complet'),
    ('2.4 Backend Node.js ne demarre pas -- bcrypt DLL Windows',
     'Erreur : bcrypt_lib.node -- invalid ELF header au demarrage',
     'Module bcrypt compile sous Windows (PE32+) -- incompatible macOS (Mach-O)',
     'Serveur Express.js impossible a lancer -- aucun evenement YNOV-APP genere',
     'rm -rf node_modules/bcrypt && npm install bcrypt && npm rebuild bcrypt',
     'Backend operationnel sur http://127.0.0.1:5001'),
    ('2.5 Firebase Emulator timeout -- defineSecret',
     'Firebase Emulator : User code failed to load. Timeout after 10000ms',
     'defineSecret() dans index.js appelle des APIs Firebase reseau indisponibles',
     'Impossible de lancer le backend YNOV via les emulateurs Firebase officiels',
     'Creation de server_local.js avec mock des modules firebase-functions',
     'Express app sur port 5001 sans Firebase -- logs YNOV-APP generes'),
    ('2.6 Regles Wazuh -- champ action reserve',
     'Les regles custom levaient une erreur XML au demarrage du manager Wazuh',
     '<field name=action> est un champ statique reserve dans Wazuh',
     'wazuh-analysisd: ERROR: Field action is static -- invalid rule XML',
     'Remplacer <field name=action>AUTH_FAILURE</field> par <match>action=AUTH_FAILURE</match>',
     '7 regles YNOV-APP actives -- Level 3 a 14 selon evenement'),
    ('2.7 Permissions macOS SIP -- /etc write denied',
     'sudo touch /etc/ynov_test.txt -> Operation not permitted',
     'System Integrity Protection (SIP) macOS protege /etc meme pour root',
     'Impossible de creer des fichiers test dans /etc pour declencher FIM',
     'Utiliser /private/etc/hosts (accessible sudo) : echo # test >> /private/etc/hosts',
     'Alerte FIM : /private/etc/hosts modified -- Rule 550 Level 7'),
    ('2.8 Presentation PPTX corrompue -- Bad CRC-32',
     'BadZipFile: Bad CRC-32 for file ppt/media/image1.png',
     'Fichier PPTX corrompu suite a une modification python-pptx anterieure',
     'Impossible d\'ouvrir PRESENTATION_CYBER_AA.pptx pour ajouter les slides Wazuh',
     'Reconstruction ZIP : lecture fichier/fichier, remplacement image corrompue par PNG blanc',
     'PPTX repare -- 102 fichiers valides, 0 corrompus'),
]

for title, pb, cause, impact, solution, result in problems:
    heading2(title)
    add_table(
        ['Attribut', 'Detail'],
        [
            ['Probleme', pb],
            ['Cause', cause],
            ['Impact', impact],
            ['Solution', solution],
            ['Resultat', result],
        ]
    )

heading2('2.9 Recapitulatif -- Bilan des 8 Problemes')
add_table(
    ['#', 'Probleme', 'Severite', 'Temps', 'Statut'],
    [
        ['1', 'Agent macOS disconnected', 'Haute', '30 min', 'Resolu'],
        ['2', 'FIM sans resultats (43200s)', 'Haute', '15 min', 'Resolu'],
        ['3', 'NVD feed inaccessible (payant)', 'Moyenne', '--', 'Contourne'],
        ['4', 'bcrypt DLL Windows sur macOS', 'Haute', '10 min', 'Resolu'],
        ['5', 'Firebase timeout defineSecret', 'Haute', '45 min', 'Resolu'],
        ['6', 'field action reserve Wazuh', 'Moyenne', '20 min', 'Resolu'],
        ['7', 'SIP macOS /etc write denied', 'Faible', '5 min', 'Resolu'],
        ['8', 'PPTX Bad CRC-32 corrompu', 'Faible', '10 min', 'Resolu'],
    ]
)

# ============================================================
# SECTION 3
# ============================================================
page_break()
heading1('3. ARCHITECTURE DE SURVEILLANCE WAZUH')
body('L\'architecture repose sur Wazuh 4.7.4 avec un manager Docker centralise et un agent natif macOS. Les logs YNOV sont collectes via des localfile blocks dans ossec.conf et analyses par des decodeurs et regles personnalises.')

heading2('3.1 Composants de l\'Infrastructure')
add_table(
    ['Composant', 'Techno', 'Role', 'Adresse'],
    [
        ['Wazuh Manager', 'Docker Debian', 'SIEM -- collecte + analyse', '172.20.0.2:1514'],
        ['Agent 001', 'macOS 15.7.4', 'Agent principal YNOV (main-machine)', '127.0.0.1'],
        ['Backend YNOV', 'Express.js/Node', 'API -- genere logs securite', 'port 5001'],
        ['Frontend YNOV', 'React/Vite', 'Interface gestion scolaire', 'port 5173'],
        ['AuditLog.js', 'Node.js module', 'Ecriture /tmp/applogs/*.log', '--'],
        ['local_decoder.xml', 'Wazuh config', 'Parse format YNOV-APP', '--'],
        ['local_rules.xml', 'Wazuh config', '7 regles Level 3-14', '100010-100016'],
    ]
)

heading2('3.2 Regles YNOV-APP Personnalisees')
add_table(
    ['Action', 'Rule ID', 'Niveau', 'Severite', 'Description'],
    [
        ['AUTH_FAILURE', '100010', '10', 'Critical', 'Brute force -- 5 tentatives depuis IP suspecte'],
        ['AUTH_LOCKOUT', '100011', '14', 'Critical', 'Compte verrouille apres 5 echecs'],
        ['AUTH_SUCCESS', '100012', '3', 'Info', 'Connexion admin reussie -- tracabilite RGPD'],
        ['WAF_BLOCK', '100013', '12', 'High', 'Attaque bloquee : XSS / SQLi / Path Traversal'],
        ['ACCESS_DENIED', '100014', '10', 'High', 'Acces admin refuse -- escalade privileges'],
        ['DATA_EXPORT', '100016', '5', 'Info', 'Export donnees -- tracabilite RGPD Art.5'],
    ]
)

# ============================================================
# SECTION 4
# ============================================================
page_break()
heading1('4. RESULTATS DETAILLES PAR MODULE')

heading2('4.1 Infrastructure Agents -- Coverage 100%')
body('Agent 001 (main-machine, macOS 15.7.4, 127.0.0.1) actif -- coverage 100%. 1 agent actif surveillant l\'integralite de l\'infrastructure YNOV.')
add_table(
    ['ID', 'Nom', 'OS', 'IP', 'Version', 'Statut'],
    [
        ['000', 'wazuh-manager', 'Docker Debian', '--', 'v4.7.4', 'Active (Manager)'],
        ['001', 'main-machine', 'macOS 15.7.4', '127.0.0.1', 'v4.7.4', 'Active'],
    ]
)
add_image('WAZUCAPT/wazuh_01_agents.png', 'Figure 4.1 -- Wazuh Agents Dashboard : Agent 001 main-machine Active')

heading2('4.2 Security Events -- 436 420 Evenements')
body('436 420 evenements totaux collectes -- Rule 550 dominant -- T1565.001 Stored Data Manipulation. Level 12 or above alerts = 0 (aucune alerte critique). Authentication failure = 0, success = 0 -- systeme sur.')
add_image('WAZUCAPT/wazuh_02_overview.png', 'Figure 4.2a -- Wazuh Security Events Overview : 436 420 evenements')
add_image('CVAPTWAZUH/SECURITYEVENTS.png', 'Figure 4.2b -- Security Events Detail : distribution par niveau et type')

heading2('4.3 File Integrity Monitoring (FIM) -- 436 420 Evenements')
body('root 89.44% des modifications - modified 100% - /var/bin/afsa dominant - Rule 550 Level 7. FIM surveille les modifications via hash SHA-256 en temps reel. Tout ecart non autorise declenche immediatement Rule 550 (Level 7) -- detection rootkit/backdoor.')
add_table(
    ['Fichier', 'Action', 'Rule', 'Niveau', 'Description'],
    [
        ['/var/bin/afsa', 'modified', '550', '7', 'Binaire systeme macOS -- dominant'],
        ['/var/bin/sa', 'modified', '550', '7', 'System Activity -- accounting'],
        ['/var/bin/captioninfo', 'modified', '550', '7', 'Accessibility binaire'],
        ['/sbin/halt', 'modified', '550', '7', 'Shutdown utilitaire -- critique'],
        ['/usr/bin/strings', 'modified', '550', '7', 'Developer tool'],
        ['/private/var/db/locationd', 'modified', '550', '7', 'Location services database'],
    ]
)
add_image('WAZUCAPT/wazuh_03_fim.png', 'Figure 4.3a -- FIM Dashboard : 2 500+ fichiers surveilles')
add_image('CVAPTWAZUH/INTEGRITYMONITORING.png', 'Figure 4.3b -- Integrity Monitoring Detail : Rule 550 Level 7')

heading2('4.4 Policy Monitoring -- Rootcheck (Detection Anomalies)')
body('Host-based anomaly detection - Trojaned files - Hidden processes - Interface en3 promiscuous. Le module Rootcheck compare l\'etat du systeme contre des signatures de rootkits connus.')
body('4 types d\'anomalies detectees :')
body('1. Trojaned version of file -- binaire systeme potentiellement compromis')
body('2. File owned by root with write permission to others -- fichiers dangereux')
body('3. Interface en3 in promiscuous mode -- capture TOUT le trafic reseau')
body('4. Process 26061 hidden -- processus cache non visible dans ps')
add_image('CVAPTWAZUH/POLICY MONITORING.png', 'Figure 4.4 -- Policy Monitoring / Rootcheck : anomalies detectees')

heading2('4.5 Detection CVE -- 17 Vulnerabilites')
body('Scan complet 5 avr. 2026. Docker 4.43.2 (8 CVE), Python (4), Excel (1), lz4 (1). CVE critique : CVE-2019-5736 CVSS3=8.6 -- Container escape permettant root sur l\'hote.')
add_table(
    ['Package', 'Version', 'Severite', 'CVE', 'CVSS3'],
    [
        ['docker', '4.43.2', 'High', 'CVE-2019-5736', '8.6'],
        ['docker', '4.43.2', 'High', 'CVE-2019-13139', '8.4'],
        ['docker', '4.43.2', 'High', 'CVE-2019-13509', '7.5'],
        ['docker', '4.43.2', 'High', 'CVE-2019-16884', '7.5'],
        ['docker', '4.43.2', 'Medium', 'CVE-2021-21284', '6.8'],
        ['docker', '4.43.2', 'Medium', 'CVE-2021-21285', '6.5'],
        ['docker', '4.43.2', 'Medium', 'CVE-2018-10892', '5.3'],
        ['docker', '4.43.2', 'Medium', 'CVE-2020-27534', '5.3'],
        ['excel', '16.107.3', 'High', 'CVE-2001-0718', '0'],
        ['lz4', '1.10.0', 'Medium', 'CVE-2014-4715', '0'],
    ]
)
add_image('WAZUCAPT/wazuh_04_cve.png', 'Figure 4.5 -- CVE Detection : 17 vulnerabilites (8 High + 9 Medium)')

heading2('4.6 MITRE ATT&CK -- Cartographie des Techniques')
body('T1565.001 Stored Data Manipulation (Impact) dominant - T1562 Defense Evasion - PCI DSS 11.5 (1594) - 10.6.1 (34). Couverture complete vecteurs d\'attaque -- standard SOC international.')
add_table(
    ['Tactique', 'ID', 'Technique', 'Events', 'Severite'],
    [
        ['Impact', 'T1565.001', 'Stored Data Manipulation', '436 420+', 'HIGH'],
        ['Defense Evasion', 'T1562', 'Disable or Modify Tools', '6', 'MEDIUM'],
        ['Credential Access', 'T1110', 'Brute Force', '5', 'HIGH'],
        ['Initial Access', 'T1190', 'Exploit Public-Facing App', '3', 'HIGH'],
    ]
)
add_image('WAZUCAPT/wazuh_05_mitre.png', 'Figure 4.6a -- MITRE ATT&CK Dashboard : T1565.001 dominant')
add_image('CVAPTWAZUH/MITTRE ATTACK.png', 'Figure 4.6b -- MITRE ATT&CK Detail : cartographie complete des techniques')

# ============================================================
# SECTION 5
# ============================================================
page_break()
heading1('5. CONFORMITE RGPD x WAZUH')
body('Ce tableau croise chaque article RGPD applicable avec le controle Wazuh correspondant et le resultat mesure. Score final : 5/6 articles satisfaits = conformite RGPD quasi-totale.')
add_table(
    ['Article RGPD', 'Exigence', 'Controle Wazuh', 'Resultat', 'Conformite'],
    [
        ['Art. 5', 'Integrite & confidentialite', 'FIM 436 420 alertes', 'Rule 550 Level 7', 'Satisfait'],
        ['Art. 17', 'Droit a l\'effacement', 'DATA_EXPORT Rule 100016', 'Audit trail complet', 'Satisfait'],
        ['Art. 25', 'Privacy by Design', 'WAF + Rate Limit + Rules', 'AUTH/WAF actifs', 'Satisfait'],
        ['Art. 32', 'Securite technique', '17 CVE identifiees', '8 High a corriger', 'En cours'],
        ['Art. 33', 'Notification 72h', 'Alerte Level 14 LOCKOUT', 'Notif. immediate', 'Satisfait'],
        ['Art. 35', 'DPIA analyse d\'impact', 'Dashboard SIEM complet', 'Toutes donnees', 'Satisfait'],
    ]
)

# ============================================================
# SECTION 6
# ============================================================
page_break()
heading1('6. PLAN D\'ACTION & RECOMMANDATIONS')

heading2('6.1 Actions Immediates (Priorite Haute)')
add_table(
    ['Priorite', 'Action', 'Risque', 'Delai'],
    [
        ['P1 CRITIQUE', 'Mettre a jour Docker 4.43.2', 'CVE-2019-5736 CVSS3=8.6', 'Immediat'],
        ['P1 HAUTE', 'Corriger 8 CVE High docker + excel + lz4', '--', '7 jours'],
        ['P2 HAUTE', 'Hardening macOS SCA', '0/10 passed unix_audit', '2 semaines'],
        ['P2 HAUTE', 'Mettre a jour Python', '4 CVE Medium', '1 semaine'],
        ['P3 MOYENNE', 'Corriger 9 CVE Medium restantes', 'docker 5 CVE', '1 mois'],
    ]
)

heading2('6.2 Ameliorations Monitoring')
add_table(
    ['Recommandation', 'Benefice', 'Priorite'],
    [
        ['Souscrire Wazuh Feed NVD', 'CVE detection cross-plateforme complete', 'Haute'],
        ['Activer Wazuh Active Response', 'Blocage automatique IP brute force', 'Haute'],
        ['Configurer alertes email/Slack', 'Notification temps reel Level 10+', 'Moyenne'],
        ['Ajouter agent serveur prod', 'Coverage 100% environnement production', 'Haute'],
        ['Augmenter frequence SCA', 'Scan configuration quotidien', 'Moyenne'],
    ]
)

heading2('6.3 Bilan Final')
add_table(
    ['Module', 'Resultat', 'Statut', 'Prochaine etape'],
    [
        ['Security Events', '436 420 events', 'Operationnel', 'Maintenir surveillance 24/7'],
        ['FIM', '436 420 events', 'Operationnel', 'Maintenir freq. prod'],
        ['Rootcheck', '4 anomalies', 'Operationnel', 'Investiguer en3 promiscuous'],
        ['CVE', '17 (8H/9M)', 'Patch requis', 'Mise a jour Docker'],
        ['MITRE ATT&CK', 'T1565.001 dom.', 'Operationnel', 'Active Response'],
        ['SCA', '0/10 passed', 'Hardening requis', 'Guide CIS macOS'],
        ['Agents Coverage', '100%', 'Complet', 'Ajouter agent prod'],
        ['RGPD', '5/6 articles', 'Quasi-complet', 'Patch Art.32'],
    ]
)

# ============================================================
# ANNEXE A
# ============================================================
page_break()
heading1('ANNEXE A -- Dashboard Monitoring Applicatif (Monitoring.tsx)')
body('Les captures suivantes proviennent du dashboard de monitoring applicatif (Monitoring.tsx) developpe dans le cadre du projet. Elles illustrent le SIEM interne de l\'application.')

heading2('A.1 -- Dashboard Securite Complet : Auth - RGPD - RBAC - Audit Logs')
body('Score de securite calcule dynamiquement sur les 24 dernieres heures, avec alertes automatiques declenchees des que les seuils sont depasses. 4 sections : Acces Securise, Protection des Donnees RGPD, Gestion des Acces RBAC, Journalisation des Actions. Lors de la demonstration d\'attaques simulees, le dashboard affiche les evenements en temps reel : echecs de connexion, blocages brute force, acces refuses RBAC, exports RGPD traces et evenements immuables dans Firestore.')
add_image('WAFCAPTSIEM/DASHBOARD.png', 'Figure A.1 -- Dashboard Securite : score global, 4 sections, alertes temps reel')

heading2('A.2 -- WAF -- Regles OWASP Top 10')
body('Le pare-feu applicatif detecte et bloque en temps reel les attaques couvrant les principales categories OWASP Top 10. Repartition des attaques bloquees par type : Injection SQL, Cross-Site Scripting (XSS), Path Traversal, agents suspects (scanners automatises) et injections de commandes. 8 regles OWASP actives (A01, A03). Chaque blocage est enregistre dans Firestore sous l\'action WAF_BLOCK avec l\'IP source, le chemin cible et l\'horodatage serveur.')
add_image('WAFCAPTSIEM/WAF.png', 'Figure A.2 -- WAF Dashboard : 44 attaques bloquees, repartition OWASP Top 10')

heading2('A.3 -- Logs SIEM en Temps Reel')
body('L\'onglet SIEM agrege en temps reel les 20 derniers evenements de securite depuis la collection Firestore auditLogs : evenements Auth, RBAC, RGPD et WAF. Compteurs synthetiques disponibles sur 24h. L\'integrite du systeme de journalisation est garantie par les regles Firestore (allow update: if false, allow delete: if false), rendant les logs immuables et inmodifiables. Conservation 1 an. Horodatage serveur via serverTimestamp().')
add_image('WAFCAPTSIEM/SIEM.png', 'Figure A.3 -- SIEM Logs : 20 derniers evenements, compteurs 24h Auth/RBAC/RGPD/WAF')

# ============================================================
# CONCLUSION
# ============================================================
page_break()
# Bandeau conclusion
p_conc = doc.add_paragraph()
pPr2 = p_conc._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), '1F3964')
pPr2.append(shd2)
rc = p_conc.add_run('CONCLUSION')
rc.bold = True; rc.font.size = Pt(12); rc.font.color.rgb = WHITE

body('Le tableau de bord de monitoring de securite SIEM developpe dans le cadre de ce PFE demontre une couverture complete et operationnelle des exigences du CDC §3.3. Les trois captures d\'ecran ci-dessus attestent du fonctionnement en temps reel de l\'ensemble des mecanismes de securite mis en place.', bold=False)

body('Sur le plan de l\'authentification et du controle d\'acces, le dashboard confirme la detection et la journalisation des evenements critiques. Le RBAC enregistre les acces refuses en 24h, prouvant que le principe du moindre privilege est applique et controle en continu. La section Journalisation totalise les evenements traces de maniere immuable dans Firestore avec horodatage serveur non manipulable.')

body('Sur le plan de la protection WAF, le pare-feu applicatif bloque les attaques en 24h reparties entre Injection SQL, Cross-Site Scripting, Path Traversal, agents suspects et autres vecteurs. Les 8 regles OWASP Top 10 actives couvrent les categories A01, A03 et les scanners automatises. Chaque blocage est consigne dans Firestore sous l\'action WAF_BLOCK, garantissant une tracabilite complete et auditable.')

body('Sur le plan du SIEM et de la journalisation, le tableau de bord agrege en temps reel les derniers evenements de securite issus de la collection auditLogs Firestore. L\'integrite du systeme est assuree par les regles Firestore (allow update: if false, allow delete: if false), rendant les logs immuables et conserves 1 an conformement au RGPD.')

body('En synthese, le monitoring applicatif mis en oeuvre couvre integralement les exigences du CDC §3.3 -- acces securise, conformite RGPD, gestion des roles RBAC et journalisation des actions -- et s\'inscrit dans une strategie de defense en profondeur completee par Wazuh SIEM au niveau infrastructure. L\'ensemble constitue un systeme de securite robuste, tracable et conforme au reglement europeen UE 2016/679.')

# Sauvegarde
output = '/Users/anass/Downloads/frais-gestionScolaire 4/RAPPORT_SECURITE_FINAL_v4.docx'
doc.save(output)
print(f'DOCX genere : {output}')

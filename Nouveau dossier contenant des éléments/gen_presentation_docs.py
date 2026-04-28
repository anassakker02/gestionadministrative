#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère 2 documents basés sur la présentation PDF (33 slides) :
1. GUIDE_PRESENTATION_20SLIDES.docx — 20 slides condensés + script 10 min
2. COMMANDES_DEMO_LIVE.docx — Annexes commandes (slides 32-33)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/anass/Downloads/frais-gestionScolaire 4"

SCREENSHOTS = {
    "dashboard_score": "CAPMONITORINGSECU/brave_screenshot_localhost (2).png",
    "waf_onglet":      "CAPMONITORINGSECU/brave_screenshot_localhost (1).png",
    "siem_logs":       "CAPMONITORINGSECU/brave_screenshot_localhost.png",
    "rbac_audit":      "CAPMONITORINGSECU/brave_screenshot_localhost (3).png",
    "security_events": "CAPMONITORINGSECU/SECURITEEVENTS.png",
    "wazuh_agents":    "WAZUCAPT/wazuh_01_agents.png",
    "wazuh_fim":       "WAZUCAPT/wazuh_03_fim.png",
    "wazuh_cve":       "WAZUCAPT/wazuh_04_cve.png",
    "wazuh_mitre":     "WAZUCAPT/wazuh_05_mitre.png",
    "policy_monitor":  "CVAPTWAZUH/POLICY MONITORING.png",
}

# ─── HELPERS ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def bandeau(doc, texte, bg_hex='1A3A6E', fg='FFFFFF', size=13):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), bg_hex)
    pPr.append(shd)
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def h1(doc, texte):
    p = doc.add_heading(texte, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1A,0x3A,0x6E)
        r.font.size = Pt(15)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def h2(doc, texte, color=RGBColor(0x1A,0x3A,0x6E)):
    p = doc.add_heading(texte, level=2)
    for r in p.runs:
        r.font.color.rgb = color
        r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p

def body(doc, texte):
    p = doc.add_paragraph(texte)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    return p

def highlight(doc, texte, fill='EEF4FF', col='1A3A6E'):
    p = doc.add_paragraph()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run("→ " + texte)
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(int(col[0:2],16), int(col[2:4],16), int(col[4:6],16))
    p.paragraph_format.space_after = Pt(2)
    return p

def add_img(doc, key, caption, width=Cm(13)):
    rel = SCREENSHOTS.get(key, "")
    full = os.path.join(BASE, rel)
    if not os.path.exists(full):
        doc.add_paragraph(f"[Image: {rel}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(full, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    doc.add_paragraph()

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, '1A3A6E')
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_bg(c, 'F5F8FF')
    doc.add_paragraph()
    return t

def code_block(doc, texte):
    p = doc.add_paragraph()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), '0D1117')
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run(texte)
    r.font.name = 'Courier New'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x00,0xFF,0x7F)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    return p

def kpi_row(doc, items):
    colors = ['1A3A6E','C62828','2E7D32','E65100','6A1B9A','00695C','AD1457','283593']
    t = doc.add_table(rows=2, cols=len(items))
    t.style = 'Table Grid'
    for i, (nb, label) in enumerate(items):
        col = colors[i % len(colors)]
        for row_idx in range(2):
            set_cell_bg(t.rows[row_idx].cells[i], col)
        r1 = t.rows[0].cells[i].paragraphs[0].add_run(str(nb))
        r1.bold = True; r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        t.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = t.rows[1].cells[i].paragraphs[0].add_run(label)
        r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        t.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def slide_card(doc, num, titre, resume, note, color='1A3A6E'):
    """Bloc visuel pour chaque slide"""
    # Numéro + titre
    bandeau(doc, f"SLIDE {num:02d} — {titre}", bg_hex=color, size=11)
    # Résumé
    p = doc.add_paragraph()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F0F4FF')
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run("RÉSUMÉ : " + resume)
    r.font.size = Pt(10); r.bold = True
    r.font.color.rgb = RGBColor(0x1A,0x3A,0x6E)
    p.paragraph_format.space_after = Pt(1)
    # Note orateur
    if note:
        pn = doc.add_paragraph()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear')
        shd2.set(qn('w:fill'), 'FFF8E1')
        pn._p.get_or_add_pPr().append(shd2)
        rn = pn.add_run("À DIRE : " + note)
        rn.font.size = Pt(9); rn.italic = True
        rn.font.color.rgb = RGBColor(0x5D,0x40,0x00)
        pn.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 : GUIDE_PRESENTATION_20SLIDES.docx
# ═══════════════════════════════════════════════════════════════════════════

def gen_guide_presentation():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # COUVERTURE
    bandeau(doc, "GUIDE PRÉSENTATION — 20 SLIDES — 10 MIN — 24 AVRIL 2026 — 10H30→10H40", bg_hex='1A3A6E', size=12)
    bandeau(doc, "Anass Akker / Amine BAHOU — YNOV Campus — PFE Cybersécurité / Cyberdéfense", bg_hex='C62828', size=10)

    kpi_row(doc, [
        ("10 CAT.", "OWASP Top 10"),
        ("6 ART.", "RGPD"),
        ("12 TESTS", "DAST 92/100"),
        ("436 420", "Events Wazuh"),
        ("17 CVE", "8H+9M"),
        ("100/100", "Score App"),
    ])

    doc.add_page_break()

    # ── PARTIE 1 : RÉSUMÉS 33 SLIDES ORIGINAUX ──
    bandeau(doc, "PARTIE 1 — RÉSUMÉS DES 31 SLIDES ORIGINAUX + 2 ANNEXES = 33 PAGES PDF (1-2 LIGNES PAR SLIDE)", bg_hex='1A3A6E')

    slides_originaux = [
        (1, "Cover — Contexte & KPI", "1A3A6E",
         "Application gestion scolaire Firebase/React sécurisée. Stack : React 18 + Node.js + Firestore + JWT HS256. SIEM Wazuh 4.7.4.",
         "On présente le projet et ses 6 KPI clés : OWASP, RGPD, DAST, FIM, CVE, MITRE."),
        (2, "Plan — 30 Slides", "283593",
         "3 parties : ① Sécurité Applicative (OWASP/RBAC/RGPD/JWT/DAST) ② Monitoring Applicatif (Dashboard 5 captures) ③ Infrastructure Wazuh SIEM.",
         "Annonce la structure : sécurité → monitoring → infrastructure."),
        (3, "Contexte du Projet", "2E7D32",
         "Stack : React 18 + TypeScript (front), Node.js + Express (back), Firestore NoSQL, JWT HS256, Firebase Hosting + Azure. SIEM Wazuh Docker.",
         "Données sensibles élèves/paiements → RGPD obligatoire. Moindre privilège RBAC."),
        (4, "Audit OWASP Top 10", "E65100",
         "10 catégories vérifiées : 8 CORRIGÉES, 1 SURVEILLÉE (CVE Wazuh), 1 N/A (SSRF). Toutes failles critiques couvertes.",
         "A01 à A10 : RBAC, AES-256, WAF, JWT, Helmet, Wazuh CVE, Rate Limit, AuditLogs, SIEM."),
        (5, "Vulnérabilités & Corrections", "C62828",
         "6 vulnérabilités résolues : SQLi CRITIQUE (WAF 403), XSS HAUTE (WAF+CSP), Brute force HAUTE (lockout 5), Escalade HAUTE (checkRole), JWT MOYENNE (exp 24h), Données clair HAUTE (AES-256-CBC).",
         "WAF = première ligne de défense appliqué avant l'authentification."),
        (6, "RBAC — Moindre Privilège", "6A1B9A",
         "4 rôles : Admin (complet), Sous-admin (opérationnel), Comptable (finance), Étudiant/Parent (lecture seule). checkRole() re-lit Firestore à CHAQUE requête.",
         "Jamais depuis JWT uniquement — vérification temps réel = blocage immédiat si rôle révoqué."),
        (7, "Conformité RGPD — 6 Articles", "1A3A6E",
         "Art.15 DATA_EXPORT, Art.16 CRUD loggué, Art.17 DATA_ANONYMIZE, Art.25 AES+HTTPS+RBAC, Art.32 WAF+JWT+bcrypt, Art.33 procédure 72h. Tous ✓.",
         "Dashboard garantit conformité. Privacy by Design dès la conception."),
        (8, "Audit Logs — 9 Types Immuables", "2E7D32",
         "9 types : AUTH_SUCCESS/FAILURE/LOCKOUT, LOGOUT, SESSION_EXPIRED, ACCESS_DENIED, DATA_EXPORT, DATA_ANONYMIZE, WAF_BLOCK. Firestore : allow update/delete = if false.",
         "Logs infalsifiables. serverTimestamp() côté serveur. Append-only."),
        (9, "Security Headers · JWT · AES-256", "E65100",
         "Helmet.js : HSTS 1an, X-Frame DENY, CSP default-src 'self'. JWT HS256 exp 24h. AES-256-CBC IV 128bits. bcrypt 12 rounds. Secrets en .env jamais en code.",
         "Anti-timing attack + anti-rainbow table. JWT_SECRET en variable d'environnement."),
        (10, "Rate Limiting & Protection DoS", "C62828",
         "Global 100 req/15min → 429. Auth login 10 req/15min. Brute force 5 échecs → AUTH_LOCKOUT. Payload max 10KB → 413. Timeout 28s → 503.",
         "En production : Redis partagera le compteur entre instances Cloud Functions."),
        (11, "Architecture Sécurité Pipeline", "6A1B9A",
         "Pipeline 8 couches : WAF → Rate Limit → verifyJWT → checkRole() → Route Handler → auditLogger → Dashboard /monitoring → Wazuh SIEM.",
         "Chaque requête filtrée par 4 couches avant d'accéder aux données."),
        (12, "Tests DAST — 12 Tests OWASP — 92/100", "1A3A6E",
         "12 tests automatisés : T01-T11 PASS (API, Headers, Auth, Rate, SQLi, XSS, Path, Agents, Payload, RBAC, /monitoring). T12 warning CORS.",
         "node scripts/security_scan.js → Score 92/100. Intégration CI/CD possible."),
        (13, "Partie 1 — Monitoring Applicatif", "2E7D32",
         "Transition : 5 captures dashboard /monitoring · WAF · auditLogs · SIEM · Score /100 · CDC §3.3.",
         "10 étapes tech, 9 audit logs, 12 tests DAST, score /100."),
        (14, "Dashboard Score 100/100", "E65100",
         "Score 100/100 calculé sur 24h : authFailures + lockouts + accessDenied + wafBlocks. JWT forcé, bcrypt 10, rate limit, timeout 30min, anti-énumération.",
         "Alertes auto : >20 auth failures → CRITIQUE, >5 lockouts → CRITIQUE, >10 WAF blocks → CRITIQUE."),
        (15, "Dashboard WAF OWASP Top 10", "C62828",
         "0 attaques bloquées (24h = système sûr). 8 règles OWASP actives : A03 SQLi, A03 XSS, A01 LFI, Cmd Injection, Scanners (sqlmap/nikto). WAF_BLOCK → Firestore.",
         "Le WAF détecte et bloque en temps réel. Chaque attaque tracée avec IP/route/méthode."),
        (16, "SIEM Logs Temps Réel", "6A1B9A",
         "Tableau 20 derniers events Firestore : IP, email, rôle, chemin, action, horodatage. Admin only. 0 events = surveillance normale. Logs immuables (allow update/delete: if false).",
         "Boîtier noir infalsifiable — personne ne peut modifier ces preuves, même pas un admin."),
        (17, "Dashboard RGPD Conformité", "1A3A6E",
         "Art.15 GET /users/:id/export → DATA_EXPORT loggué. Art.17 DELETE /users/:id/data → DATA_ANONYMIZE. AES-256-CBC + HTTPS/HSTS 1an. Procédure 72h documentée.",
         "Permet au DPO de vérifier conformité sans accès technique."),
        (18, "Dashboard RBAC & Audit Logs", "2E7D32",
         "Deny by default : Firestore Rules + ProtectedRoute /unauthorized. 4 rôles. allow update/delete: if false. 9 types events immuables.",
         "RBAC bloque tout accès par défaut et surveille tentatives en temps réel."),
        (19, "Partie 2 — Wazuh SIEM Infrastructure", "E65100",
         "Transition : 213 événements <24h, 3 conteneurs Docker, 5 modules actifs, coverage 100%.",
         "Monitoring Infrastructure · MITRE ATT&CK · GDPR · CIS Benchmark."),
        (20, "Architecture Wazuh Docker — 4 Commandes", "C62828",
         "6 composants : App Web, API Backend, Wazuh Agent (OSSEC port 1514 AES), Manager, Indexer (OpenSearch port 9200), Dashboard (Kibana fork port 443). Installation en 4 commandes git clone + docker compose.",
         "3 conteneurs Docker. Communications sécurisées TLS. Single-node suffisant pour PFE."),
        (21, "Agents Wazuh — 2 Actifs — Coverage 100%", "6A1B9A",
         "Agent 001 : frais-gestion-scolaire Debian 12 (172.20.0.5). Agent 002 : main-machine macOS 15.7.4 (127.0.0.1). v4.7.4. OSSEC port 1514 AES. 0 disconnected.",
         "Coverage 100% = tous les composants surveillés en temps réel."),
        (22, "Security Events — 436 420 Events", "1A3A6E",
         "436 420 events journalisés. Rule 550 Level 7 dominant. T1565.001 Stored Data Manipulation. Level 12+ = 0 (sûr). Auth failure/success = 0.",
         "Vue la plus exhaustive du SIEM. Corrélation MITRE automatique."),
        (23, "FIM — 436 420 Events — root 89.44%", "2E7D32",
         "Rule 550 Level 7. root 89.44% des modifications. 100% 'modified'. /var/bin/afsa dominant. SHA-256 temps réel. Risque : rootkit/backdoor si modifications non légitimes.",
         "FIM = IDS filesystem. Tout écart non autorisé → alerte immédiate."),
        (24, "CVE — 17 Vulnérabilités", "C62828",
         "0 Critical, 8 High, 9 Medium. Docker 4.43.2 (8 CVE dont CVE-2019-5736 CVSS3=8.6 container escape). Python (4), Excel (1), lz4 (1). Action P1 : mise à jour Docker.",
         "Vulnérability Detector compare packages vs NVD NIST automatiquement."),
        (25, "MITRE ATT&CK — Cartographie", "6A1B9A",
         "T1565.001 Stored Data Manipulation (Impact) dominant (~95% alertes). T1562 Disable or Modify Tools (Defense Evasion) = 6 events. Corrélation PCI DSS 11.5 × 1594. RGPD Art.25 Privacy by Design ✓.",
         "Framework universel SOC. Wazuh mappe alertes → techniques ATT&CK automatiquement."),
        (26, "Rootcheck — Policy Monitoring — Anomalies Système", "AD1457",
         "Trojaned files : détection binaires compromis. Hidden processes : processus 26061 caché. Interface en3 promiscuous = capture réseau détectée. Host-based Anomaly Detection actif temps réel.",
         "Rootcheck détecte rootkits, backdoors et sniffing réseau. 4 anomalies identifiées sur l'agent macOS."),
        (27, "5 Modules Wazuh Actifs — CDC §3.3", "1A3A6E",
         "Security Events (journalisation temps réel), MITRE ATT&CK (T1499×1594/T1562×6), GDPR module (Art.15/17/33), FIM Hash SHA-256 (2500+ events), Vulnerabilities (17 CVE/CVSS).",
         "Les 5 modules assurent surveillance complète : menaces + conformité + vulnérabilités."),
        (28, "Defense in Depth — Monitoring 2 Niveaux — CDC §3.3 100%", "2E7D32",
         "Couche 1 APP : WAF+9AuditLogs+Score/100+DAST 92/100+Dashboard 3 onglets+RGPD. Couche 2 INFRA : Agent Wazuh+213 events+CIS Benchmark+MITRE+GDPR+17CVE. CDC §3.3 : 5 exigences couvertes.",
         "Deux couches complémentaires = couverture totale applicatif ET infrastructure."),
        (29, "Synthèse CDC §3.3 — 8/8 Exigences FAITES", "E65100",
         "Accès sécurisé ✓, HTTPS ✓, RGPD conforme ✓, RBAC ✓, Journalisation ✓, Détection vulnérabilités (Wazuh) ✓, Intégrité fichiers (FIM) ✓, Détection menaces (MITRE) ✓. Toutes vérifiables en live.",
         "Chaque exigence a une implémentation technique vérifiable et démontrable devant le jury."),
        (30, "Bilan Global & Perspectives Post-Soutenance", "C62828",
         "Réalisé : WAF 5 vecteurs, 9 logs immuables, Dashboard 100/100, DAST 92/100, Wazuh 213 events, Coverage 100%, CDC §3.3 100%. Perspectives : Azure, Alertes Twilio, CI/CD GitHub Actions, Grafana, Pentest ZAP.",
         "Maîtrise complète sécurité applicative + infrastructure. Prochaines étapes maturité."),
        (31, "Merci — Questions & Démo Live", "1A3A6E",
         "Conclusion : Defense in Depth CDC §3.3 à 100% sur 2 niveaux : applicatif (WAF+9logs+SIEM+DAST) + infrastructure Wazuh (436 420 events+MITRE+GDPR+CIS). Conforme RGPD.",
         "Phrase exacte à prononcer : 'Le monitoring couvre intégralement le CDC §3.3 sur deux niveaux...'"),
        (32, "[ANNEXE] Commandes Sécurité par catégorie", "0D1117",
         "7 catégories de commandes : WAF (SQLi/XSS → 403), Auth JWT brute force (lockout 429), DAST scanner (92/100), Démarrage app, Wazuh Docker, Agent macOS, Chiffrement AES-256, RGPD Art.15/17. → COMMANDES_DEMO_LIVE.docx",
         "Slide extrait dans fichier séparé COMMANDES_DEMO_LIVE.docx"),
        (33, "[ANNEXE] Démo Live — 7 Étapes — 8 minutes — Guide Jury", "1A3A6E",
         "7 étapes chronométrées : ÉTAPE01 WAF SQLi (1min), ÉTAPE02 WAF XSS (30s), ÉTAPE03 Brute Force (1min), ÉTAPE04 DAST (1min), ÉTAPE05 Wazuh (2min), ÉTAPE06 Dashboard (1min), ÉTAPE07 RGPD (1min). Total 8 min. → COMMANDES_DEMO_LIVE.docx",
         "Slide extrait dans fichier séparé COMMANDES_DEMO_LIVE.docx"),
    ]

    for num, titre, col, resume, note in slides_originaux:
        slide_card(doc, num, titre, resume, note, color=col)

    doc.add_page_break()

    # ── PARTIE 2 : VERSION 20 SLIDES CONDENSÉE ──
    bandeau(doc, "PARTIE 2 — VERSION CONDENSÉE 20 SLIDES — 10 MINUTES (30 SEC PAR SLIDE) — SLIDES 32-33 EXTRAITS", bg_hex='C62828')
    body(doc, "Plan adapté pour 10h30→10h40 : 31 slides originaux condensés en 20. Slides 32-33 (Commandes + Démo Live) → fichier séparé COMMANDES_DEMO_LIVE.docx. Structure : Sécurité App (7) · Monitoring (5) · Wazuh SIEM (6) · Conclusion (2).")
    doc.add_paragraph()

    slides_20 = [
        # ── PARTIE 1 : SÉCURITÉ APP ──
        ("S01", "Cover + Contexte", "1A3A6E", "30s",
         "App gestion scolaire YNOV — données sensibles élèves/paiements → RGPD. Stack : React 18 + Node.js + Firebase + JWT HS256. SIEM Wazuh 4.7.4.",
         "Bonjour. Je présente la sécurisation d'une application de gestion scolaire Firebase/React. Les données sont sensibles — RGPD obligatoire. J'ai implémenté deux couches de sécurité.",
         None),
        ("S02", "Audit OWASP Top 10 — 10 catégories vérifiées", "2E7D32", "30s",
         "10 catégories : 8 CORRIGÉES (A01 RBAC, A02 AES-256, A03 WAF, A04 JWT, A05 Helmet, A07 Rate Limit, A08 AuditLogs, A09 SIEM), 1 SURVEILLÉE (A06 CVE Wazuh), 1 N/A (A10 SSRF).",
         "Sur les 10 catégories OWASP, 8 sont corrigées, 1 surveillée par Wazuh, 1 non applicable.",
         None),
        ("S03", "WAF + Vulnérabilités Corrigées", "E65100", "30s",
         "WAF waf.js : SQLi SELECT/UNION/DROP → 403, XSS <script> → 403, Path Traversal ../../ → 403, CMD Injection → 403, Scanners sqlmap/nikto → 403. Première ligne de défense avant l'auth.",
         "Le WAF est appliqué avant même l'authentification. Toutes les vulnérabilités critiques et hautes sont corrigées.",
         "waf_onglet"),
        ("S04", "RBAC — 4 Rôles — Moindre Privilège", "6A1B9A", "30s",
         "Admin (complet), Sous-admin (opérationnel), Comptable (finance), Étudiant (lecture seule). checkRole() re-lit Firestore à CHAQUE requête — jamais depuis JWT seul.",
         "La vérification des rôles est en temps réel. Si un rôle est révoqué, l'accès est bloqué immédiatement à la prochaine requête.",
         None),
        ("S05", "RGPD — 6 Articles Implémentés", "AD1457", "30s",
         "Art.15 DATA_EXPORT loggué, Art.16 CRUD tracé, Art.17 DATA_ANONYMIZE, Art.25 AES-256+HTTPS+RBAC Privacy by Design, Art.32 WAF+JWT+bcrypt 12r, Art.33 procédure 72h. Tous ✓.",
         "6 articles RGPD couverts. AES-256-CBC pour données sensibles. Procédure 72h documentée.",
         None),
        ("S06", "9 Audit Logs Immuables + Security Headers", "283593", "30s",
         "9 types events Firestore : AUTH_SUCCESS/FAILURE/LOCKOUT, LOGOUT, SESSION_EXPIRED, ACCESS_DENIED, DATA_EXPORT, DATA_ANONYMIZE, WAF_BLOCK. allow update/delete: if false. Helmet HSTS+CSP+X-Frame.",
         "Logs infalsifiables — personne ne peut les modifier, même un admin. serverTimestamp() côté serveur.",
         None),
        ("S07", "DAST Scanner — 92/100 — 12 Tests OWASP", "00695C", "30s",
         "node scripts/security_scan.js → 11/12 PASS : API, HSTS, Auth 401, Rate 429, SQLi 403, XSS 403, Path 403, Agents 403, Payload 413, RBAC, /monitoring admin. 1 warning CORS.",
         "Scanner automatisé qui teste toute la sécurité OWASP. Intégrable CI/CD. Score 92/100.",
         None),

        # ── PARTIE 2 : MONITORING APPLICATIF ──
        ("S08", "[PARTIE 2] Dashboard Score 100/100", "1A3A6E", "30s",
         "Dashboard /monitoring : Score 100/100 sur 24h. JWT forcé, brute force 5 essais → blocage, rate limit 10 req/15min, bcrypt 10 rounds, timeout 30min, anti-énumération.",
         "Voici le dashboard de monitoring applicatif. Score 100/100 en ce moment — aucun incident.",
         "dashboard_score"),
        ("S09", "Dashboard WAF + SIEM Logs", "2E7D32", "30s",
         "WAF : 0 attaques (24h). 8 règles OWASP actives. SIEM Logs : tableau 20 derniers events Firestore (IP+email+rôle+chemin+action+horodatage). Admin only. 0 events = normal.",
         "Le WAF est actif mais aucune attaque détectée — système sûr. Les logs SIEM sont le boîtier noir infalsifiable.",
         "siem_logs"),
        ("S10", "Dashboard RGPD + RBAC", "E65100", "30s",
         "RGPD : AES-256-CBC ✓, HTTPS/HSTS ✓, Art.15/16/17/33 tous couverts. RBAC : Firestore deny-by-default, 6 rôles, allow update/delete if false, 9 types AuditLogs, Conservation 1 an.",
         "Le DPO peut vérifier la conformité RGPD sans accès technique. RBAC bloque par défaut.",
         "rbac_audit"),
        ("S11", "Architecture Pipeline de Sécurité", "C62828", "30s",
         "Pipeline 8 couches en série : WAF → Rate Limit → verifyJWT → checkRole(Firestore) → Route Handler → auditLogger → Dashboard → Wazuh SIEM. Chaque requête filtrée avant les données.",
         "8 couches de sécurité. Une requête malveillante est bloquée dès le WAF sans jamais toucher la base de données.",
         None),
        ("S12", "[Transition] Wazuh SIEM Infrastructure", "6A1B9A", "15s",
         "Monitoring Infrastructure · MITRE ATT&CK · GDPR · CIS Benchmark. 3 conteneurs Docker. 5 modules actifs. Coverage 100%.",
         "Passons maintenant à la deuxième couche : l'infrastructure Wazuh SIEM.",
         None),

        # ── PARTIE 3 : WAZUH SIEM ──
        ("S13", "Architecture Wazuh + Agents Coverage 100%", "1A3A6E", "30s",
         "Manager Docker Debian (172.20.0.2:1514) + Agent macOS main-machine (127.0.0.1). Installation 4 commandes Docker. v4.7.4. OSSEC AES. Coverage 100% — 0 disconnected.",
         "Wazuh s'installe en 4 commandes Docker. L'agent macOS envoie des événements en temps réel.",
         "wazuh_agents"),
        ("S14", "Security Events — 436 420 + FIM SHA-256", "2E7D32", "30s",
         "436 420 events. Rule 550 Level 7 domine à 95%. T1565.001 Stored Data Manipulation. FIM : root 89.44%, /var/bin/afsa dominant, SHA-256 temps réel. Tout écart → alerte immédiate.",
         "436 420 événements surveillés. FIM détecterait tout rootkit ou backdoor en temps réel.",
         "wazuh_fim"),
        ("S15", "CVE — 17 Vulnérabilités — Patch P1", "C62828", "30s",
         "17 CVE : 8 High + 9 Medium. CVE-2019-5736 CVSS3=8.6 (Docker container escape → root sur l'hôte). Action P1 immédiate post-soutenance : docker update. Python 4 CVE, Excel 1, lz4 1.",
         "La CVE la plus critique est CVE-2019-5736 sur Docker. Patch P1 immédiat post-soutenance. La détection démontre l'efficacité du SIEM.",
         "wazuh_cve"),
        ("S16", "MITRE ATT&CK + Rootcheck — Slide 25-26 PDF", "6A1B9A", "40s",
         "MITRE : T1565.001 Stored Data Manipulation = 95% des alertes (Impact). T1562 Defense Evasion = 6 events. Rootcheck : 4 anomalies détectées — Trojaned files, en3 promiscuous (capture réseau), processus 26061 caché. PCI DSS 11.5 ✓.",
         "Wazuh mappe automatiquement sur MITRE ATT&CK — framework SOC international. Rootcheck a détecté de l'anomalie système réelle sur ma machine — interface promiscuous = sniffing réseau possible.",
         "wazuh_mitre"),
        ("S17", "5 Modules Wazuh + Defense in Depth", "AD1457", "30s",
         "5 modules : Security Events, MITRE ATT&CK, GDPR (Art.15/17/33), FIM SHA-256, CVE/CVSS. Defense in Depth = 2 couches : Applicatif (WAF+9logs+DAST) + Infrastructure (Wazuh). CDC §3.3 100%.",
         "Les 5 modules combinés couvrent toutes les exigences CDC §3.3. Defense in Depth = deux couches complémentaires.",
         None),
        ("S18", "Synthèse CDC §3.3 — 8/8 Exigences FAITES", "283593", "30s",
         "Accès sécurisé ✓, HTTPS ✓, RGPD ✓, RBAC ✓, Journalisation ✓, Détection vulnérabilités ✓, Intégrité fichiers ✓, Détection menaces ✓. Chaque exigence vérifiable en live.",
         "CDC §3.3 couvert à 100%. Chaque point est implémenté, mesuré et démontrable devant le jury.",
         None),

        # ── CONCLUSION ──
        ("S19", "Bilan + Perspectives Post-Soutenance", "00695C", "30s",
         "Réalisé : WAF 5 vecteurs OWASP, 9 logs immuables, Dashboard 100/100, DAST 92/100, Wazuh 436 420 events, Coverage 100%. À faire : Azure, CI/CD, Alertes email/SMS, Pentest ZAP, Grafana.",
         "Le projet démontre une maîtrise complète de la sécurité applicative et infrastructure.",
         None),
        ("S20", "Merci — Questions & Démonstration Live", "1A3A6E", "30s",
         "Le monitoring couvre le CDC §3.3 à 100% sur 2 niveaux : applicatif (WAF·9logs·SIEM·DAST) et infrastructure Wazuh (436 420 events·MITRE·GDPR·CIS). Defense in Depth conforme RGPD.",
         "Merci. Je suis disponible pour vos questions et une démonstration live si vous le souhaitez.",
         None),
    ]

    bandeau(doc, "PLAN 20 SLIDES — STRUCTURE OPTIMISÉE 10 MINUTES", bg_hex='2E7D32')
    add_table(doc, ["Slide","Titre","Durée"], [
        (s[0], s[1], s[3]) for s in slides_20
    ])

    doc.add_page_break()

    bandeau(doc, "DÉTAIL DES 20 SLIDES — RÉSUMÉS + TEXTE À PRONONCER", bg_hex='1A3A6E')

    for sid, titre, col, duree, resume, a_dire, img_key in slides_20:
        # En-tête slide
        p = doc.add_paragraph()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),col)
        p._p.get_or_add_pPr().append(shd)
        r = p.add_run(f"[{sid}] {duree} — {titre}")
        r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p.paragraph_format.space_after = Pt(1)

        # Résumé
        p2 = doc.add_paragraph()
        shd2 = OxmlElement('w:shd'); shd2.set(qn('w:val'),'clear'); shd2.set(qn('w:fill'),'EEF4FF')
        p2._p.get_or_add_pPr().append(shd2)
        r2 = p2.add_run("CONTENU : " + resume)
        r2.font.size = Pt(9); r2.bold = True
        r2.font.color.rgb = RGBColor(0x1A,0x3A,0x6E)
        p2.paragraph_format.space_after = Pt(1)

        # À dire
        p3 = doc.add_paragraph()
        shd3 = OxmlElement('w:shd'); shd3.set(qn('w:val'),'clear'); shd3.set(qn('w:fill'),'FFF8E1')
        p3._p.get_or_add_pPr().append(shd3)
        r3 = p3.add_run("À DIRE : " + a_dire)
        r3.italic = True; r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0x5D,0x40,0x00)
        p3.paragraph_format.space_after = Pt(4)

        # Screenshot si applicable
        if img_key:
            add_img(doc, img_key, f"Capture — {titre}", width=Cm(12))

    doc.add_page_break()

    # ── PARTIE 3 : SCRIPT ORAL CONDENSÉ ──
    bandeau(doc, "PARTIE 3 — SCRIPT ORAL EXACT — 10 MINUTES CHRONO", bg_hex='283593')

    scripts = [
        ("0:00–1:00", "1A3A6E", "Introduction + OWASP + WAF",
         "Bonjour. Je présente la sécurisation d'une application de gestion scolaire Firebase/React traitant des données sensibles — RGPD obligatoire.\n\n"
         "Sur les 10 catégories OWASP, 8 sont corrigées. La première ligne de défense est le WAF waf.js qui bloque SQLi, XSS, Path Traversal et Command Injection avec HTTP 403 avant même l'authentification."),
        ("1:00–2:00", "2E7D32", "RBAC + RGPD + AuditLogs",
         "Le RBAC implémente 4 rôles avec le principe du moindre privilège. checkRole() re-lit Firestore à CHAQUE requête — jamais depuis le JWT seul — blocage immédiat si rôle révoqué.\n\n"
         "6 articles RGPD couverts : AES-256-CBC pour données sensibles, Art.15 export loggué, Art.17 anonymisation, Art.33 procédure 72h. 9 types d'audit logs immuables — allow update/delete: if false."),
        ("2:00–3:00", "E65100", "DAST + Architecture Pipeline",
         "Le scanner DAST automatise 12 tests OWASP : score 92/100, 11/12 réussis. Intégrable CI/CD.\n\n"
         "Le pipeline de sécurité en 8 couches : WAF → Rate Limit → verifyJWT → checkRole(Firestore) → Route Handler → auditLogger → Dashboard → Wazuh SIEM. Chaque requête filtrée avant les données."),
        ("3:00–4:30", "6A1B9A", "Dashboard Monitoring — Score 100/100",
         "Le dashboard /monitoring affiche un score 100/100. JWT forcé, brute force bloqué après 5 tentatives, rate limiting 10 req/15min, bcrypt 10 rounds.\n\n"
         "Onglet WAF : 0 attaques détectées — système sûr. Onglet SIEM Logs : tableau temps réel des 20 derniers events Firestore — boîtier noir infalsifiable. Onglet RGPD : tous les articles couverts."),
        ("4:30–5:00", "C62828", "Transition vers Wazuh",
         "Passons à la deuxième couche de sécurité : l'infrastructure Wazuh SIEM 4.7.4. Déployé en 4 commandes Docker. L'agent macOS main-machine est actif avec coverage 100%."),
        ("5:00–6:30", "1A3A6E", "Wazuh — Agents + Security Events + FIM",
         "436 420 événements collectés. Rule 550 Level 7 domine à 95% — T1565.001 Stored Data Manipulation.\n\n"
         "Le module FIM surveille les fichiers système via SHA-256 en temps réel. root représente 89.44% des modifications sur /var/bin/*. Tout rootkit ou backdoor serait détecté immédiatement."),
        ("6:30–7:30", "2E7D32", "CVE + MITRE ATT&CK + Rootcheck",
         "17 CVE détectées : 8 High, 9 Medium. La plus critique : CVE-2019-5736 CVSS3=8.6 sur Docker — container escape permettant root sur l'hôte. Patch P1 post-soutenance.\n\n"
         "MITRE ATT&CK : T1565.001 Impact dominant. Rootcheck : 4 anomalies dont interface en3 promiscuous et processus 26061 caché."),
        ("7:30–8:30", "E65100", "Defense in Depth + CDC §3.3 100%",
         "Defense in Depth = 2 couches complémentaires :\n"
         "Couche 1 Applicatif : WAF + 9 AuditLogs immuables + Score 100/100 + DAST 92/100 + RGPD Art.15/16/17/33\n"
         "Couche 2 Infrastructure : Wazuh FIM + CVE + MITRE ATT&CK + SCA + Rootcheck\n\n"
         "CDC §3.3 couvert à 100% : 8 exigences, 8 implémentées, toutes vérifiables en live."),
        ("8:30–9:30", "6A1B9A", "Bilan + Conclusion",
         "Points forts réalisés : WAF maison 5 vecteurs OWASP, 9 logs immuables Firestore, Dashboard SIEM 100/100, Scanner DAST 92/100, Wazuh 436 420 events, Coverage 100%, CDC §3.3 100%.\n\n"
         "Perspectives : déploiement Azure Container Apps, alertes email/SMS Twilio, CI/CD GitHub Actions, Wazuh agent production, Pentest OWASP ZAP."),
        ("9:30–10:00", "C62828", "Phrase de conclusion + Questions",
         "Le monitoring couvre intégralement le CDC §3.3 sur deux niveaux :\n"
         "monitoring applicatif (WAF · 9 audit logs Firestore · dashboard SIEM · DAST 12 tests)\n"
         "et monitoring infrastructure Wazuh (436 420 événements · MITRE ATT&CK · GDPR · CIS Benchmark).\n"
         "Ensemble, ils assurent une défense en profondeur conforme RGPD.\n\n"
         "Merci. Je suis disponible pour vos questions et une démonstration live."),
    ]

    for time_str, col, titre, texte in scripts:
        bandeau(doc, f"[{time_str}] {titre}", bg_hex=col, size=10)
        for line in texte.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line)
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    # KPI à mémoriser
    doc.add_page_break()
    bandeau(doc, "CHIFFRES CLÉS À MÉMORISER — NE PAS OUBLIER", bg_hex='1A3A6E')
    add_table(doc, ["Valeur","Signification","Slide"], [
        ("436 420","Security Events Wazuh totaux","S14"),
        ("100/100","Score Dashboard Applicatif","S08"),
        ("92/100","Score DAST Scanner (12 tests)","S07"),
        ("17 CVE","8 High + 9 Medium (Docker dominant)","S15"),
        ("CVE-2019-5736","CVSS3=8.6 Docker container escape P1","S15"),
        ("T1565.001","MITRE — Stored Data Manipulation (~95%)","S16"),
        ("Rule 550 Level 7","FIM — Integrity checksum changed","S14"),
        ("89.44%","FIM root dominant /var/bin/*","S14"),
        ("5/6 ou 6/6","Articles RGPD satisfaits","S05"),
        ("9","Types AuditLogs immuables Firestore","S06"),
        ("7","Règles YNOV-APP Wazuh (100010-100016)","S13"),
        ("bcrypt 12 rounds","Hachage mdp anti-brute force","S09"),
        ("AES-256-CBC","Chiffrement données sensibles","S05"),
        ("5 tentatives","Rate Limit → AUTH_LOCKOUT 15min","S06"),
        ("CDC §3.3 100%","8/8 exigences couvertes","S18"),
        ("Defense in Depth","2 couches : App + Wazuh SIEM","S17"),
    ])

    out = os.path.join(BASE, "GUIDE_PRESENTATION_20SLIDES.docx")
    doc.save(out)
    print(f"✅ GUIDE_PRESENTATION_20SLIDES.docx → {out}")
    return out


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 : COMMANDES_DEMO_LIVE.docx (slides 32-33)
# ═══════════════════════════════════════════════════════════════════════════

def gen_commandes_demo():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    bandeau(doc, "COMMANDES DÉMO LIVE — SLIDES 32 & 33 — ANNEXES TECHNIQUES", bg_hex='0D1117', fg='00FF7F')
    bandeau(doc, "Pré-requis : Backend :5001 + Frontend :5173 + Wazuh Docker :443 + Token JWT prêt", bg_hex='C62828', size=10)

    doc.add_paragraph()
    kpi_row(doc, [
        ("T1=:5001","Backend node server_local.js"),
        ("T2=:5173","Frontend npm run dev"),
        ("T3=:443","Wazuh docker compose up -d"),
        ("8 min","Durée totale démo"),
        ("7 étapes","WAF+Rate+DAST+Wazuh+Monitoring+RGPD"),
        ("3 terminaux","Ouverts avant de commencer"),
    ])

    h1(doc, "SLIDE 32 — TOUTES LES COMMANDES PAR CATÉGORIE")

    # WAF
    h2(doc, "WAF — Test Blocage (A03 OWASP)", RGBColor(0xC6,0x28,0x28))
    code_block(doc, "# SQLi — doit retourner 403 WAF_BLOCK\ncurl -s -o /dev/null -w \"HTTP: %{http_code}\\n\" \\\n  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?id=1%20OR%201=1'\n# → HTTP: 403")
    code_block(doc, "# XSS — doit retourner 403 WAF_BLOCK\ncurl -s -o /dev/null -w \"HTTP: %{http_code}\\n\" \\\n  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?nom=<script>alert(1)</script>'\n# → HTTP: 403")

    # Auth JWT Brute Force
    h2(doc, "Auth JWT — Brute Force → Lockout (A07 OWASP)", RGBColor(0xE6,0x51,0x00))
    code_block(doc, "# Brute force → lockout après 5 essais → tentative 6 = 429\nfor i in {1..6}; do\n  curl -s -o /dev/null -w \"Tentative $i → %{http_code}\\n\" \\\n    -X POST http://localhost:5001/gestionadminastration/us-central1/api/auth/login \\\n    -H 'Content-Type: application/json' \\\n    -d '{\"email\":\"admin@school.fr\",\"password\":\"WRONG\"}'\ndone\n# → Tentative 1→401  2→401  3→401  4→401  5→401  6→429 AUTH_LOCKOUT")

    # DAST Scanner
    h2(doc, "DAST Scanner Automatique — 12 Tests OWASP", RGBColor(0x2E,0x7D,0x32))
    code_block(doc, "# Lancer le scan de vulnérabilités OWASP Top 10\ncd \"/Users/anass/Downloads/frais-gestionScolaire 4/back/functions\"\nnode scripts/security_scan.js\n# Résultat attendu :\n# [PASS] WAF actif — SQLi bloqué\n# [PASS] Rate limiting — 429 après 5 req\n# [PASS] JWT validation — 401 token invalide\n# [PASS] Headers sécurité — HSTS présent\n# [WARN] 17 CVE détectées (audit npm)\n# Score DAST final : 92/100")

    # Démarrage App
    h2(doc, "Démarrage Application (2 terminaux)", RGBColor(0x1A,0x3A,0x6E))
    code_block(doc, "# Terminal 1 — Backend\ncd \"/Users/anass/Downloads/frais-gestionScolaire 4/back/functions\"\nnode server_local.js\n# → Express running on http://localhost:5001")
    code_block(doc, "# Terminal 2 — Frontend\ncd \"/Users/anass/Downloads/frais-gestionScolaire 4/front\"\nnpm run dev\n# → http://localhost:5173")

    # Wazuh Docker
    h2(doc, "Wazuh — Démarrage Docker (3 conteneurs)", RGBColor(0x6A,0x1B,0x9A))
    code_block(doc, "# Symlink sans espace dans le chemin (1 seule fois)\nln -s \"/Users/anass/Downloads/frais-gestionScolaire 4/front/wazuh-docker/single-node\" /tmp/wazuh-node\ncd /tmp/wazuh-node\n\n# Démarrer les 3 conteneurs (manager + indexer + dashboard)\ndocker compose up -d\n\n# Vérifier (attendre ~2 min)\ndocker compose ps   # → 3 conteneurs Up\n\n# Accès dashboard\nopen https://localhost   # → admin / SecretPassword")

    # Wazuh Agent
    h2(doc, "Wazuh — Agent main-machine (si déconnecté)", RGBColor(0xAD,0x14,0x57))
    code_block(doc, "# 1. Arrêter l'agent\nsudo /Library/Ossec/bin/wazuh-control stop\n\n# 2. Pointer vers manager localhost\nsudo sed -i \"\" 's|<address>.*</address>|<address>127.0.0.1</address>|' \\\n  /Library/Ossec/etc/ossec.conf\n\n# 3. Vérifier l'adresse\nsudo grep \"<address>\" /Library/Ossec/etc/ossec.conf\n\n# 4. Re-enregistrer main-machine\nsudo /Library/Ossec/bin/agent-auth -m 127.0.0.1 -A main-machine\n\n# 5. Redémarrer + statut\nsudo /Library/Ossec/bin/wazuh-control start\nsudo /Library/Ossec/bin/wazuh-control status\n# → https://localhost → main-machine Active")

    # FIM - Déclencher alerte
    h2(doc, "FIM — Déclencher Alerte Rule 550 Level 7", RGBColor(0x2E,0x7D,0x32))
    code_block(doc, "# Modifier /private/etc/hosts → déclenche FIM Rule 550 Level 7\nsudo bash -c \"echo '# wazuh-test' >> /private/etc/hosts\"\n# → Wazuh détecte : Integrity checksum changed\n# → Rule 550, Level 7, T1565.001 Stored Data Manipulation")

    # Chiffrement AES
    h2(doc, "Chiffrement AES-256-CBC — Vérification", RGBColor(0x00,0x69,0x5C))
    code_block(doc, "# Test chiffrement (Node.js REPL)\ncd \"/Users/anass/Downloads/frais-gestionScolaire 4/back/functions\"\nnode -e \"\nconst enc = require('./src/utils/encryption');\nconst c = enc.encrypt('donnee-sensible');\nconsole.log('Chiffré:', c);\nconsole.log('Déchiffré:', enc.decrypt(c));\"\n\n# Vérifier headers sécurité\ncurl -I http://localhost:5001/api/auth/login\n# → Strict-Transport-Security: max-age=31536000")

    # RGPD
    h2(doc, "RGPD — Routes Art.15 et Art.17", RGBColor(0x28,0x3C,0x6E))
    code_block(doc, "# Obtenir un token JWT d'abord\nTOKEN=$(curl -s -X POST http://localhost:5001/api/auth/login \\\n  -H 'Content-Type: application/json' \\\n  -d '{\"email\":\"admin@school.fr\",\"password\":\"MOT_DE_PASSE\"}' \\\n  | python3 -c \"import sys,json; print(json.load(sys.stdin)['token'])\")\n\n# Art. 15 — Export données personnelles\ncurl -H \"Authorization: Bearer $TOKEN\" \\\n  http://localhost:5001/api/users/USER_ID/export\n# → DATA_EXPORT loggué dans auditLogs Firestore\n\n# Art. 17 — Anonymisation\ncurl -X DELETE -H \"Authorization: Bearer $TOKEN\" \\\n  http://localhost:5001/api/users/USER_ID/data\n# → DATA_ANONYMIZE loggué dans auditLogs Firestore")

    doc.add_page_break()

    # ── SLIDE 33 : GUIDE JURY DÉMO LIVE ──
    h1(doc, "SLIDE 33 — DÉMONSTRATION LIVE — GUIDE JURY — 7 ÉTAPES — 8 MINUTES")
    bandeau(doc, "AVANT DE COMMENCER : T1=Backend :5001 · T2=Frontend :5173 · T3=Wazuh Docker :443 · Token JWT dans presse-papier", bg_hex='C62828', size=9)

    etapes = [
        ("ÉTAPE 01", "WAF — Injection SQL bloquée", "1A3A6E", "1 min",
         "curl -s -o /dev/null -w \"Réponse HTTP: %{http_code}\\n\" \\\n  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?id=1%20OR%201=1'\n# → Réponse HTTP: 403",
         "J'essaie une injection SQL dans l'URL. Le WAF middleware intercepte la requête avant qu'elle touche la base de données → 403 Forbidden. OWASP A03."),
        ("ÉTAPE 02", "WAF — XSS bloqué", "C62828", "30s",
         "curl -s -o /dev/null -w \"Réponse HTTP: %{http_code}\\n\" \\\n  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?nom=<script>alert(1)</script>'\n# → Réponse HTTP: 403",
         "Même chose pour le XSS — la balise script est bloquée immédiatement. OWASP A03."),
        ("ÉTAPE 03", "Brute Force — Rate Limiting (OWASP A07)", "E65100", "1 min",
         "for i in {1..6}; do\n  curl -s -o /dev/null -w \"Tentative $i → %{http_code}\\n\" \\\n    -X POST http://localhost:5001/gestionadminastration/us-central1/api/auth/login \\\n    -H 'Content-Type: application/json' \\\n    -d '{\"email\":\"admin@school.fr\",\"password\":\"WRONG\"}'\ndone\n# → Tentative 1→401 2→401 3→401 4→401 5→401 6→429",
         "5 tentatives échouées = compte verrouillé 15 minutes. Protection brute force — OWASP A07."),
        ("ÉTAPE 04", "Scan DAST automatique — OWASP Top 10", "2E7D32", "1 min",
         "node \"/Users/anass/Downloads/frais-gestionScolaire 4/back/functions/scripts/security_scan.js\"\n# [PASS] WAF · JWT · Headers · Rate limiting · CORS\n# [WARN] 17 CVE npm\n# Score DAST final : 92/100",
         "Ce script automatise tous les tests OWASP Top 10. Tout passe sauf 17 CVE dans les dépendances npm — détectées mais non critiques pour la soutenance."),
        ("ÉTAPE 05", "Wazuh SIEM — Dashboard en direct", "6A1B9A", "2 min",
         "open https://localhost\n# → admin / SecretPassword\n# → Security Events · File Integrity · MITRE ATT&CK · Policy Monitoring\n# → 436 420 events · FIM 89% · T1565 Stored Data · Policy anomalies",
         "436 420 événements collectés. FIM surveille les fichiers système. Corrélation MITRE ATT&CK automatique. Vérification continue de la config OS."),
        ("ÉTAPE 06", "Dashboard Monitoring intégré", "AD1457", "1 min",
         "open http://localhost:5173\n# → Aller dans le menu : Monitoring de Sécurité\n# → Score 100/100 · Alertes temps réel · SIEM Logs · WAF stats",
         "Ce dashboard est intégré dans l'application — l'admin voit les alertes de sécurité directement depuis l'interface, sans aller sur Wazuh."),
        ("ÉTAPE 07", "RGPD — Traçabilité Art.5 & Art.32", "00695C", "1 min",
         "# 1. Onglet privé → http://localhost:5173/login\n# → Taper un MAUVAIS mot de passe → connexion échouée\n# 2. http://localhost:5173/monitoring → SIEM Logs → Rafraîchir\n# → Voir auth_failure avec timestamp + IP",
         "Je simule une tentative de connexion échouée — l'événement auth_failure est immédiatement tracé dans Firestore avec timestamp, IP, identité. Conformité RGPD Art.5 et Art.32."),
    ]

    for code, titre, col, duree, cmd, commentaire in etapes:
        bandeau(doc, f"{code} — {titre}  [{duree}]", bg_hex=col, size=10)
        code_block(doc, cmd)
        pn = doc.add_paragraph()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'FFF8E1')
        pn._p.get_or_add_pPr().append(shd)
        r = pn.add_run("COMMENTAIRE : " + commentaire)
        r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x5D,0x40,0x00)
        pn.paragraph_format.space_after = Pt(8)

    # Checklist pré-soutenance
    doc.add_page_break()
    bandeau(doc, "CHECKLIST PRÉ-SOUTENANCE — À FAIRE AVANT 10H30", bg_hex='C62828')
    checklist = [
        "[ ] Backend démarré : cd back/functions && node server_local.js → port 5001",
        "[ ] Frontend démarré : cd front && npm run dev → port 5173",
        "[ ] Wazuh Docker : cd /tmp/wazuh-node && docker compose up -d → attendre 2 min",
        "[ ] Ouvrir https://localhost → vérifier agent main-machine Active",
        "[ ] Ouvrir http://localhost:5173 → vérifier Dashboard Score 100/100",
        "[ ] Token JWT admin dans le presse-papier (pour démos RGPD)",
        "[ ] 3 terminaux ouverts et prêts",
        "[ ] Ce document ouvert en référence pendant la soutenance",
    ]
    for item in checklist:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)

    out = os.path.join(BASE, "COMMANDES_DEMO_LIVE.docx")
    doc.save(out)
    print(f"✅ COMMANDES_DEMO_LIVE.docx → {out}")
    return out


# ─── MAIN ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("Génération des 2 documents présentation...")
    print("=" * 60)

    f1 = gen_guide_presentation()
    f2 = gen_commandes_demo()

    print("\n" + "=" * 60)
    print("✅ DOCUMENTS GÉNÉRÉS :")
    for f in [f1, f2]:
        size = os.path.getsize(f) // 1024
        print(f"  📄 {os.path.basename(f)} ({size} Ko)")
    print("=" * 60)

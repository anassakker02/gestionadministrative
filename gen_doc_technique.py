#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Documentation Technique — Gestion Scolaire YNOV 2026"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = "/Users/anass/Downloads/frais-gestionScolaire 4"
doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = section.right_margin = Inches(1.0)
section.top_margin  = section.bottom_margin = Inches(1.0)

NAVY  = RGBColor(0x0D, 0x1B, 0x4B)
BLUE  = RGBColor(0x14, 0x5D, 0xA7)
CYAN  = RGBColor(0x1E, 0x88, 0xE5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
        run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def para(text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Calibri'
    if color: run.font.color.rgb = color
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00,0x40,0x80)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF4FF')
    pPr.append(shd)

def info_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run("ℹ  " + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'DCE8F8')
    pPr.append(shd)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    row0 = t.rows[0]
    for i,h in enumerate(headers):
        c = row0.cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True; r.font.size = Pt(10); r.font.name='Calibri'
        r.font.color.rgb = WHITE
        set_cell_bg(c, '145DA7')
        if widths: c.width = Inches(widths[i])
    for i,row_data in enumerate(rows):
        row = t.rows[i+1]
        bg = 'EEF4FF' if i%2==0 else 'FFFFFF'
        for j,txt in enumerate(row_data):
            c = row.cells[j]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(9.5); r.font.name='Calibri'
            set_cell_bg(c, bg)
    doc.add_paragraph()
    return t

def try_img(path, w=Inches(5.5), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.italic = True
            cp.runs[0].font.color.rgb = BLUE

def sep():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6'); bot.set(qn('w:color'),'145DA7')
    pBdr.append(bot); pPr.append(pBdr)

# ══ PAGE DE TITRE ══════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
r = p.add_run("DOCUMENTATION TECHNIQUE")
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY; r.font.name='Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Application Web — Gestion Scolaire YNOV Campus")
r2.font.size = Pt(16); r2.font.color.rgb = BLUE; r2.font.name='Calibri'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Installation · Déploiement · Sécurité")
r3.font.size = Pt(13); r3.font.italic=True; r3.font.color.rgb=CYAN; r3.font.name='Calibri'

doc.add_paragraph(); sep(); doc.add_paragraph()

add_table(["Champ","Valeur"],[
    ["Projet","Application Gestion Scolaire YNOV Campus"],
    ["Version","1.0.0 — Phase 3 Post-Production"],
    ["Date","Avril 2026"],
    ["Auteurs","Amine BAHOU / Anass Akker"],
    ["Formation","PFE Bachelor Cybersécurité / Cyberdéfense — YNOV Campus"],
    ["Stack","React + Node.js + Firebase/Firestore + Docker"],
    ["URL Production","https://frais-gestionscolaire.web.app"],
    ["Standards","OWASP Top 10 · RGPD UE 2016/679 · CDC §3.3 · Wazuh 4.7.4"],
],widths=[2.0,4.3])

doc.add_page_break()

# ══ SOMMAIRE ══════════════════════════════════════════════
heading("SOMMAIRE",1)
for num,title in [
    ("1.","Architecture Globale du Système"),
    ("2.","Prérequis et Environnement"),
    ("3.","Installation — Backend (Node.js)"),
    ("4.","Installation — Frontend (React)"),
    ("5.","Configuration Firebase / Firestore"),
    ("6.","Déploiement en Production"),
    ("7.","Déploiement Docker + Wazuh SIEM"),
    ("8.","Architecture de Sécurité — 8 Couches"),
    ("9.","WAF — Pare-feu Applicatif (waf.js)"),
    ("10.","RBAC — Gestion des Rôles et Accès"),
    ("11.","Chiffrement & Cryptographie"),
    ("12.","Conformité RGPD (UE 2016/679)"),
    ("13.","Journalisation & Audit Logs Immuables"),
    ("14.","Monitoring — Dashboard /monitoring"),
    ("15.","Wazuh SIEM — Configuration & Résultats"),
    ("16.","Variables d'Environnement"),
    ("17.","Rollback & Procédures d'Urgence"),
    ("18.","Synthèse CDC §3.3 — 8/8 Exigences"),
    ("19.","Annexes"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{num}  "); r1.font.bold=True; r1.font.color.rgb=BLUE; r1.font.size=Pt(11); r1.font.name='Calibri'
    r2 = p.add_run(title); r2.font.size=Pt(11); r2.font.name='Calibri'

doc.add_page_break()

# ══ 1. ARCHITECTURE ═══════════════════════════════════════
heading("1. Architecture Globale du Système",1)
para("L'application Gestion Scolaire YNOV est une SPA (Single Page Application) fullstack déployée en production sur Firebase, avec un backend Node.js sécurisé et un pipeline de sécurité à 8 couches défensives.")

heading("1.1 Stack Technique",2)
add_table(["Composant","Technologie","Version","Rôle"],[
    ["Frontend","React + TypeScript","18.x","SPA — Interface utilisateur"],
    ["Build tool","Vite","5.x","Compilation + bundle optimisé"],
    ["UI Library","Tailwind CSS + shadcn/ui","—","Design system responsive"],
    ["Backend","Node.js + Express","18.x","API REST sécurisée"],
    ["Base de données","Firebase Firestore","v9","NoSQL temps réel"],
    ["Authentification","JWT HS256 + bcrypt","—","Auth stateless sécurisée"],
    ["Stockage fichiers","Firebase Storage","—","Documents et PDFs"],
    ["Hosting","Firebase Hosting","—","CDN + HTTPS automatique"],
    ["WAF","waf.js (custom)","1.0","Pare-feu applicatif OWASP"],
    ["SIEM","Wazuh","4.7.4","Monitoring infrastructure"],
    ["Conteneurisation","Docker + Compose","—","Environnement reproductible"],
],widths=[1.4,1.7,0.8,2.4])

heading("1.2 Pipeline de Sécurité — 8 Couches",2)
add_table(["#","Couche","Rôle","Réponse si violation"],[
    ["1","WAF (waf.js)","SQLi · XSS · PathTraversal · CmdInj · Scanners","HTTP 403 + WAF_BLOCK loggué"],
    ["2","Rate Limiter","100 req/15min global · 5 échecs → lockout","HTTP 429 Too Many Requests"],
    ["3","verifyJWT","JWT HS256 · exp 24h · signature","HTTP 401 Unauthorized"],
    ["4","checkRole() Firestore","Re-lecture rôle BDD à chaque requête","HTTP 403 + ACCESS_DENIED"],
    ["5","Route Handler","Traitement métier sécurisé","Réponse applicative"],
    ["6","auditLogger","9 types events · Firestore immuable","Log serverTimestamp()"],
    ["7","Dashboard /monitoring","Score /100 · alertes · SIEM · 60s refresh","Alerte admin"],
    ["8","Wazuh SIEM","CIS · MITRE ATT&CK · FIM · CVE · GDPR","Alerte SIEM"],
],widths=[0.3,1.6,3.0,1.7])

doc.add_page_break()

# ══ 2. PRÉREQUIS ══════════════════════════════════════════
heading("2. Prérequis et Environnement",1)
add_table(["Outil","Version minimale","Vérification"],[
    ["Node.js","v18.x LTS","node --version"],
    ["npm","v9.x","npm --version"],
    ["Firebase CLI","v13.x","firebase --version"],
    ["Docker","v24.x","docker --version"],
    ["Docker Compose","v2.x","docker compose version"],
    ["Git","v2.x","git --version"],
    ["Python","v3.9+ (Wazuh)","python3 --version"],
],widths=[1.5,1.5,3.3])

heading("2.1 Variables d'Environnement Requises",2)
code_block("""# back/functions/.env  (NE JAMAIS COMMITTER)
JWT_SECRET=<secret_256_bits_minimum>           # openssl rand -hex 32
ENCRYPTION_KEY=<cle_aes_256_hex_64_chars>      # openssl rand -hex 32
FIREBASE_PROJECT_ID=frais-gestionscolaire
FIREBASE_SERVICE_ACCOUNT_PATH=./serviceAccountKey.json
NODE_ENV=production
PORT=3001
ALLOWED_ORIGINS=https://frais-gestionscolaire.web.app

# front/.env.production
VITE_API_BASE_URL=https://frais-gestionscolaire.web.app/api
VITE_FIREBASE_API_KEY=<firebase_api_key>
VITE_FIREBASE_AUTH_DOMAIN=frais-gestionscolaire.firebaseapp.com
VITE_FIREBASE_PROJECT_ID=frais-gestionscolaire""")

doc.add_page_break()

# ══ 3. INSTALLATION BACKEND ═══════════════════════════════
heading("3. Installation — Backend (Node.js)",1)

heading("3.1 Cloner et installer",2)
code_block("""git clone <url-repo> frais-gestionScolaire
cd frais-gestionScolaire/back/functions
npm install

# Vérifier les packages de sécurité installés
npm list helmet express-rate-limit jsonwebtoken bcrypt cors""")

heading("3.2 Dépendances de Sécurité Critiques",2)
add_table(["Package","Version","Rôle Sécurité"],[
    ["helmet","^7.x","Headers HTTP sécurité (CSP, HSTS, X-Frame-Options...)"],
    ["express-rate-limit","^7.x","Rate limiting — protection brute force"],
    ["jsonwebtoken","^9.x","JWT HS256 — authentification stateless"],
    ["bcrypt","^5.x","Hash mots de passe — 12 rounds"],
    ["cors","^2.x","CORS strict — domaines autorisés uniquement"],
    ["firebase-admin","^12.x","SDK Firebase — Firestore + Auth + Storage"],
    ["express","^4.x","Framework HTTP — routes API REST"],
],widths=[1.8,1.0,3.5])

heading("3.3 Démarrer le Serveur",2)
code_block("""# Développement (nodemon — rechargement auto)
cd back/functions && npm run dev
# → API sur http://localhost:3001

# Production
NODE_ENV=production node server_local.js

# Vérifier que le serveur répond
curl http://localhost:3001/health
# → {"status":"ok","version":"1.0.0","uptime":...}""")

heading("3.4 Routes API — Référence Complète",2)
add_table(["Route","Méthodes","Auth requise","Description"],[
    ["/api/v1/auth/login","POST","Non","Connexion — retourne JWT"],
    ["/api/v1/auth/register","POST","Non","Inscription (admin uniquement en prod)"],
    ["/api/v1/etudiants","GET/POST","JWT+RBAC","Liste et création étudiants"],
    ["/api/v1/etudiants/:id","PUT/DELETE","JWT+Admin","Modification/Suppression"],
    ["/api/v1/paiements","GET/POST","JWT+RBAC","Gestion paiements"],
    ["/api/v1/factures","GET/POST","JWT+RBAC","Génération factures PDF"],
    ["/api/v1/classes","GET/POST","JWT+RBAC","Gestion classes"],
    ["/api/v1/bourses","GET/POST","JWT+RBAC","Gestion bourses"],
    ["/api/v1/tarifs","GET/POST","JWT+RBAC","Gestion tarifs"],
    ["/api/v1/users/:id/export","GET","JWT+Admin/Soi","Export RGPD Art.15"],
    ["/api/v1/users/:id/data","DELETE","JWT+Admin","Anonymisation Art.17"],
    ["/api/v1/monitoring","GET","JWT+Admin","Dashboard SIEM sécurité"],
    ["/api/v1/backup","POST","JWT+Admin","Sauvegarde Firestore"],
    ["/api/v1/upload","POST","JWT+RBAC","Upload documents (PDF/images)"],
],widths=[2.0,1.2,1.0,2.1])

doc.add_page_break()

# ══ 4. FRONTEND ═══════════════════════════════════════════
heading("4. Installation — Frontend (React)",1)
code_block("""cd front
npm install

# Développement
npm run dev          # → http://localhost:5173

# Build production
npm run build        # → dist/ (1.08 MB / 309 KB gzip)

# Prévisualiser
npm run preview""")

heading("4.1 Structure des Répertoires Frontend",2)
add_table(["Dossier/Fichier","Description"],[
    ["src/api/","Appels HTTP vers l'API backend (axios)"],
    ["src/components/","Composants React réutilisables"],
    ["src/pages/","Pages de l'application (routing React Router)"],
    ["src/pages/Monitoring.tsx","Dashboard sécurité /monitoring (admin)"],
    ["src/services/","Services métier (factureService, studentService...)"],
    ["src/contexts/AuthContext.tsx","Gestion session JWT — login/logout/refresh"],
    ["src/utils/sanitize.ts","Sanitisation XSS — DOMPurify"],
    ["src/components/ProtectedRoute.tsx","Guard de route — vérifie JWT + rôle"],
    ["src/locales/","Traductions i18n (fr/en)"],
    ["src/types/","Types TypeScript (student, facture, user...)"],
],widths=[2.5,3.8])

doc.add_page_break()

# ══ 5. FIREBASE ═══════════════════════════════════════════
heading("5. Configuration Firebase / Firestore",1)

heading("5.1 Initialisation Firebase",2)
code_block("""npm install -g firebase-tools
firebase login
firebase init   # Sélectionner: Firestore, Hosting, Functions, Storage

# Déployer les règles de sécurité Firestore
firebase deploy --only firestore:rules

# Déployer les index de performance
firebase deploy --only firestore:indexes""")

heading("5.2 Règles Firestore — Sécurité RBAC",2)
info_box("Les règles Firestore constituent une deuxième couche RBAC côté base de données, indépendante du middleware backend.")
code_block("""// firestore.rules
rules_version = '2';
service cloud.firestore {
  match /databases/{database}/documents {

    // AUDIT LOGS — IMMUABLES (append-only OBLIGATOIRE)
    match /auditLogs/{logId} {
      allow read:   if isAdmin();
      allow create: if isAuthenticated();
      allow update: if false;   // ← JAMAIS modifiable
      allow delete: if false;   // ← JAMAIS supprimable
    }

    // Étudiants
    match /etudiants/{etudiantId} {
      allow read:   if isAdmin() || isSousAdmin() || isOwner(etudiantId);
      allow create: if isAdmin() || isSousAdmin();
      allow update: if isAdmin() || isSousAdmin();
      allow delete: if isAdmin();
    }

    // Paiements
    match /paiements/{paiementId} {
      allow read:   if isAdmin() || isSousAdmin() || isComptable();
      allow create: if isAdmin() || isSousAdmin() || isComptable();
      allow update: if isAdmin();
      allow delete: if isAdmin();
    }

    // Helpers
    function isAuthenticated() { return request.auth != null; }
    function isAdmin()    { return getUserRole() == 'admin'; }
    function isSousAdmin(){ return getUserRole() == 'sous-admin'; }
    function isComptable(){ return getUserRole() == 'comptable'; }
    function getUserRole() {
      return get(/databases/$(database)/documents/users/$(request.auth.uid)).data.role;
    }
  }
}""")

doc.add_page_break()

# ══ 6. DÉPLOIEMENT PRODUCTION ═════════════════════════════
heading("6. Déploiement en Production",1)

heading("6.1 Procédure Complète",2)
code_block("""# 1. Build frontend optimisé
cd front && npm run build
# → dist/ créé — 1.08 MB (309 KB gzip)

# 2. Deploy complet Firebase
cd .. && firebase deploy --only hosting,firestore:rules,firestore:indexes
# → https://frais-gestionscolaire.web.app ✅

# 3. Vérifier les headers de sécurité
curl -sI https://frais-gestionscolaire.web.app""")

heading("6.2 Headers de Sécurité en Production",2)
code_block("""# Résultat curl -sI https://frais-gestionscolaire.web.app

strict-transport-security: max-age=31556926; includeSubDomains; preload  ✅
content-security-policy: default-src 'self'; script-src 'self'; ...      ✅
x-frame-options: DENY                                                      ✅
x-content-type-options: nosniff                                           ✅
referrer-policy: strict-origin-when-cross-origin                          ✅
permissions-policy: camera=(), microphone=(), geolocation=()              ✅
x-xss-protection: 1; mode=block                                           ✅""")

heading("6.3 Checklist Déploiement",2)
add_table(["#","Vérification","Statut"],[
    ["1","Site accessible HTTPS — curl 200 OK","✅ PASS"],
    ["2","Redirect HTTP→HTTPS — 301 automatique","✅ PASS"],
    ["3","7 Headers sécurité présents (Helmet)","✅ PASS"],
    ["4","Firestore Rules déployées et actives","✅ PASS"],
    ["5","Firestore Indexes déployés (6 index)","✅ PASS"],
    ["6","Variables env production correctes","✅ PASS"],
    ["7","Bundle optimisé 1.08 MB / 309 KB gzip","✅ PASS"],
    ["8","Cache assets statiques max-age=31536000","✅ PASS"],
],widths=[0.3,4.7,1.2])

doc.add_page_break()

# ══ 7. DOCKER + WAZUH ═════════════════════════════════════
heading("7. Déploiement Docker + Wazuh SIEM",1)

heading("7.1 Docker Compose — Application",2)
code_block("""# compose.yaml (racine du projet)
version: '3.8'
services:
  backend:
    build: ./back
    ports: ["3001:3001"]
    environment:
      - NODE_ENV=production
      - JWT_SECRET=${JWT_SECRET}
      - ENCRYPTION_KEY=${ENCRYPTION_KEY}
    restart: unless-stopped

  frontend:
    build: ./front
    ports: ["80:80","443:443"]
    depends_on: [backend]
    restart: unless-stopped

# Lancer l'application complète
docker compose up -d
docker compose ps     # Vérifier statut conteneurs""")

heading("7.2 Installation Wazuh 4.7.4 — 4 Commandes",2)
info_box("Wazuh s'installe en 4 commandes Docker sans configuration manuelle. Coverage 100% atteint immédiatement.")
code_block("""# 1. Cloner le dépôt Wazuh Docker v4.7.4
git clone -b v4.7.4 https://github.com/wazuh/wazuh-docker.git
cd wazuh-docker/single-node

# 2. Générer les certificats TLS (obligatoire)
docker compose -f generate-indexer-certs.yml run --rm generator

# 3. Démarrer tous les services Wazuh
docker compose up -d

# 4. Accéder au Dashboard Wazuh
# URL : https://localhost
# Login : admin / SecretPassword""")

heading("7.3 Composants Wazuh",2)
add_table(["Composant","Port","Rôle"],[
    ["Wazuh Manager","1514 / 55000","Collecte events · corrélation · moteur règles"],
    ["Wazuh Indexer (OpenSearch)","9200","Stockage alertes · Full-text search"],
    ["Wazuh Dashboard","443 HTTPS TLS","Interface Kibana fork — visualisation"],
    ["Wazuh Agent macOS","1514 (sortant)","Push events vers Manager · chiffré AES"],
],widths=[2.0,1.3,3.0])

heading("7.4 Installation Agent Wazuh (macOS)",2)
code_block("""# Télécharger l'agent macOS ARM64
curl -so wazuh-agent.pkg \
  https://packages.wazuh.com/4.x/macos/wazuh-agent-4.7.4-1.arm64.pkg
sudo installer -pkg wazuh-agent.pkg -target /

# Configurer connexion vers le Manager
sudo /Library/Ossec/bin/agent-auth -m 127.0.0.1

# Démarrer l'agent
sudo /Library/Ossec/bin/wazuh-control start

# Vérifier le statut
sudo /Library/Ossec/bin/wazuh-control status
# Résultat attendu : wazuh-agentd is running...""")

try_img(os.path.join(BASE,"WAZUCAPT","wazuh_agents.png"), w=Inches(5.5),
    caption="Figure 1 — Wazuh Agents Dashboard : main-machine · macOS 15.7.4 · v4.7.4 · Active")

heading("7.5 Wazuh — 5 Modules de Sécurité Actifs",2)
add_table(["Module","Données surveillées","Exigence CDC §3.3"],[
    ["Security Events","436 420 events · Rule 550 Level 7 · Niveaux 3→15","Journalisation actions"],
    ["MITRE ATT&CK","T1565.001 (~95%) · T1562 Evasion (×6) · 1600+ alertes","Détection menaces"],
    ["GDPR Module","Art.15 accès · Art.17 effacement · Art.33 violation","Conformité RGPD"],
    ["Integrity Monitoring","FIM 2500+ events · /bin/bash · Rule 550 · SHA-256","Intégrité fichiers"],
    ["Vulnerabilities","17 CVE · 8 High + 9 Medium · Docker dominant","Détection CVE"],
],widths=[1.7,3.0,1.6])

doc.add_page_break()

# ══ 8. ARCHITECTURE SÉCURITÉ ══════════════════════════════
heading("8. Architecture de Sécurité — Captures d'Écran",1)
para("Les captures suivantes sont issues de l'application en production/local. Elles constituent les preuves visuelles du bon fonctionnement du système de sécurité.")

try_img(os.path.join(BASE,"DASH .png"), w=Inches(5.8),
    caption="Figure 2 — Dashboard Sécurité Principal : Score · Auth · RGPD · RBAC · Audit Logs")
try_img(os.path.join(BASE,"WAFCAPTSIEM","SIEM.png"), w=Inches(5.8),
    caption="Figure 3 — Dashboard SIEM : 18 Auth · 10 RBAC · 20 RGPD · 44 Blocages WAF")
try_img(os.path.join(BASE,"WAFCAPTSIEM","WAF.png"), w=Inches(5.8),
    caption="Figure 4 — WAF Monitoring : 44 attaques bloquées · SQL/XSS/PathTraversal/Scanner")
try_img(os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost.png"), w=Inches(5.8),
    caption="Figure 5 — Dashboard /monitoring complet (admin)")

doc.add_page_break()

# ══ 9. WAF ════════════════════════════════════════════════
heading("9. WAF — Pare-feu Applicatif (waf.js)",1)
para("Fichier : back/functions/src/middlewares/waf.js — Appliqué avant l'authentification.")

add_table(["Type d'attaque","Sévérité","Patterns détectés","Réponse HTTP"],[
    ["Injection SQL","CRITIQUE","SELECT/UNION/DROP/INSERT/OR 1=1","403 + WAF_BLOCK loggué"],
    ["XSS","HAUTE","<script>/eval()/onerror=/javascript:","403 · CSP bloque"],
    ["Path Traversal","HAUTE","../../ · %2e%2e · /etc/passwd","403 + WAF_BLOCK"],
    ["Command Injection","HAUTE","; && | $() · backtick","403 + WAF_BLOCK"],
    ["Scanner suspects","HAUTE","sqlmap/nikto/nmap User-Agent","403 bloqué"],
],widths=[1.5,0.9,2.3,1.9])

add_table(["Règle Brute Force","Seuil","Action","Code HTTP"],[
    ["Rate limit global","100 req/15min","Blocage temporaire","429"],
    ["Rate limit auth","10 req/15min","Blocage IP temporaire","429"],
    ["Lockout brute force","5 échecs consécutifs","AUTH_LOCKOUT Firestore + 5 min blocage","429"],
    ["Payload max","10 KB","Rejet payload","413"],
    ["Timeout requête","28 secondes","Timeout + log","503"],
],widths=[1.8,1.4,2.3,1.1])

code_block("""// Test WAF — Preuves en ligne de commande :
curl 'http://localhost:3001/api/v1/etudiants?id=1 OR 1=1'
# → HTTP 403 {"error":"Requête bloquée par le WAF","code":"WAF_BLOCK"}

curl 'http://localhost:3001/api/v1/etudiants?nom=<script>alert(1)</script>'
# → HTTP 403 {"error":"Requête bloquée par le WAF","code":"WAF_BLOCK"}

curl -H "User-Agent: sqlmap/1.0" 'http://localhost:3001/api/v1/etudiants'
# → HTTP 403 {"error":"Requête bloquée par le WAF","code":"WAF_BLOCK"}""")

doc.add_page_break()

# ══ 10. RBAC ══════════════════════════════════════════════
heading("10. RBAC — Gestion des Rôles et Accès",1)
info_box("Le rôle est relu en base Firestore à CHAQUE requête — jamais depuis le JWT uniquement. Si révoqué : bloqué immédiatement.")

add_table(["Permission","admin","sous-admin","comptable","étudiant/parent"],[
    ["Dashboard /monitoring","✅","❌","❌","❌"],
    ["Gestion utilisateurs","✅","❌","❌","❌"],
    ["Créer étudiant","✅","✅","❌","❌"],
    ["Supprimer étudiant","✅","❌","❌","❌"],
    ["Voir paiements","✅","✅","✅","❌"],
    ["Créer paiement","✅","✅","✅","❌"],
    ["Générer facture","✅","✅","✅","❌"],
    ["Export RGPD Art.15","✅","❌","❌","✅ (ses données)"],
    ["Anonymisation Art.17","✅","❌","❌","❌"],
],widths=[2.0,0.9,0.9,0.9,1.5])

code_block("""// Implémentation checkRole() — back/functions/src/middlewares/auth.js
const checkRole = (allowedRoles) => async (req, res, next) => {
  const userDoc = await db.collection('users').doc(req.user.uid).get();
  const currentRole = userDoc.data()?.role;

  if (!allowedRoles.includes(currentRole)) {
    await logAuditEvent('ACCESS_DENIED', {
      userId: req.user.uid, route: req.path,
      roleAttempted: currentRole, ip: req.ip
    });
    return res.status(403).json({ error: 'Accès refusé' });
  }
  req.userRole = currentRole;
  next();
};

// Exemple d'utilisation
router.get('/etudiants', verifyJWT, checkRole(['admin','sous-admin']), ...);
router.delete('/etudiants/:id', verifyJWT, checkRole(['admin']), ...);""")

doc.add_page_break()

# ══ 11. CHIFFREMENT ═══════════════════════════════════════
heading("11. Chiffrement & Cryptographie",1)
add_table(["Élément","Algorithme","Paramètres","Usage"],[
    ["Mots de passe","bcrypt","12 rounds","Hash one-way — stockage Firestore"],
    ["JWT tokens","HS256","exp: 24h","Authentification stateless"],
    ["Données sensibles","AES-256-CBC","IV aléatoire 16 bytes","Chiffrement champs Firestore"],
    ["Transport","HTTPS/TLS","HSTS max-age=31556926","Toutes communications"],
    ["Fichiers","Firebase Storage","Access tokens signés","Documents admin"],
],widths=[1.5,1.4,1.7,1.7])

code_block("""// Chiffrement AES-256-CBC — back/functions/src/utils/encryption.js
const KEY = Buffer.from(process.env.ENCRYPTION_KEY, 'hex'); // 32 bytes

function encrypt(text) {
  const iv = crypto.randomBytes(16);           // IV aléatoire à chaque fois
  const cipher = crypto.createCipheriv('aes-256-cbc', KEY, iv);
  return iv.toString('hex') + ':' +
    Buffer.concat([cipher.update(text), cipher.final()]).toString('hex');
}""")

doc.add_page_break()

# ══ 12. RGPD ══════════════════════════════════════════════
heading("12. Conformité RGPD (UE 2016/679)",1)
add_table(["Article","Droit","Endpoint","Implémentation"],[
    ["Art. 15","Droit d'accès","GET /users/:id/export","Export JSON · DATA_EXPORT loggué"],
    ["Art. 16","Rectification","PUT /users/:id","CRUD loggué serverTimestamp()"],
    ["Art. 17","Droit à l'effacement","DELETE /users/:id/data","Anonymisation : email→anon@deleted · nom→Anonyme"],
    ["Art. 25","Privacy by Design","—","AES-256-CBC + RBAC moindre privilège dès conception"],
    ["Art. 32","Sécurité traitement","—","WAF + JWT HS256 + bcrypt 12r + HTTPS/HSTS"],
    ["Art. 33","Notification violation","/monitoring","Alerte CRITIQUE auditLogs · Procédure 72h"],
],widths=[0.8,1.3,1.7,2.5])

doc.add_page_break()

# ══ 13. JOURNALISATION ════════════════════════════════════
heading("13. Journalisation & Audit Logs Immuables",1)
info_box("Firestore Rules : allow update: if false · allow delete: if false → LOGS IMMUABLES. serverTimestamp() généré côté serveur uniquement.")

add_table(["Type d'événement","Niveau","Données loggées","Déclencheur"],[
    ["AUTH_SUCCESS","INFO","userId · IP · timestamp","Connexion réussie"],
    ["AUTH_FAILURE","WARNING","userId · IP · tentative n°","Mauvais identifiants"],
    ["AUTH_LOCKOUT","CRITIQUE","userId · IP · durée blocage","5 échecs consécutifs"],
    ["LOGOUT","INFO","userId · durée session","Déconnexion"],
    ["SESSION_EXPIRED","INFO","userId · last_activity","Inactivité 30 min"],
    ["ACCESS_DENIED","WARNING","userId · route · rôle tenté","RBAC insuffisant"],
    ["DATA_EXPORT","INFO","userId · données exportées","Export RGPD Art.15"],
    ["DATA_ANONYMIZE","WARNING","userId · champs anonymisés","Effacement Art.17"],
    ["WAF_BLOCK","CRITIQUE","IP · type attaque · payload","Attaque WAF détectée"],
],widths=[1.6,0.9,2.4,1.4])

doc.add_page_break()

# ══ 14. MONITORING ════════════════════════════════════════
heading("14. Monitoring — Dashboard /monitoring",1)
para("Route : front/src/pages/Monitoring.tsx — Accessible admin seulement — JWT + checkRole(['admin'])")

add_table(["Section","Données","Source"],[
    ["Score /100","Calculé sur 24h (WAF+Auth+RBAC)","auditLogs Firestore"],
    ["Alertes actives","4 seuils : WAF>10 · BruteForce>5 · RBAC>20","Temps réel"],
    ["Accès Sécurisé","Connexions · Échecs · Bloqués · Sessions","auditLogs Auth"],
    ["RGPD","Exports Art.15 · Anonymisations Art.17","auditLogs RGPD"],
    ["RBAC","Accès refusés 24h · Dernière heure","auditLogs RBAC"],
    ["Audit Logs","Total · Auth · RBAC · RGPD · WAF","auditLogs all"],
    ["WAF Tab","Attaques bloquées · Répartition par type","auditLogs WAF_BLOCK"],
    ["SIEM Tab","20 derniers événements · toutes catégories","auditLogs all"],
],widths=[1.5,2.8,2.0])

for img_path,caption in [
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (1).png"),"Figure 6 — Onglet Dashboard : Score + Alertes"),
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (2).png"),"Figure 7 — Onglet WAF : Attaques bloquées"),
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (3).png"),"Figure 8 — Onglet SIEM : Journal événements"),
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (4).png"),"Figure 9 — SIEM : Détails événements sécurité"),
]:
    try_img(img_path, w=Inches(5.5), caption=caption)

doc.add_page_break()

# ══ 15. WAZUH RÉSULTATS ═══════════════════════════════════
heading("15. Wazuh SIEM — Résultats et Captures",1)

heading("15.1 Security Events — 436 420 Events",2)
try_img(os.path.join(BASE,"CVAPTWAZUH","SECURITYEVENTS.png"), w=Inches(5.5),
    caption="Figure 10 — Wazuh Security Events : 436 420 events journalisés · Rule 550 Level 7")

heading("15.2 File Integrity Monitoring (FIM)",2)
try_img(os.path.join(BASE,"WAZUCAPT","wazuh_03_fim.png"), w=Inches(5.5),
    caption="Figure 11 — Wazuh FIM : 2500+ events · SHA-256 · /bin/bash · Rule 550")
try_img(os.path.join(BASE,"CVAPTWAZUH","INTEGRITYMONITORING.png"), w=Inches(5.5),
    caption="Figure 12 — Integrity Monitoring détail")

heading("15.3 CVE — 17 Vulnérabilités Détectées",2)
add_table(["CVE","Package","Sévérité","CVSS3","Action Requise"],[
    ["CVE-2019-5736","Docker 4.43.2","High","8.6","PRIORITÉ P1 — docker pull docker:latest"],
    ["CVE-2019-13509","Docker 4.43.2","High","7.5","Mise à jour Docker Engine"],
    ["CVE-2019-16884","Docker 4.43.2","High","7.5","Mise à jour Docker Engine"],
    ["CVE-2021-21284","Docker 4.43.2","Medium","6.8","Mise à jour Docker Engine"],
    ["CVE-2001-0718","Excel 16.107.3","High","7.5","Mise à jour Microsoft Office"],
    ["CVE-2014-4715","lz4 1.10.0","Medium","5.0","brew upgrade lz4"],
],widths=[1.4,1.4,0.8,0.7,2.0])

try_img(os.path.join(BASE,"WAZUCAPT","wazuh_04_cve.png"), w=Inches(5.5),
    caption="Figure 13 — Wazuh CVE Detection : 17 vulnérabilités · 8 High + 9 Medium")

heading("15.4 MITRE ATT&CK",2)
add_table(["Technique","Tactic","Occurrences","Corrélation"],[
    ["T1565.001 — Stored Data Manipulation","Impact","~95% des events","FIM Rule 550 Level 7"],
    ["T1562 — Disable or Modify Tools","Defense Evasion","×6","Rootcheck anomalies"],
    ["T1499 — Endpoint Denial of Service","Impact","×1594","Rate limiting actif"],
],widths=[2.8,1.2,1.1,1.2])

try_img(os.path.join(BASE,"WAZUCAPT","wazuh_05_mitre.png"), w=Inches(5.5),
    caption="Figure 14 — Wazuh MITRE ATT&CK : T1565.001 dominant · Classification SOC")
try_img(os.path.join(BASE,"CVAPTWAZUH","MITTRE ATTACK.png"), w=Inches(5.5),
    caption="Figure 15 — MITRE ATT&CK Framework : Tactics et Techniques détectées")
try_img(os.path.join(BASE,"CVAPTWAZUH","POLICY MONITORING.png"), w=Inches(5.5),
    caption="Figure 16 — Rootcheck : 4 anomalies · processus caché · interface promiscuous")

doc.add_page_break()

# ══ 16. VARIABLES ENV ═════════════════════════════════════
heading("16. Variables d'Environnement — Référence Complète",1)
add_table(["Variable","Requis","Description","Exemple"],[
    ["JWT_SECRET","✅ OUI","Secret JWT HS256 — 256 bits min","openssl rand -hex 32"],
    ["ENCRYPTION_KEY","✅ OUI","Clé AES-256-CBC — 64 hex chars","openssl rand -hex 32"],
    ["FIREBASE_PROJECT_ID","✅ OUI","ID projet Firebase","frais-gestionscolaire"],
    ["FIREBASE_SERVICE_ACCOUNT_PATH","✅ OUI","Service account JSON","./serviceAccountKey.json"],
    ["NODE_ENV","✅ OUI","Environnement Node.js","production"],
    ["PORT","Non","Port API REST","3001"],
    ["ALLOWED_ORIGINS","✅ OUI","CORS — domaines autorisés","https://frais-gestionscolaire.web.app"],
    ["VITE_API_BASE_URL","✅ OUI","URL API backend (frontend)","https://.../api"],
    ["VITE_FIREBASE_API_KEY","✅ OUI","Clé API Firebase public","AIzaSy..."],
    ["WEBHOOK_SECRET","Non","Secret HMAC webhooks","openssl rand -hex 32"],
],widths=[2.2,0.8,2.0,1.3])

info_box("Ne jamais committer .env dans Git. Utiliser .gitignore et les secrets CI/CD.")

doc.add_page_break()

# ══ 17. ROLLBACK ══════════════════════════════════════════
heading("17. Rollback & Procédures d'Urgence",1)

heading("17.1 Rollback Firebase Hosting",2)
code_block("""firebase hosting:channel:list
# Lister les versions disponibles

firebase hosting:clone SOURCE_SITE_ID:SOURCE_CHANNEL TARGET_SITE_ID:live
# Restaurer une version précédente""")

heading("17.2 Rollback Base de Données",2)
code_block("""# Restaurer depuis sauvegarde Firestore
gcloud firestore databases restore \
  --source-backup=projects/.../backups/BACKUP_ID \
  --destination-database=frais-gestionscolaire

# Sauvegarde manuelle via l'API admin
POST /api/v1/backup   (Bearer <admin_token>)""")

heading("17.3 Procédure Violation de Données Art.33 RGPD",2)
add_table(["Étape","Action","Délai","Responsable"],[
    ["1","Détection via /monitoring — alerte CRITIQUE","0h","Système auto"],
    ["2","Isolation service compromis (docker stop)","< 1h","Admin tech"],
    ["3","Analyse auditLogs — identifier l'étendue","< 2h","Admin tech"],
    ["4","Notification CNIL si > 50 personnes concernées","< 72h","DPO"],
    ["5","Notification personnes concernées","< 72h","Direction"],
    ["6","Correction et redéploiement sécurisé","< 72h","Admin tech"],
    ["7","Rapport post-incident","< 7j","Équipe tech"],
],widths=[0.3,2.8,0.7,1.5])

doc.add_page_break()

# ══ 18. SYNTHÈSE CDC §3.3 ═════════════════════════════════
heading("18. Synthèse CDC §3.3 — 8/8 Exigences Couvertes",1)
add_table(["Exigence CDC §3.3","Implémentation","Module Wazuh","Statut"],[
    ["Accès sécurisé id/mdp","JWT HS256 · bcrypt 12r · WAF · lockout 5 échecs","Security Events","✅ FAIT"],
    ["HTTPS obligatoire","Firebase HTTPS auto · Helmet HSTS max-age=31556926","—","✅ FAIT"],
    ["RGPD conforme","DATA_EXPORT · DATA_ANONYMIZE · AES-256 · Art.15/16/17/33","Module GDPR","✅ FAIT"],
    ["RBAC — Gestion rôles","checkRole() Firestore · 4 rôles · Firestore Rules deny","Security Events","✅ FAIT"],
    ["Journalisation actions","9 events immuables · SIEM dashboard · serverTimestamp()","213+ events","✅ FAIT"],
    ["Détection vulnérabilités","17 CVE identifiées · Plan correctif prioritisé","Vuln Detector","✅ FAIT"],
    ["Intégrité fichiers","FIM 2500+ · Rule 550 · Hash SHA-256","Integrity Mon.","✅ FAIT"],
    ["Détection menaces","MITRE ATT&CK T1499/T1562 · 1600+ alertes","MITRE ATT&CK","✅ FAIT"],
],widths=[2.0,2.5,1.3,0.9])

doc.add_page_break()

# ══ 19. ANNEXES ═══════════════════════════════════════════
heading("19. Annexes",1)

heading("19.1 OWASP Top 10 — Couverture",2)
add_table(["Catégorie OWASP","Implémentation","Statut"],[
    ["A01 — Broken Access Control","RBAC · checkRole() · Firestore Rules","✅ CORRIGÉ"],
    ["A02 — Cryptographic Failures","AES-256-CBC · HTTPS · bcrypt 12 rounds","✅ CORRIGÉ"],
    ["A03 — Injection SQLi·XSS·Cmd","WAF waf.js → HTTP 403 + WAF_BLOCK","✅ CORRIGÉ"],
    ["A04 — Insecure Design","Architecture JWT + RBAC dès la conception","✅ CORRIGÉ"],
    ["A05 — Security Misconfiguration","Helmet · CORS strict · Headers HTTP","✅ CORRIGÉ"],
    ["A06 — Vulnerable Components","Wazuh Vulnerability Detection — 17 CVE","⚠️ SURVEILLÉ"],
    ["A07 — Auth & Session Failures","Rate limiting · lockout 5 échecs · JWT 24h","✅ CORRIGÉ"],
    ["A08 — Data Integrity Failures","auditLogs immuables · JWT signé HS256","✅ CORRIGÉ"],
    ["A09 — Security Logging Failures","9 types audit logs Firestore · SIEM","✅ CORRIGÉ"],
    ["A10 — SSRF","Non applicable — pas d'appels serveur→serveur","N/A"],
],widths=[2.2,2.8,1.2])

heading("19.2 Commandes de Référence",2)
code_block("""# DÉVELOPPEMENT
cd front && npm run dev                          # Frontend :5173
cd back/functions && npm run dev                 # Backend :3001

# TESTS
cd front && npm test                             # Tests unitaires
cd front && npm run test:coverage                # Avec couverture
node back/functions/scripts/security_scan.js    # DAST 12 tests

# DÉPLOIEMENT
firebase deploy --only hosting                   # Frontend
firebase deploy --only firestore:rules           # Règles BDD

# DOCKER
docker compose up -d                             # Lancer app
docker compose ps                                # Statut
docker compose logs -f backend                   # Logs

# WAZUH
sudo /Library/Ossec/bin/wazuh-control status    # Agent status
sudo /Library/Ossec/bin/wazuh-control start     # Démarrer agent""")

sep()
p_f = doc.add_paragraph()
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_f.add_run("Documentation Technique v1.0 — Amine BAHOU / Anass Akker — PFE Cybersécurité — YNOV Campus 2026")
r_f.font.size=Pt(9); r_f.font.italic=True; r_f.font.color.rgb=BLUE

out = os.path.join(BASE, "DOCUMENTATION_TECHNIQUE_YNOV_2026.docx")
doc.save(out)
print(f"✅ Documentation Technique : {out}")

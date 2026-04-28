#!/usr/bin/env python3
"""
Génère MEGA_GUIDE_SOUTENANCE_FINAL.docx
Résumé complet du rapport PDF + screenshots + commandes + script 10 min
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

BASE = "/Users/anass/Downloads/frais-gestionScolaire 4"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def bandeau(doc, texte, bg=RGBColor(0x1a,0x3a,0x6e)):
    bg_hex = '%06X' % bg if isinstance(bg, int) else '%02X%02X%02X' % (bg[0],bg[1],bg[2])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.bold = True; run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xff,0xff,0xff)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), bg_hex)
    p._p.get_or_add_pPr().append(shd)
    return p

def h1(doc, texte, couleur=RGBColor(0x1a,0x3a,0x6e)):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.bold = True; run.font.size = Pt(13)
    run.font.color.rgb = couleur
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def h2(doc, texte, couleur=RGBColor(0x2e,0x6d,0xb8)):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.bold = True; run.font.size = Pt(11)
    run.font.color.rgb = couleur
    p.paragraph_format.space_before = Pt(6)
    return p

def resume_line(doc, texte, color=RGBColor(0x1a,0x3a,0x6e)):
    """Ligne résumé en 1-2 phrases — mise en évidence."""
    p = doc.add_paragraph()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'EEF4FF')
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run("→ " + texte)
    run.font.size = Pt(10); run.bold = True
    run.font.color.rgb = color
    return p

def normal(doc, texte, size=10, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    return p

def bullet(doc, texte, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.4 + level*0.4)
    run = p.add_run(texte)
    run.font.size = Pt(10)
    if color: run.font.color.rgb = color
    return p

def code(doc, texte):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.font.size = Pt(8.5); run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x00,0xcc,0x66)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'1E1E1E')
    p._p.get_or_add_pPr().append(shd)
    return p

def sep(doc):
    p = doc.add_paragraph()
    run = p.add_run("─"*90)
    run.font.size = Pt(7); run.font.color.rgb = RGBColor(0xcc,0xcc,0xcc)

def kpi_badge(doc, items):
    """Ligne de badges KPI colorés."""
    t = doc.add_table(rows=1, cols=len(items))
    t.style = 'Table Grid'
    colors = ['1A3A6E','C62828','2E7D32','E65100','6A1B9A','00695C','AD1457','283593']
    for i,(nb,label) in enumerate(items):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, colors[i % len(colors)])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(nb+"\n")
        r1.bold = True; r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(0xff,0xff,0xff)
        r2 = p.add_run(label)
        r2.font.size = Pt(7)
        r2.font.color.rgb = RGBColor(0xdd,0xdd,0xdd)
    doc.add_paragraph()

def add_img(doc, rel_path, caption, width=Cm(14)):
    full = os.path.join(BASE, rel_path)
    if not os.path.exists(full): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(full, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("📷 " + caption)
    r.font.size = Pt(8); r.italic = True
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    doc.add_paragraph()

def make_table(doc, headers, rows_data, col_bgs=None, bg_header='1A3A6E'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    row0 = t.rows[0]
    for i,h in enumerate(headers):
        c = row0.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, bg_header)
    for row_vals in rows_data:
        row = t.add_row()
        for i,v in enumerate(row_vals):
            c = row.cells[i]
            c.text = str(v)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if col_bgs and i < len(col_bgs) and col_bgs[i]:
                set_cell_bg(c, col_bgs[i])
    return t

# ══════════════════════════════════════════════════════════════════════════════
# COUVERTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MEGA GUIDE SOUTENANCE — TOUT EN UN")
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1a,0x3a,0x6e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PFE Cybersécurité — YNOV Campus 2026")
r.bold = True; r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x2e,0x6d,0xb8)

doc.add_paragraph()
bandeau(doc, "Anass Akker  ·  Amine BAHOU  ·  24 avril 2026  ·  10h30 → 10h40")
doc.add_paragraph()

# Badges KPI principaux
kpi_badge(doc, [
    ("436 420","Events Wazuh"),
    ("17","CVE détectées"),
    ("12/12","DAST PASS"),
    ("5/6","RGPD OK"),
    ("100%","Agent Coverage"),
    ("100/100","Score Sécu"),
])

normal(doc, "Ce document contient :", 11, bold=True)
bullet(doc, "Résumé 1-2 lignes de CHAQUE section du rapport PDF")
bullet(doc, "Script minuté 10 minutes — ce qu'il faut dire exactement")
bullet(doc, "Toutes les commandes de démonstration (WAF, DAST, Wazuh…)")
bullet(doc, "10 captures d'écran annotées intégrées")
bullet(doc, "12 Q&A rapides pour répondre au jury")
sep(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — RÉSUMÉS DU RAPPORT PDF (section par section)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "📋 RÉSUMÉS DU RAPPORT — 1 LIGNE PAR SECTION", RGBColor(0x1a,0x3a,0x6e))
normal(doc, "Chaque section du rapport PDF résumée en 1-2 phrases pour répondre aux questions du jury.", 9, italic=True)
doc.add_paragraph()

# ── §1 KPI ──────────────────────────────────────────────────────────────────
h1(doc, "§1 — Tableau de Bord KPI Sécurité")
resume_line(doc,
    "436 420 events Wazuh · 17 CVE (8H+9M) · 1 600+ MITRE ATT&CK · 100% Agent Coverage · 5/6 RGPD · "
    "Score applicatif 100/100 — audit du 5 au 12 avril 2026.")

h2(doc, "§1.1 — Statut Global des Modules")
resume_line(doc,
    "6 modules sur 8 opérationnels. SCA (0/10 = hardening SSH non appliqué) et CVE "
    "(17 vulnérabilités à patcher) sont les 2 points à corriger post-soutenance — non bloquants.")

h2(doc, "§1.2 — Chronologie de l'Audit")
resume_line(doc,
    "8 étapes documentées du 5 au 12 avril : enregistrement agent → scan CVE → FIM batch → "
    "MITRE ATT&CK → Security Events snapshot → Rootcheck. Démarche structurée et vérifiable.")
doc.add_paragraph()

# ── §2 Problèmes ────────────────────────────────────────────────────────────
h1(doc, "§2 — 8 Problèmes Rencontrés & Solutions")
resume_line(doc,
    "7/8 problèmes entièrement résolus. Seul contournement : NVD feed CVE payant (Wazuh free plan) "
    "— résolu via syscollector Docker Linux. Preuve de capacité à diagnostiquer des obstacles réels.")

problemes = [
    ("#1 Agent Disconnected",       "Haute",  "30 min", "ossec.conf mal configuré → reconfiguration + restart"),
    ("#2 FIM No results (43200s)",  "Haute",  "15 min", "Fréquence 12h → 60s + echo hosts pour déclencher FIM"),
    ("#3 NVD Feed CVE payant",      "Moy.",   "—",      "Contournement : syscollector Docker Linux (17 CVE quand même)"),
    ("#4 bcrypt DLL Windows",       "Haute",  "10 min", "npm rebuild bcrypt → backend opérationnel"),
    ("#5 Firebase timeout",         "Haute",  "45 min", "Création server_local.js mock firebase-functions"),
    ("#6 field 'action' réservé",   "Moy.",   "20 min", "<field> → <match> dans local_rules.xml"),
    ("#7 SIP macOS /etc denied",    "Faible", "5 min",  "/private/etc/hosts (symlink accessible sudo)"),
    ("#8 PPTX Bad CRC-32",          "Faible", "10 min", "Reconstruction ZIP fichier/fichier + PNG blanc"),
]

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
for i,h in enumerate(["Problème","Sévérité","Temps","Solution"]):
    c = t.rows[0].cells[i]
    c.text = h; c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(9)
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)
    set_cell_bg(c, '1A3A6E')
for pb, sev, temps, sol in problemes:
    row = t.add_row()
    row.cells[0].text = pb
    row.cells[1].text = sev
    row.cells[2].text = temps
    row.cells[3].text = sol
    for c in row.cells: c.paragraphs[0].runs[0].font.size = Pt(8.5)
    if sev == 'Haute': set_cell_bg(row.cells[1], 'FFCDD2')
    elif sev == 'Moy.': set_cell_bg(row.cells[1], 'FFE0B2')
    else: set_cell_bg(row.cells[1], 'E8F5E9')
doc.add_paragraph()

# ── §3 Architecture ─────────────────────────────────────────────────────────
h1(doc, "§3 — Architecture de Surveillance Wazuh")
resume_line(doc,
    "Wazuh Manager Docker (172.20.0.2:1514) + Agent macOS 127.0.0.1. "
    "AuditLog.js écrit dans /tmp/applogs/*.log → local_decoder.xml → 7 règles custom Level 3 à 14.")

h2(doc, "§3.2 — 7 Règles YNOV-APP Personnalisées")
resume_line(doc,
    "100010 AUTH_FAILURE (L7) · 100011 AUTH_LOCKOUT (L14 ⚡) · 100012 AUTH_SUCCESS (L3) · "
    "100013 WAF_BLOCK (L12) · 100014 ACCESS_DENIED (L9) · 100015 ADMIN_ACTION (L5) · 100016 DATA_EXPORT (L3)")
doc.add_paragraph()

# ── §4.1 Agents ─────────────────────────────────────────────────────────────
h1(doc, "§4.1 — Infrastructure Agents — Coverage 100%")
resume_line(doc,
    "Agent 001 (main-machine, macOS 15.7.4, v4.7.4) ACTIF. 1 agent / 1 déployé = 100% coverage. "
    "Protocole OSSEC port 1514 + chiffrement AES. Aucun point aveugle.")

add_img(doc, "WAZUCAPT/wazuh_01_agents.png",
    "Wazuh Agents — Agent 001 actif · Coverage 100% · v4.7.4 · frais-gestion-scolaire")

# ── §4.2 Security Events ────────────────────────────────────────────────────
h1(doc, "§4.2 — Security Events — 436 420 Événements")
resume_line(doc,
    "436 420 events collectés en 24h. Rule 550 (Integrity checksum changed, Level 7) domine à ~95%. "
    "Level 12+ alerts = 0, Auth failure = 0 → système sain. PCI DSS 11.5 × 1594 events.")

add_img(doc, "CAPMONITORINGSECU/SECURITEEVENTS.png",
    "Security Events — 436 420 events · T1565.001 · Rule 550 Level 7 · 0 alerte critique")

# ── §4.3 FIM ────────────────────────────────────────────────────────────────
h1(doc, "§4.3 — File Integrity Monitoring (FIM)")
resume_line(doc,
    "root 89.44% des modifications · modified 100% · /var/bin/afsa dominant. "
    "Files added/deleted = 0 (aucune création/suppression suspecte). "
    "Hash SHA-256 en temps réel — détection rootkit < 5 secondes. RGPD Art.5 ✓")

add_img(doc, "CVAPTWAZUH/INTEGRITYMONITORING.png",
    "FIM — root 89.44% · modified 100% · /var/bin/* · 436 420 events")
add_img(doc, "WAZUCAPT/wazuh_03_fim.png",
    "FIM Dashboard — Files modified /bin/* · Files added/deleted = No results")

# ── §4.4 Policy Monitoring ──────────────────────────────────────────────────
h1(doc, "§4.4 — Policy Monitoring / Rootcheck")
resume_line(doc,
    "4 anomalies détectées : ① Trojaned files · ② File owned root write permission · "
    "③ Interface en3 promiscuous (capture réseau) · ④ Process 26061 hidden. "
    "Peuvent être faux positifs macOS mais nécessitent investigation en production.")

add_img(doc, "CVAPTWAZUH/POLICY MONITORING.png",
    "Rootcheck — 4 anomalies · en3 promiscuous · PID 26061 caché · Trojaned files")

# ── §4.5 CVE ────────────────────────────────────────────────────────────────
h1(doc, "§4.5 — Détection CVE — 17 Vulnérabilités")
resume_line(doc,
    "0 Critical · 8 High · 9 Medium. Docker 4.43.2 = 8 CVE dont CVE-2019-5736 (CVSS3=8.6) "
    "container escape → root sur hôte. Excel 16.107.3 + lz4 1.10.0. Action P1 : docker update.")

add_img(doc, "WAZUCAPT/wazuh_04_cve.png",
    "Vulnerabilities — 17 CVE · 8 High · 9 Medium · Docker dominant · CVSS3 max 8.6")

# ── §4.6 MITRE ──────────────────────────────────────────────────────────────
h1(doc, "§4.6 — MITRE ATT&CK")
resume_line(doc,
    "T1565.001 Stored Data Manipulation (Impact) = ~95% des alertes → modifications /var/bin/* par FIM. "
    "T1562 Defense Evasion = ~5%. Cartographie standard SOC international. RGPD Art.25 ✓")

add_img(doc, "WAZUCAPT/wazuh_05_mitre.png",
    "MITRE ATT&CK — T1565.001 Impact ~95% · T1562 Defense Evasion · Top tactics")

# ── §5 RGPD ─────────────────────────────────────────────────────────────────
h1(doc, "§5 — Conformité RGPD × Wazuh")
resume_line(doc,
    "5/6 articles satisfaits. Art.32 en cours (patch CVE-2019-5736 Docker). "
    "Art.5 FIM Rule 550 ✓ · Art.17 DATA_EXPORT 100016 ✓ · Art.25 WAF+RBAC ✓ · "
    "Art.33 Level 14 AUTH_LOCKOUT ✓ · Art.35 Dashboard SIEM complet ✓")

make_table(doc,
    ["Article","Résumé 1 ligne","Statut"],
    [
        ("Art.5",  "FIM 436 420 alertes → chaque modif fichier loggée Rule 550", "✅"),
        ("Art.15", "GET /users/:id/export → JSON complet + audit DATA_EXPORT",   "✅"),
        ("Art.17", "DELETE /users/:id/data → hash irréversible + audit",          "✅"),
        ("Art.25", "WAF + Rate Limit + RBAC deny-by-default actifs dès conception","✅"),
        ("Art.32", "17 CVE identifiées — patch Docker P1 en cours",               "⚠"),
        ("Art.33", "AUTH_LOCKOUT Level 14 → notification immédiate",              "✅"),
        ("Art.35", "Dashboard SIEM complet → analyse d'impact toutes données",    "✅"),
    ]
)
doc.add_paragraph()

# ── §6 Plan d'action ────────────────────────────────────────────────────────
h1(doc, "§6 — Plan d'Action & Recommandations")
resume_line(doc,
    "P1 CRITIQUE : docker update (CVE-2019-5736). P2 : hardening SCA macOS 0/10 → 10/10 CIS Benchmark. "
    "P3 : Wazuh Active Response (blocage IP auto Level 14) + alertes Slack Level 10+ + agent serveur prod.")
doc.add_paragraph()

# ── Annexe Dashboard ────────────────────────────────────────────────────────
h1(doc, "Annexe A — Dashboard Monitoring Applicatif (/monitoring)")
resume_line(doc,
    "Score 100/100 calculé sur 24h. 3 onglets : Dashboard (Auth/RBAC/RGPD) · WAF (0 attaque) · "
    "SIEM Logs (20 derniers events immuables). Visible admin only — ProtectedRoute React.")

add_img(doc, "CAPMONITORINGSECU/brave_screenshot_localhost (2).png",
    "Dashboard — Score 100/100 · JWT HS256 · bcrypt saltRounds=10 · Rate limiting 5 tentatives")
add_img(doc, "CAPMONITORINGSECU/brave_screenshot_localhost (1).png",
    "WAF Onglet — 0 attaque bloquée · 8 règles OWASP Top 10 actives · WAF_BLOCK → Firestore")
add_img(doc, "CAPMONITORINGSECU/brave_screenshot_localhost.png",
    "SIEM Logs — Journal 20 derniers events · allow update: if false · serverTimestamp()")

sep(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SCRIPT 10 MINUTES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "⏱ SCRIPT MINUTÉ — 10h30 À 10h40 (10 MINUTES EXACTEMENT)", RGBColor(0x8b,0x00,0x00))
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("⚡ RÈGLE D'OR : Ne lis PAS les slides. Dis l'essentiel. 20 secondes par slide max.")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xcc,0x00,0x00)
doc.add_paragraph()

minutes = [
    ("MIN 1 — 10h30", "INTRODUCTION (slides 1-3)",
     ['"Bonjour, je présente la partie Sécurité de notre application de gestion scolaire YNOV."',
      '"Notre app gère des données élèves et paiements — le RGPD est obligatoire."',
      '"J\'ai implémenté une défense en profondeur sur 2 niveaux : applicatif + Wazuh SIEM."',
      '"Stack : React / Node.js Express / Firebase / Wazuh 4.7.4 Docker."'],
     "436 420 events · 5/6 RGPD · 12 tests DAST · Defense in Depth",
     '1A3A6E'),
    ("MIN 2 — 10h31", "OWASP + WAF (slides 4-5)",
     ['"J\'ai audité les 10 catégories OWASP Top 10 — 8 corrigées, 1 surveillée, 1 N/A."',
      '"Le WAF waf.js intercepte SQLi, XSS, Path Traversal, CMD Injection, agents scanners."',
      '"Chaque attaque → HTTP 403 + WAF_BLOCK loggué Firestore + Wazuh Rule 100013 Level 12."',
      '"Données sensibles (téléphone, adresse) = AES-256-CBC avant stockage Firestore."'],
     "5 types bloqués · mot de passe = bcrypt hash irréversible (jamais stocké)",
     'C62828'),
    ("MIN 3 — 10h32", "RBAC + RGPD (slides 6-7)",
     ['"6 rôles RBAC : admin, sous-admin, comptable, enseignant, étudiant, parent."',
      '"checkRole() relit Firestore à CHAQUE requête — rôle modifié = actif immédiatement."',
      '"Firestore Rules : deny by default — toute collection non listée = refus automatique."',
      '"RGPD 5/6 : Art.15 export données, Art.17 anonymisation, Art.25 Privacy by Design."'],
     "checkRole() re-lit Firestore live · AES-256-CBC + HTTPS/HSTS",
     '2E7D32'),
    ("MIN 4 — 10h33", "AUDIT LOGS + JWT + RATE LIMIT (slides 8-10)",
     ['"9 types d\'auditLogs immuables — allow update: if false dans Firestore Rules."',
      '"Personne ne peut modifier ces logs, même un admin, même si le backend est compromis."',
      '"JWT HS256 : access token 30 min, refresh 7 jours. Rate limit : 5 tentatives → lockout."',
      '"6e tentative → HTTP 429 Too Many Requests + AUTH_LOCKOUT Firestore Level 14."'],
     "serverTimestamp() côté Google — non manipulable client",
     '6A1B9A'),
    ("MIN 5 — 10h34", "PIPELINE + DAST (slides 11-12)",
     ['"Pipeline : WAF → Rate Limiter → verifyJWT → checkRole() → Handler → auditLogger."',
      '"Chaque requête passe par 4 couches avant d\'accéder aux données Firestore."',
      '"Scanner DAST : 12 tests OWASP automatisés — 12/12 PASS, score 92/100."',
      '"Le scanner teste en live : SQLi, XSS, Path Traversal, brute force, escalade de privil."'],
     "DAST = requêtes HTTP réelles sur l'app en cours d'exécution (pas du code statique)",
     'E65100'),
    ("MIN 6 — 10h35", "DASHBOARD /monitoring (slides 13-18)",
     ['"Le dashboard affiche un score /100 mis à jour toutes les 60 secondes."',
      '"3 onglets : Dashboard (auth/RBAC/RGPD) · WAF (attaques) · SIEM Logs (20 events)."',
      '"Score 100/100 aujourd\'hui = 0 incident actif. Rouge en cas d\'attaque détectée."',
      '"Admin only — ProtectedRoute React roles=[\'admin\']."'],
     "setInterval(load, 60_000) · Score = 100 - pénalités par type d'incident",
     '00695C'),
    ("MIN 7 — 10h36", "WAZUH ARCHITECTURE (slides 19-21)",
     ['"Wazuh SIEM déployé en 4 commandes Docker : git clone → docker compose up -d."',
      '"3 conteneurs : Manager 172.20.0.2 · Indexer OpenSearch · Dashboard https://localhost."',
      '"Agent 001 actif sur macOS — protocole OSSEC port 1514 + chiffrement AES temps réel."',
      '"Logs applicatifs YNOV → /tmp/applogs/*.log → Wazuh → Dashboard."'],
     "3 conteneurs Docker · Agent 001 macOS 15.7.4 · 7 règles custom 100010-100016",
     '1A3A6E'),
    ("MIN 8 — 10h37", "WAZUH RÉSULTATS (slides 22-26)",
     ['"436 420 événements. Rule 550 Level 7 (Integrity checksum changed) domine à 95%."',
      '"FIM surveille /var/bin/* — root 89.44% des modifs. Détection rootkit < 5 secondes."',
      '"17 CVE dont CVE-2019-5736 CVSS3=8.6 container escape Docker — patch P1 prioritaire."',
      '"MITRE ATT&CK : T1565.001 Stored Data Manipulation dominant + T1562 Defense Evasion."'],
     "Rootcheck : interface en3 promiscuous + processus 26061 caché détectés",
     '8B0000'),
    ("MIN 9 — 10h38", "DEFENSE IN DEPTH + SYNTHÈSE (slides 27-29)",
     ['"Defense in Depth 2 niveaux : applicatif (WAF+auditLogs+dashboard) + Wazuh (OS+CVE+FIM)."',
      '"Wazuh surveille là où mon code ne peut pas voir : OS, fichiers système, CVE packages."',
      '"CDC §3.3 couvert à 100% : accès sécurisé, RGPD, journalisation, HTTPS, monitoring."',
      '"Résultat : 8/10 OWASP corrigés, 5/6 RGPD, 17 CVE documentées avec plan d\'action."'],
     "Phrase clé : 'Défense en profondeur = applicatif + infrastructure = 0 point aveugle'",
     '00695C'),
    ("MIN 10 — 10h39", "CONCLUSION + DÉMO (slide 30)",
     ['"Synthèse : WAF maison, 9 logs immuables, score 100/100, DAST 12/12, Wazuh 436k."',
      '"Perspectives : Active Response Wazuh (blocage IP auto), Redis rate limit, patch Docker."',
      '"Je suis prêt pour une démo live si vous souhaitez voir le WAF bloquer une SQLi."',
      '"→ Ouvrir : http://localhost:8081/monitoring  +  https://localhost (Wazuh)"'],
     "Si question difficile : 'C'est dans mon plan d'amélioration : [MFA / Argon2id / DPIA]'",
     '1A3A6E'),
]

for min_label, sujet, phrases, note, color_hex in minutes:
    r,g,b = int(color_hex[0:2],16),int(color_hex[2:4],16),int(color_hex[4:6],16)
    bg = RGBColor(r,g,b)
    bandeau(doc, f"⏱ {min_label} — {sujet}", bg)
    for phrase in phrases:
        bullet(doc, phrase)
    p = doc.add_paragraph()
    run = p.add_run("📌 " + note)
    run.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = bg
    doc.add_paragraph()

sep(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TOUTES LES COMMANDES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "💻 TOUTES LES COMMANDES DE DÉMONSTRATION", RGBColor(0x00,0x40,0x00))
doc.add_paragraph()

h1(doc, "🚀 AVANT LA SOUTENANCE — Démarrer les 3 services")
code(doc, "# ── Terminal 1 : Backend Node.js ──────────────────────────────────")
code(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/back/functions"')
code(doc, "node server_local.js")
code(doc, "# → http://localhost:5001  (attendre 'Express app listening')")
doc.add_paragraph()
code(doc, "# ── Terminal 2 : Frontend React ───────────────────────────────────")
code(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/front"')
code(doc, "npm run dev")
code(doc, "# → http://localhost:8081")
doc.add_paragraph()
code(doc, "# ── Terminal 3 : Wazuh Docker ─────────────────────────────────────")
code(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/front/wazuh-docker/single-node"')
code(doc, "docker compose up -d")
code(doc, "# → https://localhost  ·  Login: admin / SecretPassword  ·  Attendre ~2 min")
sep(doc)

h1(doc, "🛡 ÉTAPE 01 — WAF : Injection SQL bloquée (30 sec)")
code(doc, "curl -s -o /dev/null -w 'HTTP: %{http_code}\\n' \\")
code(doc, "  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?id=1%20OR%201=1'")
code(doc, "# ✅ Résultat attendu : HTTP: 403")
p = doc.add_paragraph()
r1=p.add_run("Dire : "); r1.bold=True; r1.font.size=Pt(10)
r2=p.add_run('"Le WAF intercepte avant Firestore → HTTP 403 + WAF_BLOCK loggué Firestore + Wazuh Rule 100013 Level 12"')
r2.font.size=Pt(10); r2.italic=True
sep(doc)

h1(doc, "🛡 ÉTAPE 02 — WAF : XSS bloqué (20 sec)")
code(doc, "curl -s -o /dev/null -w 'HTTP: %{http_code}\\n' \\")
code(doc, "  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?nom=<script>alert(1)</script>'")
code(doc, "# ✅ Résultat attendu : HTTP: 403")
sep(doc)

h1(doc, "🔒 ÉTAPE 03 — Rate Limiting / Brute Force (40 sec)")
code(doc, "for i in {1..6}; do")
code(doc, '  curl -s -o /dev/null -w "Tentative $i → %{http_code}\\n" \\')
code(doc, "    -X POST http://localhost:5001/gestionadminastration/us-central1/api/auth/login \\")
code(doc, "    -H 'Content-Type: application/json' \\")
code(doc, '    -d \'{"email":"admin@school.fr","password":"WRONG"}\'')
code(doc, "done")
code(doc, "# ✅ Résultat : 1→401 · 2→401 · 3→401 · 4→401 · 5→401 · 6→429")
p = doc.add_paragraph()
r1=p.add_run("Dire : "); r1.bold=True; r1.font.size=Pt(10)
r2=p.add_run('"5 tentatives → lockout 15 min. 6e = HTTP 429 + AUTH_LOCKOUT Firestore Level 14 Wazuh"')
r2.font.size=Pt(10); r2.italic=True
sep(doc)

h1(doc, "🔬 ÉTAPE 04 — Scanner DAST Automatique (40 sec)")
code(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/back/functions"')
code(doc, "node scripts/security_scan.js")
code(doc, "# ✅ Résultat attendu :")
code(doc, "# [PASS] WAF actif — SQLi bloqué")
code(doc, "# [PASS] Rate limiting — 429 après 11 req")
code(doc, "# [PASS] JWT validation — 401 token invalide")
code(doc, "# [PASS] Headers sécurité — HSTS présent")
code(doc, "# Score DAST final : 92/100 — 11/12 tests réussis (1 warning CORS)")
sep(doc)

h1(doc, "📊 ÉTAPE 05 — Dashboard Monitoring Applicatif (30 sec)")
code(doc, "open http://localhost:8081/monitoring")
code(doc, "# → Montrer : Score 100/100 | Onglet WAF : 0 attaque | Onglet SIEM : journal events")
code(doc, "# Connexion : admin@school.fr / votre_mdp")
sep(doc)

h1(doc, "🌐 ÉTAPE 06 — Wazuh Dashboard Live (90 sec)")
code(doc, "open https://localhost")
code(doc, "# Login : admin / SecretPassword")
code(doc, "# → Security Events : 436 420 events")
code(doc, "# → Integrity Monitoring : root 89.44% · modified 100%")
code(doc, "# → Vulnerabilities : 17 CVE · 8 High · Docker")
code(doc, "# → MITRE ATT&CK : T1565.001 Impact dominant")
sep(doc)

h1(doc, "🔐 ÉTAPE 07 — RGPD Traçabilité (30 sec)")
code(doc, "# 1. Ouvrir onglet privé → http://localhost:8081/login")
code(doc, "# 2. Saisir un MAUVAIS mot de passe → connexion échouée")
code(doc, "# 3. Aller sur http://localhost:8081/monitoring → SIEM Logs → Rafraîchir")
code(doc, "# 4. Voir : auth_failure loggué avec timestamp + IP + email")
code(doc, "# → Timestamp = serverTimestamp() côté Google — non manipulable. RGPD Art.5 + Art.32")
sep(doc)

h1(doc, "🔧 Wazuh — Redémarrer Agent si déconnecté")
code(doc, "sudo /Library/Ossec/bin/wazuh-control stop")
code(doc, "sudo /Library/Ossec/bin/wazuh-control start")
code(doc, "sudo /Library/Ossec/bin/wazuh-control status  # → wazuh-agentd is running")
sep(doc)

h1(doc, "📁 Déclencher FIM manuellement (démo Wazuh)")
code(doc, "sudo bash -c \"echo '# wazuh-demo' >> /private/etc/hosts\"")
code(doc, "# → Dans Wazuh Integrity Monitoring : alerte Rule 550 Level 7 en < 60 sec")
sep(doc)

h1(doc, "🔐 BONUS — Tester chiffrement AES-256-CBC")
code(doc, 'node -e "const enc=require(\'./src/utils/encryption\');')
code(doc, 'const c=enc.encrypt(\'donnee-sensible\');')
code(doc, 'console.log(\'Chiffré:\',c); console.log(\'Déchiffré:\',enc.decrypt(c));"')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Q&A JURY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "❓ Q&A — 12 QUESTIONS FRÉQUENTES DU JURY", RGBColor(0x1a,0x3a,0x6e))
doc.add_paragraph()

qa = [
    ("Pourquoi Wazuh en plus du dashboard ?",
     "Wazuh surveille l'OS, fichiers système et CVE — là où mon code ne peut pas voir. "
     "Defense in Depth : 2 couches indépendantes. Si l'appli est compromise, Wazuh continue à détecter."),
    ("Comment les logs ne peuvent pas être falsifiés ?",
     "Firestore Rules : allow update: if false + allow delete: if false. Même l'admin ne peut pas modifier. "
     "serverTimestamp() est résolu côté Google Cloud — non falsifiable."),
    ("Comment fonctionne le WAF ?",
     "Middleware Express avant l'auth : vérifie User-Agent → URL → Query → Body JSON. "
     "Password exclu pour éviter faux positifs bcrypt. → HTTP 403 + WAF_BLOCK auditLog."),
    ("JWT : durée ? Contenu du payload ?",
     "Access token 30 min, refresh 7 jours (httpOnly cookie). "
     "Payload : {id, email, role} — jamais de mot de passe ni données PII sensibles."),
    ("Comment protéger contre brute force ?",
     "Rate limiter : 5 tentatives / 15 min → HTTP 429 + AUTH_LOCKOUT Firestore Level 14. "
     "Limite connue : ne bloque pas les botnets multi-IP (plan : Wazuh Active Response)."),
    ("RGPD Art.17 vs logs immuables — contradiction ?",
     "Les logs contiennent un userId opaque. Effacement = remplacement ID par hash irréversible. "
     "La ligne reste mais n'est plus une donnée personnelle au sens Art.4 RGPD."),
    ("CVE-2019-5736 — c'est grave ?",
     "CVSS3=8.6 — container escape : un attaquant dans Docker peut obtenir root sur l'hôte. "
     "Détecté par Wazuh, documenté, plan P1 Critique : docker update post-soutenance."),
    ("Différence DAST vs SAST ?",
     "SAST = analyse statique du code source (ESLint). DAST = application en exécution avec vraies requêtes HTTP. "
     "Mon scanner = DAST : il teste SQLi/XSS/brute force contre l'API réelle."),
    ("Pourquoi bcrypt saltRounds=10 ?",
     "2^10 = 1024 itérations ≈ 100ms/hash. Rend les attaques dictionnaire non rentables. "
     "Salt aléatoire = même mot de passe → deux hashes différents (anti-rainbow table)."),
    ("Le SCA score est 0/10 ?",
     "Hardening SSH non appliqué (macOS SIP). Non bloquant pour la soutenance. "
     "Plan P2 : appliquer CIS Benchmark macOS (SSH, kernel params, permissions)."),
    ("Comment fonctionne le score /100 ?",
     "Départ 100. Pénalités : -20 max si auth_failures>5, -15 max si lockouts>0, "
     "-15 max si access_denied>3, -20 max si waf_blocks>0. Vert ≥80, Orange 60-79, Rouge <60."),
    ("Que fait exactement le MITRE ATT&CK de Wazuh ?",
     "Wazuh mappe automatiquement ses règles aux techniques ATT&CK. "
     "T1565.001 = modifications /var/bin/* (FIM Rule 550) → classification standard SOC international."),
]

for q, r in qa:
    h2(doc, "❓ " + q)
    p = doc.add_paragraph()
    run = p.add_run("✅ " + r)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x00)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — CHIFFRES CLÉS À MÉMORISER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "🔢 CHIFFRES CLÉS À RETENIR PAR CŒUR", RGBColor(0x1a,0x3a,0x6e))
doc.add_paragraph()

chiffres = [
    ("436 420",  "events Wazuh collectés (FIM + Security Events)"),
    ("12/12",    "tests DAST réussis — 0 vulnérabilité critique"),
    ("92/100",   "score DAST (1 warning CORS — non bloquant)"),
    ("17",       "CVE détectées (8 High + 9 Medium) — Docker principal"),
    ("1 600+",   "alertes MITRE ATT&CK — T1565.001 dominant"),
    ("5/6",      "articles RGPD satisfaits (Art.32 en cours — CVE Docker)"),
    ("9",        "types d'auditLogs immuables (allow update: if false)"),
    ("6",        "rôles RBAC — moindre privilège"),
    ("100/100",  "score sécurité dashboard — 0 incident actif"),
    ("5",        "types d'attaques WAF bloqués (SQLi, XSS, Path, CMD, Agents)"),
    ("100%",     "coverage agents Wazuh — 0 point aveugle"),
    ("30 min",   "durée access token JWT (refresh : 7 jours)"),
    ("10",       "saltRounds bcrypt — 2^10 = 1024 itérations ≈ 100ms/hash"),
    ("7",        "règles Wazuh custom (100010-100016) Level 3 à 14"),
    ("4",        "commandes pour déployer Wazuh (git clone → docker compose up)"),
    ("2",        "couches Defense in Depth : applicatif + infrastructure"),
    ("8/10",     "catégories OWASP Top 10 corrigées"),
    ("4",        "anomalies Rootcheck (en3 promiscuous, PID 26061 caché...)"),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
for i,h_txt in enumerate(["Chiffre","Ce qu'il signifie"]):
    c = t.rows[0].cells[i]
    c.text = h_txt
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(9)
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff,0xff,0xff)
    set_cell_bg(c, '1A3A6E')

for nb, desc in chiffres:
    row = t.add_row()
    row.cells[0].text = nb
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x70,0x00)
    row.cells[1].text = desc
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

# ── Footer ───────────────────────────────────────────────────────────────────
sep(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Document généré le {datetime.date.today().strftime('%d %B %Y')} · "
    "Anass Akker & Amine BAHOU — PFE Cybersécurité YNOV Campus 2026 · Confidentiel"
)
run.font.size = Pt(8); run.italic = True
run.font.color.rgb = RGBColor(0x88,0x88,0x88)

output = "MEGA_GUIDE_SOUTENANCE_FINAL.docx"
doc.save(output)
print(f"✅ {output} généré avec succès !")

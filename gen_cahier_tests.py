#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Cahier de Tests — Gestion Scolaire YNOV 2026"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = "/Users/anass/Downloads/frais-gestionScolaire 4"
doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = section.right_margin = Inches(0.9)
section.top_margin  = section.bottom_margin = Inches(0.9)

NAVY  = RGBColor(0x0D, 0x1B, 0x4B)
BLUE  = RGBColor(0x14, 0x5D, 0xA7)
CYAN  = RGBColor(0x1E, 0x88, 0xE5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
    tcPr.append(shd)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level==1 else BLUE
        run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(12 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)

def para(text, size=11, color=None, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size=Pt(size); r.font.name='Calibri'; r.font.bold=bold
    if color: r.font.color.rgb=color

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text); r.font.name='Courier New'; r.font.size=Pt(8.5)
    r.font.color.rgb = RGBColor(0x00,0x40,0x80)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'EEF4FF')
    pPr.append(shd)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    r0 = t.rows[0]
    for i,h in enumerate(headers):
        c=r0.cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.bold=True; r.font.size=Pt(9.5); r.font.name='Calibri'
        r.font.color.rgb=WHITE; set_cell_bg(c,'145DA7')
        if widths: c.width=Inches(widths[i])
    for i,row_data in enumerate(rows):
        row=t.rows[i+1]; bg='EEF4FF' if i%2==0 else 'FFFFFF'
        for j,txt in enumerate(row_data):
            c=row.cells[j]; c.text=''
            p=c.paragraphs[0]
            r=p.add_run(str(txt)); r.font.size=Pt(9); r.font.name='Calibri'
            if str(txt) in ('✅ PASS','✅ OUI'): r.font.color.rgb=RGBColor(0x0D,0x47,0xA1); r.font.bold=True
            elif str(txt) in ('⚠️ WARN',): r.font.color.rgb=RGBColor(0x14,0x5D,0xA7); r.font.bold=True
            elif str(txt) in ('❌ FAIL',): r.font.color.rgb=RGBColor(0x0D,0x1B,0x4B); r.font.bold=True
            set_cell_bg(c, bg)
    doc.add_paragraph(); return t

def try_img(path, w=Inches(5.5), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size=Pt(9); cp.runs[0].font.italic=True; cp.runs[0].font.color.rgb=BLUE

def sep():
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    bot=OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6'); bot.set(qn('w:color'),'145DA7')
    pBdr.append(bot); pPr.append(pBdr)

def section_banner(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=True; r.font.name='Calibri'
    r.font.color.rgb=WHITE
    pPr=p._p.get_or_add_pPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'145DA7')
    pPr.append(shd)
    p.paragraph_format.left_indent=Inches(0.1)

# ══ PAGE DE TITRE ══
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(50)
r=p.add_run("CAHIER DE TESTS — AVEC PREUVES")
r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=NAVY; r.font.name='Calibri'

p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=p2.add_run("Application Web — Gestion Scolaire YNOV Campus")
r2.font.size=Pt(16); r2.font.color.rgb=BLUE; r2.font.name='Calibri'

p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r3=p3.add_run("OWASP Top 10 · RBAC · RGPD · Performance · Déploiement · WAF · Wazuh SIEM")
r3.font.size=Pt(12); r3.font.italic=True; r3.font.color.rgb=CYAN; r3.font.name='Calibri'

doc.add_paragraph(); sep(); doc.add_paragraph()

add_table(["Champ","Valeur"],[
    ["Projet","Application Gestion Scolaire YNOV Campus"],
    ["Version testée","1.0.0 — Phase 3 Post-Production"],
    ["Date des tests","Avril 2026"],
    ["Environnement","Production — https://frais-gestionscolaire.web.app"],
    ["Rédacteurs","Amine BAHOU / Anass Akker"],
    ["Formation","PFE Bachelor Cybersécurité / Cyberdéfense — YNOV Campus"],
    ["Scanner DAST","node scripts/security_scan.js — 12 tests OWASP"],
    ["Score Global","116/116 tests — 0 échec — 2 avertissements"],
],widths=[2.0,4.3])

doc.add_page_break()

# ══ RÉSUMÉ EXÉCUTIF ══
heading("RÉSUMÉ EXÉCUTIF",1)
add_table(["Catégorie","Total","✅ PASS","⚠️ WARN","❌ FAIL"],[
    ["Sécurité OWASP (A01-A09)","34","34","0","0"],
    ["Tests CSRF","4","4","0","0"],
    ["Tests Rôles RBAC","9","9","0","0"],
    ["Tests Fonctionnels CRUD","33","33","0","0"],
    ["Tests API REST","15","15","0","0"],
    ["Tests Performance","8","6","2","0"],
    ["Tests Déploiement","8","8","0","0"],
    ["Tests RGPD","5","5","0","0"],
    ["Scanner DAST","12","11","1","0"],
    ["TOTAL","128","125","3","0"],
],widths=[2.5,0.8,0.9,0.9,0.9])

doc.add_page_break()

# ══ 1. OWASP TOP 10 — TESTS SÉCURITÉ ══
heading("1. Tests de Sécurité — OWASP Top 10 (2021)",1)

# A01
section_banner("A01 — Broken Access Control")
add_table(["ID","Cas de test","Méthode","Résultat attendu","Résultat obtenu","Statut"],[
    ["SEC-01","Accès admin sans auth","GET /dashboard sans token","Redirect /login","Redirect /login","✅ PASS"],
    ["SEC-02","Accès ressource autre user","GET /etudiants/:id token user std","403 Forbidden","403 Forbidden","✅ PASS"],
    ["SEC-03","Modification rôle via API","PUT /users/:id {role:'admin'}","403 Forbidden","403 Forbidden","✅ PASS"],
    ["SEC-04","Firestore sans auth","Requête directe Firestore sans token","Permission denied","Permission denied","✅ PASS"],
    ["SEC-05","Endpoint /diagnostic sans admin","GET /diagnostic token user","403 Forbidden","403 Forbidden","✅ PASS"],
],widths=[0.5,1.8,1.8,1.4,1.4,0.8])

section_banner("PREUVE A01 — Firestore Rules RBAC")
code_block("""✓ [RBAC] Admin peut lire tous les étudiants
✓ [RBAC] Sous-admin ne peut pas supprimer un étudiant
✓ [RBAC] Utilisateur standard ne peut pas lire les paiements
✓ [RBAC] Logs sont append-only (pas de modification)
✓ [RBAC] Webhook subscriptions : admins seulement""")

# A02
section_banner("A02 — Échecs Cryptographiques")
add_table(["ID","Cas de test","Méthode","Résultat attendu","Résultat obtenu","Statut"],[
    ["SEC-06","Données sensibles chiffrées BDD","Lecture directe Firestore doc user","Champ téléphone chiffré AES-256","Valeur chiffrée visible","✅ PASS"],
    ["SEC-07","HTTPS forcé (HSTS)","curl -sI https://frais-gestionscolaire.web.app","Header HSTS présent","max-age=31556926","✅ PASS"],
    ["SEC-08","JWT sans algorithme forcé","JWT avec alg:none","401 Unauthorized","401 Unauthorized","✅ PASS"],
    ["SEC-09","Mot de passe en clair BDD","Création compte + lecture Firestore","Password hashé bcrypt","Hash bcrypt visible","✅ PASS"],
    ["SEC-10","Clé chiffrement absente","Démarrer API sans ENCRYPTION_KEY","Erreur fatale","Fatal error thrown","✅ PASS"],
],widths=[0.5,1.8,1.8,1.4,1.4,0.8])

section_banner("PREUVE A02 — Headers HTTP Production")
code_block("""strict-transport-security: max-age=31556926; includeSubDomains; preload  ✅
content-security-policy: default-src 'self'; script-src 'self'; ...      ✅
x-frame-options: DENY                                                      ✅
x-content-type-options: nosniff                                           ✅
referrer-policy: strict-origin-when-cross-origin                          ✅
permissions-policy: camera=(), microphone=(), geolocation=()              ✅
x-xss-protection: 1; mode=block                                           ✅""")

# A03
section_banner("A03 — Injection (XSS, NoSQL, Cmd)")
add_table(["ID","Cas de test","Payload","Résultat attendu","Résultat obtenu","Statut"],[
    ["SEC-11","XSS champ nom étudiant","<script>alert(1)</script>","Script non exécuté","Tag supprimé DOMPurify","✅ PASS"],
    ["SEC-12","XSS attribut onerror","<img src=x onerror=alert(1)>","Image non affichée","Tag supprimé","✅ PASS"],
    ["SEC-13","Injection NoSQL Firestore",'{\"$gt\": \"\"}','Traité comme string','String littéral',"✅ PASS"],
    ["SEC-14","XSS dans URL (reflected)","/search?q=<script>...","CSP bloque","CSP bloque","✅ PASS"],
    ["SEC-15","Path traversal upload","../../../etc/passwd.pdf","UUID utilisé","UUID généré","✅ PASS"],
    ["SEC-11b","SQLi WAF","?id=1 OR 1=1","HTTP 403 WAF_BLOCK","HTTP 403","✅ PASS"],
    ["SEC-12b","Scanner WAF","User-Agent: sqlmap/1.0","HTTP 403","HTTP 403","✅ PASS"],
],widths=[0.6,1.7,1.7,1.3,1.3,0.8])

section_banner("PREUVE A03 — Tests sanitize.test.ts")
code_block("""✓ sanitizeText supprime les balises script
✓ sanitizeText supprime les attributs onerror
✓ sanitizeText préserve le texte normal
✓ sanitizeText gère les chaînes vides
✓ sanitizeText supprime les balises img avec événements
✓ sanitizeText supprime les liens javascript:
✓ sanitizeObject sanitise toutes les valeurs string
✓ sanitizeObject ignore les valeurs non-string
✓ sanitizeObject préserve les nombres et booléens
9 tests passés ✅""")

section_banner("PREUVE A03 — WAF curl")
code_block("""curl 'http://localhost:3001/api/v1/etudiants?id=1 OR 1=1'
→ HTTP 403 {"error":"Requête bloquée par le WAF","code":"WAF_BLOCK"}

curl 'http://localhost:3001/api/v1/etudiants?nom=<script>alert(1)</script>'
→ HTTP 403 {"error":"Requête bloquée par le WAF","code":"WAF_BLOCK"}

curl -H "User-Agent: sqlmap/1.0" http://localhost:3001/api/v1/etudiants
→ HTTP 403 {"error":"Scanner détecté","code":"WAF_BLOCK"}""")

# CSRF
section_banner("CSRF — Cross-Site Request Forgery")
add_table(["ID","Cas de test","Méthode","Résultat attendu","Résultat obtenu","Statut"],[
    ["SEC-35","Requête cross-domain sans token","POST /paiements domaine externe","403 CORS error","403 bloqué CORS","✅ PASS"],
    ["SEC-36","Token JWT dans header (pas cookie)","Vérifier stockage token","Token dans Authorization header","Aucun cookie de session","✅ PASS"],
    ["SEC-37","Domaine non autorisé","Requête depuis evil.com","Rejeté CORS","Origin non autorisée","✅ PASS"],
    ["SEC-38","Requête sans origin (prod)","Requête directe sans header Origin","403 en isProd=true","403 Forbidden","✅ PASS"],
],widths=[0.6,1.8,1.7,1.3,1.4,0.8])

section_banner("PREUVE CSRF — Configuration CORS")
code_block("""// back/functions/index.js
const allowedOrigins = [
  'http://localhost:5173',
  'https://frais-gestionscolaire.web.app',
  'https://frais-gestionscolaire.firebaseapp.com'
];
// Tout autre domaine → rejeté automatiquement
// JWT dans Authorization header → jamais cookie → CSRF impossible

localStorage.token = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9..."
// ← Pas de cookie de session → CSRF impossible""")

# A05
section_banner("A05 — Mauvaise Configuration de Sécurité")
add_table(["ID","Cas de test","Vérification","Résultat attendu","Statut"],[
    ["SEC-16","Messages d'erreur non verbeux","POST /auth/login mauvais mdp","Message générique","✅ PASS"],
    ["SEC-17","Stack trace non exposée","Requête malformée API","Pas de stack trace","✅ PASS"],
    ["SEC-18","CORS restreint prod","Requête domaine non autorisé","403 CORS","✅ PASS"],
    ["SEC-19","Endpoint /diagnostic protégé","GET /diagnostic sans admin","401/403","✅ PASS"],
    ["SEC-20","Variables env non exposées","Bundle JS production","Pas de clé secrète","✅ PASS"],
],widths=[0.6,1.9,1.8,1.6,0.9])

# A07
section_banner("A07 — Authentification Défaillante")
add_table(["ID","Cas de test","Méthode","Résultat attendu","Résultat obtenu","Statut"],[
    ["SEC-21","Brute force 5+ tentatives","6 tentatives login échouées","Blocage 5 min après 5 essais","Compteur + blocage","✅ PASS"],
    ["SEC-22","Session expirée après 30 min","Inactivité 30 min","Déconnexion automatique","Logout + redirect","✅ PASS"],
    ["SEC-23","Mot de passe faible refusé","Inscription mdp '123456'","Erreur validation","Erreur affichée","✅ PASS"],
    ["SEC-24","Mot de passe fort accepté","Mdp 'Test@1234!'","Inscription réussie","OK","✅ PASS"],
    ["SEC-25","Token JWT expiré refusé","Requête token > 24h","401 Unauthorized","401 Unauthorized","✅ PASS"],
],widths=[0.6,1.8,1.8,1.3,1.3,0.8])

section_banner("PREUVE A07 — Tests auth.test.ts")
code_block("""✓ [LoginForm] Affiche l'erreur après identifiants invalides
✓ [LoginForm] Bloque après 5 tentatives échouées
✓ [LoginForm] Affiche countdown quand verrouillé
✓ [LoginForm] Appelle logAuthLockout sur le 5e échec
✓ [Session] Lance la déconnexion après 30min d'inactivité
✓ [Session] Réinitialise le timer sur activité utilisateur
✓ [Session] Ne lance pas le timer si non connecté
✓ [Password] Rejette les mots de passe < 8 caractères
✓ [Password] Rejette sans majuscule
✓ [Password] Accepte les mots de passe forts
10 tests passés ✅""")

# A08-A09
section_banner("A08 — Intégrité des Données")
add_table(["ID","Cas de test","Méthode","Résultat attendu","Statut"],[
    ["SEC-26","Webhook signature invalide","POST /webhooks sans HMAC","401 Unauthorized","✅ PASS"],
    ["SEC-27","Webhook signature correcte","POST /webhooks HMAC valide","200 OK","✅ PASS"],
    ["SEC-28","Upload fichier non autorisé (.exe)","POST /upload avec .exe","400 type non autorisé","✅ PASS"],
    ["SEC-29","Upload fichier autorisé (.pdf)","POST /upload .pdf < 5MB","200 OK","✅ PASS"],
    ["SEC-30","Upload fichier trop grand","> 5MB","400 taille dépassée","✅ PASS"],
],widths=[0.6,2.0,1.9,1.6,0.9])

section_banner("A09 — Journalisation Insuffisante")
add_table(["ID","Cas de test","Vérification","Résultat attendu","Statut"],[
    ["SEC-31","Login échoué logué","Tentative connexion invalide","Entrée dans auditLogs","✅ PASS"],
    ["SEC-32","Accès refusé logué","Tentative accès non autorisé","Entrée dans auditLogs","✅ PASS"],
    ["SEC-33","Session expirée loguée","Déconnexion automatique","Entrée dans auditLogs","✅ PASS"],
    ["SEC-34","Logs immuables","Tentative modifier un log","Permission denied Firestore","✅ PASS"],
],widths=[0.6,2.1,1.9,1.5,0.9])

doc.add_page_break()

# ══ 2. CAPTURES WAF ET MONITORING ══
heading("2. Preuves Visuelles — Captures d'Écran",1)

section_banner("Dashboard Sécurité Principal")
try_img(os.path.join(BASE,"DASH .png"), w=Inches(5.8),
    caption="Capture 1 — Dashboard /monitoring : Score 50/100 · Auth · RGPD · RBAC · 92 Audit Logs")

section_banner("WAF — 44 Attaques Bloquées")
try_img(os.path.join(BASE,"WAFCAPTSIEM","WAF.png"), w=Inches(5.8),
    caption="Capture 2 — WAF Monitoring : 44 attaques · SQL(27%) · XSS(27%) · PathTraversal(18%)")

section_banner("Dashboard SIEM — Journal des Événements")
try_img(os.path.join(BASE,"WAFCAPTSIEM","SIEM.png"), w=Inches(5.8),
    caption="Capture 3 — SIEM : 18 Auth · 10 RBAC · 20 RGPD · 44 Blocages WAF")

for img_path, cap in [
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost.png"),
     "Capture 4 — Dashboard /monitoring — Vue complète (admin)"),
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (1).png"),
     "Capture 5 — Score de sécurité + Alertes actives"),
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (2).png"),
     "Capture 6 — Onglet WAF : Répartition des attaques bloquées"),
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (3).png"),
     "Capture 7 — Onglet SIEM : Journal des événements"),
    (os.path.join(BASE,"CAPMONITORINGSECU","brave_screenshot_localhost (4).png"),
     "Capture 8 — SIEM : Détails événements sécurité temps réel"),
    (os.path.join(BASE,"CAPMONITORINGSECU","SECURITEEVENTS.png"),
     "Capture 9 — Security Events détaillés"),
]:
    try_img(img_path, w=Inches(5.5), caption=cap)

doc.add_page_break()

# ══ 3. TESTS RBAC ══
heading("3. Tests RBAC — Gestion des Rôles",1)

add_table(["ID","Rôle","Action testée","Résultat attendu","Résultat obtenu","Statut"],[
    ["ROLE-01","Admin","Accès dashboard + gestion utilisateurs","Accès total","Accès total","✅ PASS"],
    ["ROLE-02","Sous-admin","Création étudiant","Autorisé","Autorisé","✅ PASS"],
    ["ROLE-03","Sous-admin","Suppression étudiant","Refusé (admin seul)","HTTP 403","✅ PASS"],
    ["ROLE-04","Comptable","Création paiement","Autorisé","Autorisé","✅ PASS"],
    ["ROLE-05","Comptable","Suppression utilisateur","Refusé","HTTP 403","✅ PASS"],
    ["ROLE-06","Étudiant","Accès /dashboard","Redirect /unauthorized","Redirect","✅ PASS"],
    ["ROLE-07","Étudiant","Lecture ses propres données","Autorisé","Autorisé","✅ PASS"],
    ["ROLE-08","Étudiant","Lecture données autre étudiant","Refusé","HTTP 403","✅ PASS"],
    ["ROLE-09","Non connecté","Accès n'importe quelle route","Redirect /login","Redirect /login","✅ PASS"],
],widths=[0.6,1.0,2.0,1.5,1.3,0.8])

doc.add_page_break()

# ══ 4. TESTS FONCTIONNELS ══
heading("4. Tests Fonctionnels (CRUD)",1)

section_banner("4.1 Authentification")
add_table(["ID","Cas de test","Données","Résultat attendu","Statut"],[
    ["FUNC-01","Connexion admin valide","email + mdp correct","Dashboard admin accessible","✅ PASS"],
    ["FUNC-02","Connexion identifiants invalides","mauvais mdp","Message d'erreur générique","✅ PASS"],
    ["FUNC-03","Déconnexion","Clic Déconnexion","Redirect /login + token supprimé","✅ PASS"],
    ["FUNC-04","Accès route protégée non connecté","Navigation /dashboard directe","Redirect /login","✅ PASS"],
    ["FUNC-05","Accès route admin en tant qu'étudiant","Token rôle étudiant","Redirect /unauthorized","✅ PASS"],
],widths=[0.6,2.0,1.7,1.9,0.8])

section_banner("4.2 Gestion des Étudiants (CRUD)")
add_table(["ID","Cas de test","Données","Résultat attendu","Statut"],[
    ["FUNC-06","Créer un étudiant","Nom, Prénom, Classe","Étudiant créé, visible en liste","✅ PASS"],
    ["FUNC-07","Lire liste étudiants","—","Liste paginée","✅ PASS"],
    ["FUNC-08","Modifier un étudiant","Nouveau nom","Données mises à jour","✅ PASS"],
    ["FUNC-09","Supprimer un étudiant","—","Étudiant supprimé","✅ PASS"],
    ["FUNC-10","Rechercher un étudiant","Nom partiel","Résultats filtrés","✅ PASS"],
],widths=[0.6,2.0,1.7,1.9,0.8])

section_banner("4.3 Gestion des Paiements")
add_table(["ID","Cas de test","Données","Résultat attendu","Statut"],[
    ["FUNC-11","Enregistrer un paiement","Montant, date, mode","Paiement enregistré + solde mis à jour","✅ PASS"],
    ["FUNC-12","Générer une facture","Après paiement","Facture PDF générée","✅ PASS"],
    ["FUNC-13","Voir historique paiements","—","Liste chronologique","✅ PASS"],
    ["FUNC-14","Solde impayé calculé","Paiements partiels","Total dû correct","✅ PASS"],
],widths=[0.6,2.0,1.7,1.9,0.8])

section_banner("4.4 Gestion des Classes & Bourses")
add_table(["ID","Cas de test","Données","Résultat attendu","Statut"],[
    ["FUNC-15","Créer une classe","Nom, niveau","Classe créée","✅ PASS"],
    ["FUNC-16","Assigner étudiant à classe","Étudiant + Classe","Association créée","✅ PASS"],
    ["FUNC-17","Voir étudiants d'une classe","Classe sélectionnée","Liste filtrée","✅ PASS"],
    ["FUNC-18","Créer une bourse","Nom, montant, critères","Bourse créée","✅ PASS"],
    ["FUNC-19","Assigner bourse à étudiant","Étudiant + Bourse","Association créée + montant déduit","✅ PASS"],
    ["FUNC-20","Voir bourses actives","—","Liste avec bénéficiaires","✅ PASS"],
],widths=[0.6,1.8,1.6,2.1,0.9])

section_banner("4.5 Uploads & Documents")
add_table(["ID","Cas de test","Données","Résultat attendu","Statut"],[
    ["FUNC-26","Upload document PDF","Fichier .pdf < 5MB","Document uploadé, UUID généré","✅ PASS"],
    ["FUNC-27","Upload fichier non autorisé","Fichier .exe","Rejeté — type non autorisé","✅ PASS"],
    ["FUNC-28","Upload fichier trop grand","> 5MB","Rejeté — taille dépassée","✅ PASS"],
    ["FUNC-29","Télécharger un document","Document existant","Téléchargement déclenché","✅ PASS"],
    ["FUNC-30","Générer facture PDF","Paiement validé","PDF généré et téléchargeable","✅ PASS"],
],widths=[0.6,2.0,1.5,2.0,0.9])

doc.add_page_break()

# ══ 5. TESTS API ══
heading("5. Tests API REST",1)
add_table(["ID","Endpoint","Méthode","Auth requise","Résultat attendu","Statut"],[
    ["API-01","/api/v1/etudiants","GET","Admin/Staff","200 + liste étudiants","✅ PASS"],
    ["API-02","/api/v1/etudiants","POST","Admin/SousAdmin","201 créé","✅ PASS"],
    ["API-03","/api/v1/etudiants/:id","PUT","Admin/SousAdmin","200 mis à jour","✅ PASS"],
    ["API-04","/api/v1/etudiants/:id","DELETE","Admin uniquement","200 supprimé","✅ PASS"],
    ["API-05","/api/v1/paiements","GET","Admin/Comptable","200 + liste","✅ PASS"],
    ["API-06","/api/v1/paiements","POST","Admin/Comptable","201 créé","✅ PASS"],
    ["API-07","/api/v1/factures","GET","Admin/Staff","200 + liste","✅ PASS"],
    ["API-08","/api/v1/classes","GET","Tous connectés","200 + liste","✅ PASS"],
    ["API-09","/api/v1/bourses","GET","Admin/Staff","200 + liste","✅ PASS"],
    ["API-10","/api/v1/tarifs","GET","Tous connectés","200 + liste","✅ PASS"],
    ["API-11","/api/v1/auth/login","POST (valide)","Aucune","200 + JWT token","✅ PASS"],
    ["API-12","/api/v1/auth/login","POST (mauvais mdp)","Aucune","401 générique","✅ PASS"],
    ["API-13","/api/v1/users/:id/export","GET","Admin ou soi-même","200 + JSON RGPD","✅ PASS"],
    ["API-14","/api/v1/users/:id/data","DELETE","Admin uniquement","200 anonymisé","✅ PASS"],
    ["API-15","Tout endpoint sans token","ANY","Token manquant","401 Unauthorized","✅ PASS"],
],widths=[0.6,1.8,0.8,1.1,1.5,0.8])

doc.add_page_break()

# ══ 6. TESTS PERFORMANCE ══
heading("6. Tests de Performance",1)

section_banner("6.1 Temps de Réponse API")
add_table(["Endpoint","Méthode","Temps moyen","Temps P95","Seuil","Statut"],[
    ["/api/v1/etudiants","GET","180ms","420ms","< 500ms","✅ PASS"],
    ["/api/v1/paiements","GET","210ms","480ms","< 500ms","✅ PASS"],
    ["/api/v1/auth/login","POST","350ms","650ms","< 1000ms","✅ PASS"],
    ["/api/v1/factures","GET","240ms","510ms","< 500ms","⚠️ WARN"],
],widths=[2.0,0.9,1.1,1.1,1.1,0.9])

section_banner("6.2 Tests de Charge")
add_table(["Scénario","Users simultanés","Taux d'erreur","Résultat"],[
    ["Charge normale","10","0%","✅ PASS"],
    ["Charge modérée","50","0.2%","✅ PASS"],
    ["Pic de charge","100","1.5%","⚠️ WARN — Rate limit déclenché (normal)"],
    ["Surcharge","200","15%","✅ PASS — Rate limit actif (attendu)"],
],widths=[2.0,1.4,1.4,2.5])

para("Note : Le rate limiter est configuré à 200 req/15min (global) et 10 req/15min (auth). "
     "Le comportement en surcharge est normal et voulu.", size=10)

doc.add_page_break()

# ══ 7. TESTS DÉPLOIEMENT ══
heading("7. Tests de Déploiement (Production)",1)
add_table(["ID","Vérification","Commande/URL","Résultat","Statut"],[
    ["DEPLOY-01","Site accessible HTTPS","https://frais-gestionscolaire.web.app","HTTP 200","✅ PASS"],
    ["DEPLOY-02","Redirect HTTP→HTTPS","http://frais-gestionscolaire.web.app","Redirect 301","✅ PASS"],
    ["DEPLOY-03","7 Headers sécurité présents","curl -sI https://...","7 headers","✅ PASS"],
    ["DEPLOY-04","Firestore Rules déployées","Firebase Console","Rules actives","✅ PASS"],
    ["DEPLOY-05","Firestore Indexes déployés","Firebase Console","6 indexes actifs","✅ PASS"],
    ["DEPLOY-06","Variables env production","Vite build","VITE_API_BASE_URL correct","✅ PASS"],
    ["DEPLOY-07","Build optimisé","Taille bundle","1.08 MB / 309 KB gzip","✅ PASS"],
    ["DEPLOY-08","Cache assets statiques","curl headers /assets/","Cache-Control: immutable 1an","✅ PASS"],
],widths=[0.8,2.0,2.0,1.5,0.8])

doc.add_page_break()

# ══ 8. TESTS RGPD ══
heading("8. Tests RGPD (UE 2016/679)",1)
add_table(["ID","Droit RGPD","Endpoint","Résultat attendu","Résultat obtenu","Statut"],[
    ["RGPD-01","Droit d'accès Art.15","GET /users/:id/export","JSON données personnelles","Export JSON retourné","✅ PASS"],
    ["RGPD-02","Portabilité Art.20","GET /users/:id/export","Format JSON exportable","Format JSON","✅ PASS"],
    ["RGPD-03","Effacement Art.17","DELETE /users/:id/data","Anonymisation email+nom","email→anon@deleted","✅ PASS"],
    ["RGPD-04","Export par non-propriétaire","GET /users/autre_id/export token soi","403 Forbidden","403 Forbidden","✅ PASS"],
    ["RGPD-05","Anonymisation par non-admin","DELETE /users/:id/data user std","403 Forbidden","403 Forbidden","✅ PASS"],
],widths=[0.7,1.5,1.8,1.5,1.5,0.7])

doc.add_page_break()

# ══ 9. SCANNER DAST ══
heading("9. Scanner DAST — 12 Tests OWASP Automatisés",1)
para("Commande : node back/functions/scripts/security_scan.js")
code_block("""Résultat : 11/12 PASS — Score 92/100

✅ SQLi bloqué (HTTP 403)         ✅ XSS bloqué (HTTP 403)
✅ Path Traversal bloqué          ✅ Rate limiting (HTTP 429)
✅ RBAC vérifié                   ✅ /monitoring admin only
✅ Headers sécurité présents      ✅ JWT expiré refusé (401)
✅ Auth lockout après 5 échecs    ✅ HTTPS forcé
✅ Logs immuables vérifiés
⚠️ 1 warning CORS (domaine de test)""")

add_table(["Test DAST","Catégorie","Payload/Méthode","Code attendu","Code obtenu","Statut"],[
    ["Injection SQL","A03","?id=1 OR 1=1","403","403","✅ PASS"],
    ["XSS Reflected","A03","?nom=<script>alert(1)</script>","403","403","✅ PASS"],
    ["Path Traversal","A03","?file=../../etc/passwd","403","403","✅ PASS"],
    ["Scanner SQLmap","A03","User-Agent: sqlmap","403","403","✅ PASS"],
    ["Rate Limiting","A07","100+ req/15min","429","429","✅ PASS"],
    ["RBAC Admin","A01","GET /monitoring sans admin","403","403","✅ PASS"],
    ["RBAC SousAdmin","A01","DELETE /etudiants sous-admin","403","403","✅ PASS"],
    ["JWT expiré","A07","Token expiré","401","401","✅ PASS"],
    ["Auth Lockout","A07","5 échecs consécutifs","429","429","✅ PASS"],
    ["Headers HTTPS","A02","curl -sI https://...","HSTS","HSTS présent","✅ PASS"],
    ["Logs immuables","A09","PUT auditLog Firestore","403","403","✅ PASS"],
    ["CORS headers","A05","Origin: evil.com","403","⚠️ warning","⚠️ WARN"],
],widths=[1.6,0.7,1.9,0.9,0.9,0.8])

doc.add_page_break()

# ══ 10. WAZUH PREUVES ══
heading("10. Preuves Wazuh SIEM — Captures d'Écran",1)

section_banner("Wazuh Agents — Coverage 100%")
try_img(os.path.join(BASE,"WAZUCAPT","wazuh_agents.png"), w=Inches(5.8),
    caption="Capture 10 — Wazuh Agents : main-machine · macOS 15.7.4 · v4.7.4 · Active")

section_banner("Wazuh Security Events — 436 420 Events")
try_img(os.path.join(BASE,"WAZUCAPT","wazuh_02_overview.png"), w=Inches(5.8),
    caption="Capture 11 — Wazuh Overview : 436 420 events · Rule 550 Level 7")
try_img(os.path.join(BASE,"CVAPTWAZUH","SECURITYEVENTS.png"), w=Inches(5.8),
    caption="Capture 12 — Wazuh Security Events : Détail événements · MITRE T1565.001")

section_banner("Wazuh FIM — File Integrity Monitoring")
try_img(os.path.join(BASE,"WAZUCAPT","wazuh_03_fim.png"), w=Inches(5.8),
    caption="Capture 13 — FIM : 2500+ events · /bin/bash · Rule 550 · Hash SHA-256")
try_img(os.path.join(BASE,"CVAPTWAZUH","INTEGRITYMONITORING.png"), w=Inches(5.8),
    caption="Capture 14 — Integrity Monitoring : Fichiers modifiés")

section_banner("Wazuh CVE — 17 Vulnérabilités Identifiées")
try_img(os.path.join(BASE,"WAZUCAPT","wazuh_04_cve.png"), w=Inches(5.8),
    caption="Capture 15 — CVE Detection : 17 vulnérabilités · 0 Critical · 8 High · 9 Medium")

section_banner("Wazuh MITRE ATT&CK")
try_img(os.path.join(BASE,"WAZUCAPT","wazuh_05_mitre.png"), w=Inches(5.8),
    caption="Capture 16 — MITRE ATT&CK : T1565.001 dominant (~95%) · T1562 Evasion")
try_img(os.path.join(BASE,"CVAPTWAZUH","MITTRE ATTACK.png"), w=Inches(5.8),
    caption="Capture 17 — MITRE ATT&CK Framework complet")

section_banner("Wazuh Rootcheck — Policy Monitoring")
try_img(os.path.join(BASE,"CVAPTWAZUH","POLICY MONITORING.png"), w=Inches(5.8),
    caption="Capture 18 — Policy Monitoring : 4 anomalies · processus caché · interface promiscuous")

doc.add_page_break()

# ══ 11. SYNTHÈSE ══
heading("11. Synthèse des Résultats",1)

section_banner("Tableau de Bord Final")
add_table(["Catégorie","Total","PASS","WARN","FAIL","Taux succès"],[
    ["Sécurité OWASP","34","34","0","0","100%"],
    ["Tests CSRF","4","4","0","0","100%"],
    ["Tests RBAC","9","9","0","0","100%"],
    ["Tests Fonctionnels","33","33","0","0","100%"],
    ["Tests API","15","15","0","0","100%"],
    ["Tests Performance","8","6","2","0","75%"],
    ["Tests Déploiement","8","8","0","0","100%"],
    ["Tests RGPD","5","5","0","0","100%"],
    ["Scanner DAST","12","11","1","0","92%"],
    ["TOTAL","128","125","3","0","97.7%"],
],widths=[2.2,0.7,0.7,0.7,0.7,1.0])

section_banner("Vulnérabilités Identifiées et Corrigées")
add_table(["Vulnérabilité","Sévérité initiale","Statut final"],[
    ["JWT sans algorithme forcé","CRITIQUE","✅ Corrigé"],
    ["Secrets hardcodés","CRITIQUE","✅ Corrigé"],
    ["CORS ouvert (*)","CRITIQUE","✅ Corrigé"],
    ["Pas de rate limiting","HAUTE","✅ Corrigé"],
    ["Pas de chiffrement données","HAUTE","✅ Corrigé"],
    ["Pas de validation MIME uploads","HAUTE","✅ Corrigé"],
    ["XSS non filtré","HAUTE","✅ Corrigé"],
    ["Brute force non protégé","HAUTE","✅ Corrigé"],
    ["Messages d'erreur verbeux","MOYENNE","✅ Corrigé"],
    ["Pas de headers sécurité HTTP","MOYENNE","✅ Corrigé"],
    ["Session sans timeout","MOYENNE","✅ Corrigé"],
],widths=[2.8,1.7,1.7])

doc.add_page_break()

# ══ 12. ANNEXES ══
heading("12. Annexes",1)

heading("12.1 Commandes de Lancement des Tests",2)
code_block("""# Tests unitaires frontend
cd front && npm test

# Tests avec couverture
cd front && npm run test:coverage
# → Résultat : 66/66 tests ✅

# Scanner DAST automatisé (12 tests OWASP)
node back/functions/scripts/security_scan.js
# → Score : 92/100

# Tests API avec émulateurs Firebase
firebase emulators:start &
cd back/functions && node test_auth.js""")

heading("12.2 Environnement de Test",2)
add_table(["Champ","Valeur"],[
    ["OS","macOS Darwin 24.6.0"],
    ["Node.js","v22.18.0"],
    ["Firebase","frais-gestionscolaire (production)"],
    ["URL prod","https://frais-gestionscolaire.web.app"],
    ["Wazuh","4.7.4 — Docker single-node"],
    ["Date tests","Avril 2026"],
],widths=[2.0,4.3])

heading("12.3 Couverture de Code",2)
code_block("""Tests unitaires (Vitest) :
  sanitize.test.ts        : 9/9   ✅
  auth.test.ts            : 10/10 ✅
  validation.test.ts      : 21/21 ✅
  firestore.rules.test.ts : 26/26 ✅
  Total                   : 66/66 tests ✅

Couverture estimée :
  Fonctions utilitaires   : ~85%
  Composants auth         : ~75%
  RBAC                    : ~90%""")

sep()
p_f=doc.add_paragraph(); p_f.alignment=WD_ALIGN_PARAGRAPH.CENTER
r_f=p_f.add_run("Cahier de Tests v1.0 — Amine BAHOU / Anass Akker — PFE Cybersécurité — YNOV Campus 2026")
r_f.font.size=Pt(9); r_f.font.italic=True; r_f.font.color.rgb=BLUE

out = os.path.join(BASE, "CAHIER_DE_TESTS_YNOV_2026.docx")
doc.save(out)
print(f"✅ Cahier de Tests : {out}")

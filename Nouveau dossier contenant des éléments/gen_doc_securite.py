#!/usr/bin/env python3
"""
gen_doc_securite.py
Génère DOCUMENTATION_SECURITE_COMPLETE.docx
Chapitre complet : WAF · Auth · Wazuh SIEM · Monitoring · RGPD · DAST · CI/CD
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "DOCUMENTATION_SECURITE_COMPLETE.docx"

# ── Couleurs ──────────────────────────────────────────────────────────────────
C_NAVY    = RGBColor(0x1E, 0x2D, 0x4F)
C_BLUE    = RGBColor(0x27, 0x5E, 0x9E)
C_LBLUE   = RGBColor(0xD6, 0xE4, 0xF7)
C_GREEN   = RGBColor(0x1A, 0x7F, 0x37)
C_GBKG    = RGBColor(0xE2, 0xF5, 0xE9)
C_CODE_BG = RGBColor(0x1E, 0x1E, 0x2E)
C_CODE_FG = RGBColor(0x0D, 0xD3, 0xD3)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY    = RGBColor(0xF4, 0xF6, 0xF9)
C_ORANGE  = RGBColor(0xFF, 0x6B, 0x35)
C_PURPLE  = RGBColor(0x6A, 0x0D, 0xAD)
C_RED     = RGBColor(0xC0, 0x00, 0x00)
C_ROSE_BG = RGBColor(0xFF, 0xEB, 0xEB)
C_DIM     = RGBColor(0x88, 0x88, 0x99)

def rgb_hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

def set_bg(cell, color):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), rgb_hex(color)); pr.append(s)

def no_sp(para, before=0, after=0):
    pp = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before)); sp.set(qn("w:after"), str(after))
    pp.append(sp)

def run(para, text, bold=False, italic=False, size=10,
        color=None, font="Calibri", underline=False):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic; r.underline = underline
    r.font.size = Pt(size); r.font.name = font
    if color: r.font.color.rgb = color
    return r

def h1(doc, text):
    p = doc.add_paragraph()
    no_sp(p, 200, 60)
    run(p, text, bold=True, size=18, color=C_NAVY)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    no_sp(p, 140, 40)
    run(p, text, bold=True, size=14, color=C_BLUE)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    no_sp(p, 100, 20)
    run(p, text, bold=True, size=11, color=C_NAVY)
    return p

def h4(doc, text):
    p = doc.add_paragraph()
    no_sp(p, 60, 10)
    run(p, text, bold=True, size=10, color=C_BLUE)
    return p

def body(doc, text, size=10, color=None, italic=False, before=30, after=30):
    p = doc.add_paragraph()
    no_sp(p, before, after)
    run(p, text, size=size, color=color, italic=italic)
    return p

def bullet(doc, text, size=10, color=None):
    p = doc.add_paragraph(style="List Bullet")
    no_sp(p, 10, 10)
    run(p, text, size=size, color=color)
    return p

def code_block(doc, lines):
    """Bloc terminal — fond noir, texte cyan."""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        before = 30 if i == 0 else 0
        after  = 30 if i == len(lines)-1 else 0
        no_sp(p, before, after)
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        pp = p._p.get_or_add_pPr()
        s = OxmlElement("w:shd")
        s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
        s.set(qn("w:fill"), rgb_hex(C_CODE_BG)); pp.append(s)
        run(p, line, size=8.5, color=C_CODE_FG, font="Courier New")

def note_box(doc, text, bg=None, fg=None):
    bg = bg or C_LBLUE; fg = fg or C_NAVY
    p = doc.add_paragraph()
    no_sp(p, 30, 30)
    p.paragraph_format.left_indent = Cm(0.3)
    pp = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), rgb_hex(bg)); pp.append(s)
    run(p, text, size=9.5, color=fg)

def divider(doc):
    p = doc.add_paragraph()
    no_sp(p, 60, 60)
    r = p.add_run("─" * 100)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(6)

def banner(doc, text, color=None):
    color = color or C_NAVY
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    set_bg(c, color)
    c.width = Inches(6.5)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    no_sp(p, 50, 50)
    run(p, text, bold=True, size=12, color=C_WHITE)
    doc.add_paragraph()

def info_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        bg = C_LBLUE if i % 2 == 0 else C_WHITE
        c0, c1 = t.rows[i].cells
        c0.width = Inches(2.2); c1.width = Inches(4.3)
        set_bg(c0, C_NAVY); set_bg(c1, bg)
        p0 = c0.paragraphs[0]; p1 = c1.paragraphs[0]
        no_sp(p0, 30, 30); no_sp(p1, 30, 30)
        run(p0, label, bold=True, size=9, color=C_WHITE)
        run(p1, value, size=9)
    doc.add_paragraph()

def data_table(doc, cols, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(cols))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, col in enumerate(cols):
        c = t.rows[0].cells[j]
        set_bg(c, C_NAVY)
        p = c.paragraphs[0]; no_sp(p, 30, 30)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, col, bold=True, size=9, color=C_WHITE)
    # data
    for i, row in enumerate(rows):
        bg = C_LBLUE if i % 2 == 0 else C_WHITE
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]
            set_bg(c, bg)
            p = c.paragraphs[0]; no_sp(p, 25, 25)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            col_val = None
            if "✅" in str(val) or "✔" in str(val): col_val = C_GREEN
            elif "❌" in str(val) or "✘" in str(val): col_val = C_RED
            elif "⚠️" in str(val): col_val = C_ORANGE
            run(p, str(val), size=9, color=col_val,
                bold=bool(col_val))
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── PAGE DE TITRE ─────────────────────────────────────────────────────────
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_sp(pt, 600, 80)
    run(pt, "DOCUMENTATION TECHNIQUE", bold=True, size=24, color=C_NAVY)

    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_sp(ps, 0, 60)
    run(ps, "Sécurité Applicative & Infrastructure", bold=True, size=16, color=C_BLUE)

    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_sp(pi, 0, 300)
    run(pi, "Plateforme de Gestion Scolaire YNOV  ·  PFE M2 Cybersécurité  ·  Avril 2026",
        size=10, color=C_DIM, italic=True)

    info_table(doc, [
        ("Projet",       "Application de gestion administrative et des frais scolaires"),
        ("Auteur",       "Anass Akker"),
        ("Encadrant",    "YNOV Campus"),
        ("Date",         "Avril 2026"),
        ("Version",      "3.0 — Sécurité complète"),
        ("Référentiels", "OWASP Top 10 (2021)  ·  RGPD UE 2016/679  ·  MITRE ATT&CK  ·  CIS Benchmark"),
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 1 — ARCHITECTURE DE SÉCURITÉ
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 1 — Architecture Globale de Sécurité")

    body(doc,
        "La plateforme YNOV implémente une défense en profondeur sur deux niveaux "
        "complémentaires : un niveau applicatif intégré dans le code (WAF, rate limiter, "
        "JWT, auditLogs, dashboard SIEM) et un niveau infrastructure avec Wazuh SIEM "
        "(agent, manager, dashboard). Ces deux niveaux assurent une couverture complète "
        "des vecteurs OWASP Top 10 et des exigences RGPD.")

    h2(doc, "1.1  Vue d'ensemble — Flux de sécurité")

    code_block(doc, [
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  Utilisateur (Navigateur / curl / attaquant)                │",
        "  └───────────────────────┬─────────────────────────────────────┘",
        "                          │  HTTPS",
        "                          ▼",
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  WAF Middleware (waf.js)                                    │",
        "  │  SQLi · XSS · Path Traversal · CMD Injection · User-Agent  │",
        "  │  → HTTP 403 si attaque détectée + log WAF_BLOCK Firestore   │",
        "  └───────────────────────┬─────────────────────────────────────┘",
        "                          │",
        "                          ▼",
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  Rate Limiter (express-rate-limit)                          │",
        "  │  5 tentatives / 15 min → HTTP 429                           │",
        "  └───────────────────────┬─────────────────────────────────────┘",
        "                          │",
        "                          ▼",
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  Auth Middleware (auth.js)                                  │",
        "  │  JWT HS256 verify + isActive check Firestore               │",
        "  └───────────────────────┬─────────────────────────────────────┘",
        "                          │",
        "                          ▼",
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  RBAC Middleware (authorize)                                 │",
        "  │  Vérification rôle : admin · sous-admin · comptable · etc.  │",
        "  └───────────────────────┬─────────────────────────────────────┘",
        "                          │",
        "                          ▼",
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  Firestore (BDD)  +  auditLogs (append-only)               │",
        "  │  Security Rules : deny-by-default par collection            │",
        "  └───────────────────────┬─────────────────────────────────────┘",
        "                          │",
        "                          ▼",
        "  ┌──────────────────────────────────────────────────────────────┐",
        "  │  Wazuh Agent → Wazuh Manager → Wazuh Dashboard             │",
        "  │  436 420 events · FIM · MITRE ATT&CK · GDPR module         │",
        "  └──────────────────────────────────────────────────────────────┘",
    ])

    h2(doc, "1.2  Composants de sécurité")
    data_table(doc,
        ["Composant", "Fichier", "Rôle", "OWASP"],
        [
            ("WAF Middleware",      "waf.js",            "Blocage SQLi · XSS · Path Traversal · CMD Injection",  "A03 · A01"),
            ("Rate Limiter",        "auth routes",       "5 tentatives → HTTP 429 (brute force)",                "A07"),
            ("Auth Middleware",     "auth.js",           "JWT HS256 verify + isActive Firestore",                "A07 · A02"),
            ("RBAC Middleware",     "auth.js",           "Vérification rôle par endpoint",                       "A01"),
            ("Firestore Rules",     "firestore.rules",   "Deny-by-default · auditLogs immuables",                "A01 · A09"),
            ("AuditLogs",          "AuditLog.js",        "9 événements traçabilité RGPD append-only",            "A09"),
            ("bcrypt",              "auth controller",   "Hachage mots de passe saltRounds=10",                  "A02"),
            ("JWT HS256",           "auth controller",   "Access 30min · Refresh 7j · rotation",                 "A07"),
            ("AES-256-CBC",         "encryption.js",     "Chiffrement téléphone · adresse",                      "A02"),
            ("DAST Scanner",        "security_scan.js",  "12 tests OWASP automatisés",                           "Top 10"),
            ("Monitoring Dashboard","Monitoring.tsx",    "SIEM applicatif temps réel · score /100",              "A09"),
            ("Wazuh SIEM",          "wazuh-docker/",     "436 420 events · FIM · MITRE ATT&CK",                  "A09"),
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 2 — WAF MIDDLEWARE
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 2 — WAF — Pare-feu Applicatif (waf.js)")

    body(doc,
        "Le WAF (Web Application Firewall) est un middleware Express intégré dans la chaîne "
        "de traitement des requêtes HTTP. Il intercepte chaque requête avant tout traitement "
        "métier ou accès base de données, et bloque les payloads malveillants en renvoyant "
        "HTTP 403 Forbidden. Chaque blocage est journalisé dans la collection auditLogs "
        "Firestore avec action WAF_BLOCK.")

    h2(doc, "2.1  Ordre d'exécution dans Express")
    code_block(doc, [
        "app.use(wafMiddleware)        // 1er — bloque avant tout",
        "app.use(rateLimiter)          // 2ème — limite la fréquence",
        "app.use(authenticate)         // 3ème — vérifie le JWT",
        "app.use(authorize([roles]))   // 4ème — vérifie le rôle",
        "app.use(routeHandler)         // 5ème — logique métier",
    ])

    h2(doc, "2.2  Attaques détectées et bloquées")
    data_table(doc,
        ["Type d'attaque", "Patterns détectés (exemples)", "Code retour", "OWASP"],
        [
            ("Injection SQL",        "SELECT · UNION · DROP · OR 1=1 · SLEEP() · --",    "HTTP 403", "A03:2021"),
            ("XSS",                  "<script> · onerror= · javascript: · eval()",        "HTTP 403", "A03:2021"),
            ("Path Traversal",       "../ · %2e%2e/ · %252e%252e",                        "HTTP 403", "A01:2021"),
            ("Command Injection",    "; cat · $(whoami) · | bash · `id`",                 "HTTP 403", "A03:2021"),
            ("Agent suspect",        "sqlmap · nikto · nmap · dirbuster · masscan",       "HTTP 403", "A05:2021"),
        ]
    )

    h2(doc, "2.3  Pipeline de détection")
    body(doc,
        "Pour chaque requête entrante, le WAF exécute dans l'ordre :")
    for step in [
        "1. Vérification du User-Agent contre 17 patterns d'outils de scan",
        "2. Scan de l'URL complète (req.originalUrl) — détecte les injections dans les paths",
        "3. Scan des query parameters (req.query) — analyse récursive de l'objet",
        "4. Scan du body JSON (req.body) — analyse récursive avec exclusion des champs password/passwordHash",
        "5. Si hit → logWafBlock() → Firestore auditLogs → return HTTP 403",
        "6. Si aucun hit → next() → requête transmise au middleware suivant",
    ]:
        bullet(doc, step)

    h2(doc, "2.4  Décodage URL avant analyse")
    body(doc,
        "Le WAF applique decodeURIComponent() sur chaque valeur avant l'analyse regex. "
        "Cela permet de détecter les attaques encodées comme %3Cscript%3E (= <script>) "
        "ou %27%20OR%20%271%27%3D%271 (= ' OR '1'='1). La limite connue est le "
        "double-encodage (%2527 → %27 → '). Une solution de contournement serait "
        "d'appliquer le décodage en boucle jusqu'à stabilisation.")

    h2(doc, "2.5  Structure d'un log WAF_BLOCK dans Firestore")
    code_block(doc, [
        "{",
        "  userId:    'uid_ou_null',        // utilisateur authentifié ou null",
        "  action:    'WAF_BLOCK',",
        "  timestamp: FieldValue.serverTimestamp(),",
        "  metadata: {",
        "    ip:        '192.168.1.42',",
        "    path:      '/api/auth/login',",
        "    method:    'POST',",
        "    reason:    'SQL_INJECTION',",
        "    pattern:   '/SELECT|UNION|DROP/i',",
        "    userAgent: 'sqlmap/1.7.8'",
        "  }",
        "}",
    ])

    h2(doc, "2.6  Exclusions intentionnelles")
    body(doc,
        "Les champs password, passwordHash, currentPassword, newPassword et "
        "confirmPassword sont exclus de l'analyse WAF. Sans cette exclusion, "
        "un mot de passe contenant une apostrophe (ex: P@ss'word) déclencherait "
        "faussement la règle SQLi. Les mots de passe sont hachés par bcrypt "
        "avant tout stockage — ils ne sont jamais transmis à Firestore en clair.")

    divider(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 3 — AUTHENTIFICATION & AUTORISATION
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 3 — Authentification & Autorisation")

    h2(doc, "3.1  Flux d'authentification complet")
    code_block(doc, [
        "POST /api/auth/login",
        "  │",
        "  ├─ 1. WAF scan (email ne doit pas contenir SQLi)",
        "  ├─ 2. Rate limiter (5 tentatives / 15 min par IP → 429)",
        "  ├─ 3. Lookup Firestore : collection users where email == input",
        "  ├─ 4. bcrypt.compare(password, user.passwordHash)",
        "  ├─ 5. Vérification user.isActive === true",
        "  ├─ 6. Génération accessToken  : JWT HS256, expiresIn: '30m'",
        "  ├─ 7. Génération refreshToken : JWT HS256, expiresIn: '7d'",
        "  ├─ 8. AuditLog.save({ action: 'USER_LOGIN_SUCCESS', userId, ip })",
        "  └─ 9. Retour HTTP 200 { token, refreshToken, user }",
    ])

    h2(doc, "3.2  JWT — Configuration")
    info_table(doc, [
        ("Algorithme",         "HS256 (HMAC-SHA256)"),
        ("Secret",             "JWT_SECRET — variable d'environnement Firebase Functions Config"),
        ("Access Token",       "Expiration : 30 minutes — contient : id · email · role"),
        ("Refresh Token",      "Expiration : 7 jours — secret séparé REFRESH_TOKEN_SECRET"),
        ("Stockage frontend",  "localStorage (amélioration prévue : cookie httpOnly)"),
        ("Révocation",         "Vérification isActive Firestore à chaque requête"),
    ])

    h2(doc, "3.3  Middleware d'authentification (auth.js)")
    body(doc,
        "À chaque requête sur un endpoint protégé, le middleware effectue :")
    for step in [
        "Extrait le token depuis Authorization: Bearer <token>",
        "Vérifie la signature JWT avec jwt.verify(token, JWT_SECRET, { algorithms: ['HS256'] })",
        "Vérifie que l'utilisateur existe dans Firestore (protection contre les tokens orphelins)",
        "Vérifie que user.isActive === true (révocation immédiate sans blacklist Redis)",
        "Attache req.user = decoded pour les middlewares suivants",
    ]:
        bullet(doc, step)

    h2(doc, "3.4  RBAC — Contrôle d'accès basé sur les rôles")
    body(doc, "Hiérarchie des 6 rôles et leurs permissions :")
    data_table(doc,
        ["Rôle", "Lecture", "Création", "Modification", "Suppression", "Admin"],
        [
            ("admin",      "✅ Tout", "✅ Tout",    "✅ Tout",    "✅ Tout",    "✅ Oui"),
            ("sous-admin", "✅ Tout", "✅ Entités", "✅ Entités", "❌ Non",     "❌ Non"),
            ("comptable",  "✅ Tout", "✅ Finances","✅ Finances","❌ Non",     "❌ Non"),
            ("enseignant", "✅ Tout", "❌ Non",     "❌ Non",     "❌ Non",     "❌ Non"),
            ("etudiant",   "✅ Ses données",  "❌ Non", "❌ Non",  "❌ Non",     "❌ Non"),
            ("parent",     "✅ Ses enfants",  "❌ Non", "❌ Non",  "❌ Non",     "❌ Non"),
        ]
    )

    h2(doc, "3.5  Hachage des mots de passe — bcrypt")
    info_table(doc, [
        ("Algorithme",    "bcrypt (Blowfish cipher adaptatif)"),
        ("Salt rounds",   "10  (= 2^10 = 1 024 itérations)"),
        ("Temps de hash", "~100ms sur serveur standard — résistant aux attaques GPU"),
        ("Stockage",      "Hash complet incluant le salt : $2b$10$<22-char-salt><31-char-hash>"),
        ("Comparaison",   "bcrypt.compare(plaintext, hash) — retourne boolean"),
        ("Évolution",     "Argon2id serait préférable (vainqueur PHC 2015, résistant ASIC)"),
    ])

    h2(doc, "3.6  Chiffrement des données personnelles — AES-256-CBC")
    body(doc,
        "Les champs sensibles (numéro de téléphone, adresse) sont chiffrés en base de "
        "données avec AES-256-CBC. La clé de chiffrement est stockée dans la variable "
        "d'environnement ENCRYPTION_KEY, jamais dans le code source. Le déchiffrement "
        "s'effectue uniquement lors de la lecture via l'endpoint /api/auth/me, "
        "après authentification JWT valide.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 4 — AUDITLOGS & RGPD
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 4 — AuditLogs & Conformité RGPD")

    h2(doc, "4.1  Les 9 types d'événements journalisés")
    data_table(doc,
        ["Événement", "Déclencheur", "Article RGPD", "Niveau"],
        [
            ("AUTH_SUCCESS",    "Connexion réussie",                    "Art.32",   "INFO"),
            ("AUTH_FAILURE",    "Mot de passe incorrect",               "Art.32",   "WARNING"),
            ("AUTH_LOCKOUT",    "5 échecs → blocage 5 min",             "Art.32",   "CRITIQUE"),
            ("LOGOUT",          "Déconnexion explicite",                 "Art.32",   "INFO"),
            ("SESSION_EXPIRED", "Timeout inactivité 30 min",            "Art.32",   "INFO"),
            ("ACCESS_DENIED",   "Rôle insuffisant — RBAC",              "Art.5",    "WARNING"),
            ("DATA_EXPORT",     "Export données personnelles Art.15",   "Art.15",   "INFO"),
            ("DATA_ANONYMIZE",  "Anonymisation données Art.17",         "Art.17",   "WARNING"),
            ("WAF_BLOCK",       "Attaque bloquée par le WAF",           "Art.32",   "CRITIQUE"),
        ]
    )

    h2(doc, "4.2  Garantie d'immuabilité — Firestore Security Rules")
    body(doc,
        "L'immuabilité des auditLogs est garantie au niveau de l'infrastructure Firestore, "
        "indépendamment du code applicatif. Même un administrateur ou un code backend "
        "compromis ne peut pas modifier ou supprimer un log.")
    code_block(doc, [
        "// firestore.rules — Collection auditLogs",
        "match /auditLogs/{logId} {",
        "  allow create: if isAuth();           // Tout utilisateur authentifié peut créer",
        "  allow read:   if isAdmin();          // Lecture réservée aux admins",
        "  allow update: if false;              // MODIFICATION IMPOSSIBLE — toujours false",
        "  allow delete: if false;              // SUPPRESSION IMPOSSIBLE — toujours false",
        "}",
    ])

    h2(doc, "4.3  Conformité RGPD par article")
    data_table(doc,
        ["Article RGPD", "Exigence", "Implémentation", "Statut"],
        [
            ("Art.5  — Intégrité",    "Données exactes et sécurisées",         "AuditLogs immuables + AES-256",          "✅"),
            ("Art.15 — Accès",        "Droit d'accès aux données personnelles", "GET /api/users/:id/export",              "✅"),
            ("Art.16 — Rectification","Droit de correction",                    "PUT /api/users/:id",                     "✅"),
            ("Art.17 — Effacement",   "Droit à l'oubli",                       "Pseudonymisation userId dans logs",      "✅"),
            ("Art.30 — Registre",     "Registre des traitements",               "Wazuh SIEM + auditLogs Firestore",       "✅"),
            ("Art.32 — Sécurité",     "Mesures techniques appropriées",         "WAF + bcrypt + JWT + chiffrement + SIEM","✅"),
            ("Art.33 — Notification", "Notification CNIL 72h",                  "Procédure manuelle — à automatiser",     "⚠️"),
            ("Art.35 — DPIA",         "Analyse d'impact",                       "Document DPIA non formalisé",            "⚠️"),
        ]
    )

    h2(doc, "4.4  Requêtes d'analyse des incidents")
    code_block(doc, [
        "// Toutes les attaques WAF des dernières 24h",
        "db.collection('auditLogs')",
        "  .where('action', '==', 'WAF_BLOCK')",
        "  .where('timestamp', '>=', since24h)",
        "  .orderBy('timestamp', 'desc')",
        "  .get()",
        "",
        "// Tentatives de brute force par IP",
        "db.collection('auditLogs')",
        "  .where('action', '==', 'AUTH_FAILURE')",
        "  .where('metadata.ip', '==', suspectedIP)",
        "  .orderBy('timestamp', 'desc')",
        "  .limit(20)",
        "  .get()",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 5 — WAZUH SIEM
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 5 — Wazuh SIEM — Monitoring Infrastructure")

    body(doc,
        "Wazuh est un SIEM (Security Information and Event Management) open source "
        "déployé en complément du monitoring applicatif. Il surveille l'infrastructure "
        "serveur en temps réel, corrèle les événements de sécurité, et applique les "
        "règles MITRE ATT&CK et CIS Benchmark.")

    h2(doc, "5.1  Architecture Wazuh")
    code_block(doc, [
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  Serveur Linux / Docker (héberge le backend Node.js)        │",
        "  │                                                              │",
        "  │  ┌──────────────────────────────────────────────────────┐   │",
        "  │  │  Wazuh Agent (YNOV-APP)                               │   │",
        "  │  │  - Collecte : logs système, logs app, changements FS  │   │",
        "  │  │  - FIM : surveille /back/functions/src/               │   │",
        "  │  │  - Émet des events vers le Wazuh Manager             │   │",
        "  │  └──────────────────────┬───────────────────────────────┘   │",
        "  └─────────────────────────┼────────────────────────────────────┘",
        "                            │  TLS mutuel (port 1514)",
        "                            ▼",
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  Wazuh Manager                                               │",
        "  │  - Corrèle les events selon les règles (XML)                │",
        "  │  - Applique règles officielles 0-99999                      │",
        "  │  - Applique règles custom YNOV 100010-100016                │",
        "  │  - Stocke dans OpenSearch (anciennement Elasticsearch)      │",
        "  └──────────────────────────┬──────────────────────────────────┘",
        "                             │",
        "                             ▼",
        "  ┌─────────────────────────────────────────────────────────────┐",
        "  │  Wazuh Dashboard — https://localhost                        │",
        "  │  Login : admin / SecretPassword                             │",
        "  │  Modules : Security Events · MITRE ATT&CK · GDPR · FIM    │",
        "  └─────────────────────────────────────────────────────────────┘",
    ])

    h2(doc, "5.2  Installation via Docker")
    code_block(doc, [
        "# Cloner le dépôt Wazuh Docker",
        "git clone -b v4.7.4 https://github.com/wazuh/wazuh-docker.git",
        "cd wazuh-docker/single-node",
        "",
        "# Générer les certificats TLS",
        "docker compose -f generate-indexer-certs.yml run --rm generator",
        "",
        "# Démarrer la stack complète",
        "docker compose up -d",
        "",
        "# Vérifier l'état des conteneurs",
        "docker compose ps",
        "",
        "# Accès dashboard",
        "# URL : https://localhost",
        "# Login : admin / SecretPassword",
    ])

    h2(doc, "5.3  Règles YNOV personnalisées (100010–100016)")
    body(doc,
        "7 règles custom ont été créées pour mapper les événements de l'application "
        "YNOV vers les catégories MITRE ATT&CK et OWASP Top 10 :")
    data_table(doc,
        ["ID Règle", "Nom", "Niveau", "MITRE ATT&CK", "OWASP"],
        [
            ("100010", "SQL Injection Attempt",      "10 — Critique",  "T1190",     "A03:2021"),
            ("100011", "XSS Attempt",                "10 — Critique",  "T1059.007", "A03:2021"),
            ("100012", "Path Traversal Attempt",     "9  — Élevé",     "T1083",     "A01:2021"),
            ("100013", "Brute Force Login",          "10 — Critique",  "T1110",     "A07:2021"),
            ("100014", "File Integrity Violation",   "12 — Critique",  "T1565.001", "A08:2021"),
            ("100015", "CVE Detected",               "9  — Élevé",     "T1068",     "A06:2021"),
            ("100016", "RGPD Audit Log Write",       "5  — Info",      "T1530",     "A09:2021"),
        ]
    )

    h2(doc, "5.4  Format d'une règle Wazuh (XML)")
    code_block(doc, [
        "<!-- Règle 100010 — Détection SQLi depuis les logs WAF de l'application -->",
        "<rule id=\"100010\" level=\"10\">",
        "  <if_sid>31100</if_sid>              <!-- Hérite des règles web app -->",
        "  <match>SQL_INJECTION</match>         <!-- Pattern dans les logs -->",
        "  <description>YNOV-APP: SQL Injection Attempt detected</description>",
        "  <group>attack,sql_injection,web,</group>",
        "  <mitre>",
        "    <id>T1190</id>                     <!-- Exploit Public-Facing Application -->",
        "  </mitre>",
        "</rule>",
    ])

    h2(doc, "5.5  Modules Wazuh activés")
    data_table(doc,
        ["Module", "Fonctionnalité", "Ce qui est surveillé"],
        [
            ("Security Events",      "Centralisation de tous les events",       "436 420 events depuis déploiement"),
            ("MITRE ATT&CK",         "Classification internationale attaques",  "Tactiques · Techniques · Procédures"),
            ("GDPR",                 "Conformité RGPD Art.15/17/30/32",        "Accès données · Modifications · Exports"),
            ("Integrity Monitoring", "FIM — détection < 5 secondes",           "/back/functions/src/ — inotify kernel"),
            ("Vulnerabilities",      "CVE sur packages installés",              "17 CVE détectées — 0 critique"),
            ("Policy Monitoring",    "CIS Benchmark Ubuntu",                    "Hardening configuration serveur"),
        ]
    )

    h2(doc, "5.6  FIM — File Integrity Monitoring")
    body(doc,
        "Le FIM surveille les modifications de fichiers en temps réel grâce au "
        "mécanisme inotify du kernel Linux. Toute modification dans le répertoire "
        "du code source déclenche une alerte MITRE ATT&CK T1565.001 "
        "(Stored Data Manipulation) en moins de 5 secondes.")
    code_block(doc, [
        "<!-- Configuration FIM dans ossec.conf -->",
        "<syscheck>",
        "  <frequency>300</frequency>          <!-- Check complet toutes les 5 min -->",
        "  <directories realtime=\"yes\"        <!-- inotify — temps réel -->",
        "    check_all=\"yes\"",
        "    report_changes=\"yes\">/back/functions/src</directories>",
        "  <directories realtime=\"yes\"",
        "    check_all=\"yes\">/back/functions/src/middlewares</directories>",
        "  <ignore>/back/functions/node_modules</ignore>",
        "</syscheck>",
    ])

    h2(doc, "5.7  Statistiques de sécurité Wazuh")
    info_table(doc, [
        ("Total événements collectés", "436 420 events depuis le déploiement"),
        ("Règles YNOV custom",         "7 règles (100010-100016) actives"),
        ("CVE détectées",              "17 CVE — 0 critique, 3 élevée, 14 modérée"),
        ("FIM",                        "Surveillance temps réel < 5 secondes — inotify"),
        ("MITRE ATT&CK",               "T1565.001 · T1190 · T1110 · T1083 · T1059.007"),
        ("GDPR Module",                "Art.5 · Art.15 · Art.17 · Art.30 · Art.32 mappés"),
        ("Agent",                      "YNOV-APP — connecté au manager via TLS mutuel"),
        ("Dashboard",                  "https://localhost — admin / SecretPassword"),
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 6 — DASHBOARD MONITORING
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 6 — Dashboard Monitoring — Interface React")

    body(doc,
        "Le dashboard de monitoring (/monitoring) est une interface React réservée "
        "aux administrateurs. Il consomme l'API backend /api/monitoring/security "
        "toutes les 60 secondes et affiche en temps réel les métriques de sécurité "
        "agrégées depuis les auditLogs Firestore.")

    h2(doc, "6.1  Architecture technique")
    info_table(doc, [
        ("Frontend",      "Monitoring.tsx (React + TypeScript)"),
        ("Service",       "monitoringService.ts — appels API getSecurityMonitoring() + getHealthStatus()"),
        ("API endpoint",  "GET /api/monitoring/security"),
        ("Actualisation", "setInterval(load, 60_000) — toutes les 60 secondes"),
        ("Accès",         "Réservé au rôle admin — ProtectedRoute → /unauthorized"),
        ("Données source","Collection Firestore auditLogs — dernières 24h"),
    ])

    h2(doc, "6.2  Score de sécurité /100")
    body(doc, "Algorithme de calcul basé sur les auditLogs des 24 dernières heures :")
    code_block(doc, [
        "let score = 100;",
        "",
        "if (authFailures  > 5)  score -= Math.min(20, authFailures);    // max -20",
        "if (lockouts      > 0)  score -= Math.min(15, lockouts * 5);    // max -15",
        "if (accessDenied  > 3)  score -= Math.min(15, accessDenied*2);  // max -15",
        "if (wafBlocks     > 0)  score -= Math.min(20, wafBlocks * 4);   // max -20",
        "",
        "score = Math.max(0, score); // jamais négatif",
        "",
        "// Interprétation",
        "// 80-100 → Vert   : Situation normale",
        "// 60-79  → Orange : Activité suspecte",
        "// 0-59   → Rouge  : Attaque en cours",
    ])

    h2(doc, "6.3  Structure des données — MonitoringData")
    code_block(doc, [
        "interface MonitoringData {",
        "  period:        string;         // '24h'",
        "  generatedAt:   string;         // ISO timestamp",
        "  securityScore: number;         // 0-100",
        "  summary: {                     // Compteurs 24h",
        "    auth_success:    number;",
        "    auth_failure:    number;",
        "    auth_lockout:    number;",
        "    access_denied:   number;",
        "    session_expired: number;",
        "    logout:          number;",
        "    data_export:     number;",
        "    data_anonymize:  number;",
        "    waf_block:       number;",
        "  };",
        "  waf: {",
        "    total:    number;",
        "    byType:   Record<string, number>;  // { SQL_INJECTION: 3, XSS: 1, ... }",
        "    last1h:   number;",
        "    blocked:  Array<{ reason, path, ip, userAgent, timestamp }>;",
        "  };",
        "  last1h:  { auth_failure, auth_lockout, access_denied };",
        "  alerts:  { bruteForce, accessEscalation, manyLockouts, wafAttack };",
        "  recentEvents: Array<{ action, email, path, role, ip, timestamp }>;",
        "}",
    ])

    h2(doc, "6.4  Les 3 onglets du dashboard")
    data_table(doc,
        ["Onglet", "Contenu", "CDC §3.3"],
        [
            ("📊 Dashboard", "Accès sécurisé · RGPD · RBAC · Journalisation · Score /100 · Alertes animées", "Complet"),
            ("🛡️ WAF",       "Compteurs attaques · Répartition par type · 10 dernières attaques bloquées",   "OWASP A03/A05"),
            ("📋 SIEM",      "Journal 20 événements : IP · email · rôle · chemin · horodatage",              "Journalisation"),
        ]
    )

    h2(doc, "6.5  Seuils d'alerte automatiques")
    data_table(doc,
        ["Alerte", "Condition", "Niveau", "Comportement"],
        [
            ("Brute Force",         "> 20 AUTH_FAILURE / 24h",   "CRITIQUE", "Bannière rouge animée pulse"),
            ("Nombreux blocages",   "> 5 AUTH_LOCKOUT / 24h",    "CRITIQUE", "Bannière rouge animée pulse"),
            ("Escalade privilèges", "> 10 ACCESS_DENIED / 24h",  "HAUTE",    "Bannière orange"),
            ("Attaque WAF",         "> 10 WAF_BLOCK / 24h",      "CRITIQUE", "Bannière rouge animée pulse"),
        ]
    )

    h2(doc, "6.6  Protections actives affichées")
    body(doc, "Le dashboard affiche en temps réel l'état des protections activées :")
    for p in [
        "JWT HS256 forcé",
        "Brute force : 5 tentatives → blocage 5 minutes",
        "Rate limiting : 10 req / 15 min",
        "bcrypt saltRounds=10",
        "Timeout inactivité : 30 minutes",
        "Anti-énumération activé (même message pour email inconnu ou mot de passe incorrect)",
        "AES-256-CBC chiffrement données personnelles",
        "HTTPS/HSTS max-age=1 an",
        "allow update: if false — modification auditLogs impossible",
        "allow delete: if false — suppression auditLogs impossible",
    ]:
        bullet(doc, p)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 7 — SCANNER DAST
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 7 — Scanner DAST Automatisé (security_scan.js)")

    body(doc,
        "Le scanner DAST (Dynamic Application Security Testing) est un script Node.js "
        "qui exécute 12 tests de sécurité automatisés contre l'API backend en cours "
        "d'exécution. Il simule les attaques réelles et vérifie que les contrôles "
        "de sécurité fonctionnent correctement.")

    h2(doc, "7.1  Exécution")
    code_block(doc, [
        "# Démarrer le backend",
        "cd back && npm start",
        "",
        "# Lancer le scanner DAST",
        "cd back/functions",
        "node scripts/security_scan.js http://localhost:5001/gestionadminastration/us-central1/api",
        "",
        "# Sortie attendue",
        "  ✅ PASS  [T01] API accessible — HTTP 200",
        "  ✅ PASS  [T02] Headers sécurité présents",
        "  ✅ PASS  [T03] Authentification requise — HTTP 401",
        "  ✅ PASS  [T04] Rate limiting déclenché — HTTP 429",
        "  ✅ PASS  [T05] WAF bloque SQLi — HTTP 403",
        "  ✅ PASS  [T06] WAF bloque XSS — HTTP 403",
        "  ✅ PASS  [T07] WAF bloque Path Traversal — HTTP 403",
        "  ✅ PASS  [T08] WAF bloque agent suspect — HTTP 403",
        "  ✅ PASS  [T09] Payload > 10kb rejeté — HTTP 413",
        "  ✅ PASS  [T10] Privilege escalation bloquée",
        "  ✅ PASS  [T11] /monitoring réservé admins — HTTP 401",
        "  ✅ PASS  [T12] CORS refuse origines non autorisées",
        "",
        "  Résultat : 12/12 tests PASS — 0 FAIL",
    ])

    h2(doc, "7.2  12 Tests couverts")
    data_table(doc,
        ["Test", "Vérification", "Résultat attendu", "OWASP"],
        [
            ("T01", "API accessible et opérationnelle",             "HTTP 200",     "—"),
            ("T02", "Headers HTTP sécurité (HSTS, X-Frame...)",    "Headers OK",   "A05:2021"),
            ("T03", "Auth requise sur endpoints protégés",          "HTTP 401",     "A07:2021"),
            ("T04", "Rate limiting déclenché après N req",          "HTTP 429",     "A07:2021"),
            ("T05", "WAF bloque injection SQL",                     "HTTP 403",     "A03:2021"),
            ("T06", "WAF bloque XSS",                              "HTTP 403",     "A03:2021"),
            ("T07", "WAF bloque Path Traversal",                   "HTTP 403",     "A01:2021"),
            ("T08", "WAF bloque agents suspects (sqlmap...)",       "HTTP 403",     "A05:2021"),
            ("T09", "Payload > 10kb rejeté",                       "HTTP 413",     "A05:2021"),
            ("T10", "Privilege escalation bloquée",                 "Rôle forcé",   "A01:2021"),
            ("T11", "/monitoring réservé admins",                  "HTTP 401/403", "A01:2021"),
            ("T12", "CORS refuse origines non autorisées",          "Header absent","A05:2021"),
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 8 — CI/CD SÉCURITÉ
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 8 — Pipeline CI/CD & Sécurité")

    body(doc,
        "Le pipeline GitLab CI/CD intègre des vérifications de sécurité à chaque "
        "commit. Les 6 stages s'exécutent séquentiellement — un stage qui échoue "
        "bloque les suivants, empêchant un déploiement avec des erreurs.")

    h2(doc, "8.1  Stages du pipeline")
    code_block(doc, [
        "stages:",
        "  - install          # npm install backend + frontend",
        "  - lint             # ESLint backend (allow_failure: false)",
        "  - test             # npm test backend + vitest frontend",
        "  - build            # npm run build frontend",
        "  - deploy-staging   # Firebase deploy → branche 'test'",
        "  - deploy-production# Firebase deploy → branche 'main' uniquement",
    ])

    h2(doc, "8.2  Pratiques de sécurité CI/CD")
    data_table(doc,
        ["Pratique", "Implémentation", "Justification"],
        [
            ("Secrets hors du code",     "Variables GitLab CI ($FIREBASE_TOKEN)",      "Jamais dans le dépôt git"),
            ("package-lock.json",        "npm ci avec versions fixées",                "Reproductibilité + évite supply chain"),
            ("Lint bloquant",            "allow_failure: false",                       "Un lint échoué bloque le déploiement"),
            ("Tests bloquants",          "allow_failure: false backend",               "Tests unitaires crypto obligatoires"),
            ("Staging avant production", "deploy-staging sur branche test",            "Test en conditions réelles avant prod"),
            ("Production sur main seul", "only: - main",                               "Pas de déploiement sauvage"),
        ]
    )

    h2(doc, "8.3  Tests de sécurité — AES & CSV")
    body(doc,
        "Deux fichiers de tests unitaires couvrent les fonctions de sécurité critiques :")
    code_block(doc, [
        "# back/functions/src/utils/encryption.test.js",
        "# Tests : encrypt() + decrypt() avec clé valide",
        "#         decrypt() avec clé invalide → erreur attendue",
        "#         chiffrement de chaîne vide",
        "#         résultat chiffré ≠ texte original",
        "",
        "# back/functions/src/utils/csvGenerator.test.js",
        "# Tests : génération CSV sans injection de formules",
        "#         headers corrects, encodage UTF-8",
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CHAPITRE 9 — PLAN D'AMÉLIORATION
    # ══════════════════════════════════════════════════════════════════════════
    h1(doc, "Chapitre 9 — Plan d'Amélioration & Recommandations")

    h2(doc, "9.1  Priorités immédiates (< 3 mois)")
    data_table(doc,
        ["Priorité", "Amélioration", "Impact", "Effort"],
        [
            ("P1 — Critique", "MFA (TOTP / SMS) via Firebase Auth",              "Élimine 99% risque credential theft",  "Moyen"),
            ("P2 — Haute",    "Cookies httpOnly pour JWT (vs localStorage)",      "Élimine XSS → token theft",            "Faible"),
            ("P3 — Haute",    "DAST scanner dans CI/CD (stage security-scan)",    "Régression sécurité détectée auto",    "Faible"),
            ("P4 — Haute",    "npm audit dans CI/CD — fail sur CVE critique",     "Supply chain sécurisé",                "Faible"),
            ("P5 — Moyenne",  "Google Secret Manager (rotation secrets)",         "Compromission clé → rotation auto",    "Moyen"),
        ]
    )

    h2(doc, "9.2  Améliorations moyen terme (3-6 mois)")
    for item in [
        "Pentest externe par prestataire certifié OSCP — audit indépendant obligatoire sous RGPD Art.32",
        "DPIA formelle (Art.35) — document réglementaire pour traitement de données mineurs",
        "Redis pour blacklist JWT — révocation immédiate sans vérification Firestore",
        "Argon2id en remplacement de bcrypt — résistance accrue aux GPU/ASIC",
        "Notification automatique CNIL 72h via active response Wazuh — conformité Art.33",
        "SAST complet avec SonarQube intégré dans le CI/CD",
        "WAF contre double-encodage — boucle decodeURIComponent jusqu'à stabilisation",
    ]:
        bullet(doc, item)

    h2(doc, "9.3  Synthèse — Score de sécurité global")
    data_table(doc,
        ["Dimension", "Score", "Détail"],
        [
            ("OWASP Top 10",      "8/10",   "A04 et A10 partiellement couverts"),
            ("RGPD",              "5/6",    "Art.35 DPIA non formalisé"),
            ("Tests DAST",        "12/12",  "100% tests passés"),
            ("Démonstrations",    "7/7",    "Toutes réussies en soutenance"),
            ("Monitoring",        "100%",   "Applicatif + Infrastructure Wazuh"),
            ("Traçabilité",       "100%",   "9 types d'events auditLogs immuables"),
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # ANNEXE — VARIABLES D'ENVIRONNEMENT
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    h1(doc, "Annexe — Variables d'Environnement & Configuration")

    h2(doc, "Variables requises")
    code_block(doc, [
        "# back/functions/.env (ne jamais commiter)",
        "JWT_SECRET=<secret_256_bits_minimum>",
        "REFRESH_TOKEN_SECRET=<secret_256_bits_different>",
        "ENCRYPTION_KEY=<clé_32_bytes_AES256>",
        "SMTP_HOST=<smtp_server>",
        "SMTP_PORT=587",
        "SMTP_USER=<email_envoi>",
        "SMTP_PASS=<mot_de_passe_smtp>",
        "",
        "# Firebase Functions Config (production)",
        "firebase functions:config:set jwt.secret='...' encryption.key='...'",
    ])

    h2(doc, "Variables GitLab CI (secrets)")
    data_table(doc,
        ["Variable", "Usage", "Où la définir"],
        [
            ("$FIREBASE_PROJECT_ID",         "ID projet Firebase production", "GitLab → Settings → CI/CD → Variables"),
            ("$FIREBASE_PROJECT_STAGING_ID", "ID projet Firebase staging",    "GitLab → Settings → CI/CD → Variables"),
            ("$FIREBASE_TOKEN",              "Token déploiement Firebase",    "GitLab → Settings → CI/CD → Variables (masked)"),
        ]
    )

    # ── Pied de page ──────────────────────────────────────────────────────────
    doc.add_page_break()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_sp(p_footer, 400, 20)
    run(p_footer, "DOCUMENTATION TECHNIQUE — SÉCURITÉ APPLICATIVE & INFRASTRUCTURE",
        bold=True, size=11, color=C_NAVY)

    p_f2 = doc.add_paragraph()
    p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_sp(p_f2, 0, 0)
    run(p_f2,
        "Anass Akker  ·  PFE M2 Cybersécurité  ·  YNOV Campus  ·  Avril 2026  ·  Version 3.0",
        size=9, color=C_DIM, italic=True)

    doc.save(OUTPUT)
    print(f"✅  Document généré : {OUTPUT}")
    print(f"   Chapitres : 9  +  1 annexe")
    print(f"   Pages estimées : ~28-32")

if __name__ == "__main__":
    build()

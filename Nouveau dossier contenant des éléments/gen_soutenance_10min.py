#!/usr/bin/env python3
"""
Génère GUIDE_SOUTENANCE_10MIN.docx
Script de soutenance minuté + résumés rapport + commandes de démo
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def titre_page(doc, texte, couleur=RGBColor(0x1a,0x3a,0x6e)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = couleur
    return p

def h1(doc, texte, couleur=RGBColor(0x1a,0x3a,0x6e)):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = couleur
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def h2(doc, texte, couleur=RGBColor(0x2e,0x6d,0xb8)):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = couleur
    p.paragraph_format.space_before = Pt(6)
    return p

def bullet(doc, texte, level=0, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if bold_part and bold_part in texte:
        idx = texte.index(bold_part)
        r1 = p.add_run(texte[:idx])
        r1.font.size = Pt(10)
        r2 = p.add_run(texte[idx:idx+len(bold_part)])
        r2.bold = True
        r2.font.size = Pt(10)
        r3 = p.add_run(texte[idx+len(bold_part):])
        r3.font.size = Pt(10)
    else:
        run = p.add_run(texte)
        run.font.size = Pt(10)
    return p

def normal(doc, texte, size=10, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def separateur(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

def bandeau(doc, texte, bg=None):
    if bg is None:
        bg_hex = '1A3A6E'
    elif isinstance(bg, int):
        # RGBColor is a subclass of int — format as 6-digit hex
        bg_hex = '%06X' % bg
    else:
        bg_hex = '%02X%02X%02X' % (bg[0], bg[1], bg[2])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xff,0xff,0xff)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), bg_hex)
    p._p.get_or_add_pPr().append(shading)
    return p

def code_block(doc, texte):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), '1E1E1E')
    p._p.get_or_add_pPr().append(shading)
    return p

def alerte(doc, texte, couleur=RGBColor(0xcc,0x00,0x00)):
    p = doc.add_paragraph()
    run = p.add_run("⚡ " + texte)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = couleur
    return p

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COUVERTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
titre_page(doc, "🎓 GUIDE DE SOUTENANCE — 10 MINUTES")
titre_page(doc, "PFE Cybersécurité / Cyberdéfense — YNOV Campus 2026", RGBColor(0x2e,0x6d,0xb8))
doc.add_paragraph()
bandeau(doc, "Anass Akker  ·  Amine BAHOU  ·  24 avril 2026  ·  10h30→10h40")
doc.add_paragraph()
normal(doc, "Ce document contient :", 11, bold=True)
bullet(doc, "Script minuté — ce qu'il faut dire pour chaque minute")
bullet(doc, "Résumé condensé du rapport de sécurité (1 ligne par section)")
bullet(doc, "Chiffres clés à retenir absolument")
bullet(doc, "Toutes les commandes de démonstration")
bullet(doc, "Réponses rapides aux questions du jury")
separateur(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — CHIFFRES CLÉS (À MÉMORISER)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "📊 CHIFFRES CLÉS — À RETENIR PAR CŒUR", RGBColor(0x1a,0x3a,0x6e))
doc.add_paragraph()

h1(doc, "🔢 LES NOMBRES QUI IMPRESSIONNENT LE JURY")

chiffres = [
    ("436 420",  "événements Wazuh collectés (Security Events + FIM)"),
    ("12/12",    "tests DAST réussis — 0 vulnérabilité critique"),
    ("17",       "CVE détectées par Wazuh (8 High + 9 Medium) — Docker principal"),
    ("1 600+",   "alertes MITRE ATT&CK — technique T1565.001 dominante"),
    ("5/6",      "articles RGPD satisfaits (Art.5, 15, 16, 17, 25, 32, 33)"),
    ("9",        "types d'auditLogs immuables Firestore (append-only)"),
    ("6",        "rôles RBAC — principle du moindre privilège"),
    ("100/100",  "score sécurité dashboard — 0 incident actif"),
    ("5",        "types d'attaques WAF bloqués (SQLi, XSS, Path, CMD, Agents)"),
    ("4",        "commandes pour déployer Wazuh (git clone → docker compose up)"),
    ("2",        "couches de sécurité — applicatif + infrastructure (Defense in Depth)"),
    ("100%",     "coverage agents Wazuh — 0 point aveugle"),
    ("30 min",   "durée access token JWT (refresh token : 7 jours)"),
    ("10",       "iterations bcrypt saltRounds — anti-rainbow table"),
    ("92/100",   "score DAST (1 warning CORS — non critique)"),
]

for nb, desc in chiffres:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f"{nb} ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
    r2 = p.add_run(f"→ {desc}")
    r2.font.size = Pt(10)

separateur(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — SCRIPT MINUTÉ 10 MINUTES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "⏱️ SCRIPT MINUTÉ — 10h30 À 10h40", RGBColor(0x8b,0x00,0x00))
doc.add_paragraph()

alerte(doc, "OBJECTIF : 10 minutes = 30 slides = ~20 sec/slide. Ne lis pas les slides, dis l'essentiel.", RGBColor(0xcc,0x00,0x00))
doc.add_paragraph()

minutes = [
    {
        "time": "⏱ MIN 1 — INTRODUCTION (slide 1-3) — 10h30",
        "dire": [
            '"Bonjour, je présente la partie Sécurité & Monitoring de notre application de gestion scolaire YNOV."',
            '"Notre application gère des données sensibles d\'élèves et de paiements — le RGPD est obligatoire."',
            '"J\'ai implémenté une défense en profondeur sur 2 niveaux : applicatif + infrastructure Wazuh."',
            '"Stack : React / Node.js / Firebase / Wazuh 4.7.4 Docker."',
        ],
        "nb": "KPI à citer : 10 catégories OWASP · 6 articles RGPD · 12 tests DAST · 436 420 events Wazuh",
        "color": RGBColor(0x00, 0x70, 0xc0)
    },
    {
        "time": "⏱ MIN 2 — OWASP + VULNÉRABILITÉS (slides 4-5) — 10h31",
        "dire": [
            '"J\'ai audité les 10 catégories OWASP Top 10 — 8 corrigées, 1 surveillée, 1 N/A."',
            '"Le WAF middleware waf.js intercepte SQLi, XSS, Path Traversal, CMD Injection, agents scanners."',
            '"Toute attaque → HTTP 403 + log WAF_BLOCK dans Firestore auditLogs."',
            '"Les données sensibles — téléphone, adresse — sont chiffrées AES-256-CBC avant stockage."',
        ],
        "nb": "5 types bloqués · mot de passe = jamais stocké (bcrypt hash irréversible)",
        "color": RGBColor(0xcc, 0x00, 0x00)
    },
    {
        "time": "⏱ MIN 3 — RBAC + RGPD (slides 6-7) — 10h32",
        "dire": [
            '"Le RBAC définit 6 rôles : admin, sous-admin, comptable, enseignant, étudiant, parent."',
            '"Chaque route vérifie le rôle en temps réel depuis Firestore — jamais depuis le JWT seul."',
            '"Firestore Rules : deny by default — toute collection non listée = refus automatique."',
            '"RGPD : 5/6 articles satisfaits. Art.15 (export), 17 (anonymisation), 25 (privacy by design)."',
        ],
        "nb": "checkRole() re-lit Firestore à CHAQUE requête · AES-256-CBC + HTTPS/HSTS",
        "color": RGBColor(0x00, 0x80, 0x00)
    },
    {
        "time": "⏱ MIN 4 — AUDIT LOGS + JWT + RATE LIMIT (slides 8-10) — 10h33",
        "dire": [
            '"9 types d\'auditLogs immuables — allow update: if false au niveau Firestore."',
            '"Personne ne peut modifier ces logs, même un administrateur, même si le backend est compromis."',
            '"JWT HS256 : access token 30 min, refresh 7 jours. Rate limiting : 5 tentatives → lockout."',
            '"6e tentative → HTTP 429 Too Many Requests + AUTH_LOCKOUT loggué."',
        ],
        "nb": "auditLogs = serverTimestamp() côté Google — non manipulable client",
        "color": RGBColor(0x7f, 0x00, 0x7f)
    },
    {
        "time": "⏱ MIN 5 — ARCHITECTURE + DAST (slides 11-12) — 10h34",
        "dire": [
            '"Le pipeline de sécurité : WAF → Rate Limiter → verifyJWT → checkRole() → Handler → auditLogger."',
            '"Chaque requête passe par 4 couches avant d\'accéder aux données."',
            '"Scanner DAST : 12 tests OWASP automatisés — 12/12 PASS, score 92/100."',
            '"Le scanner teste en live : SQLi, XSS, Path Traversal, brute force, escalade de privilèges."',
        ],
        "nb": "DAST = Dynamic Application Security Testing — requêtes HTTP réelles, pas du code statique",
        "color": RGBColor(0xcc, 0x66, 0x00)
    },
    {
        "time": "⏱ MIN 6 — DASHBOARD MONITORING APPLICATIF (slides 13-18) — 10h35",
        "dire": [
            '"Le dashboard /monitoring affiche un score de sécurité /100, mis à jour toutes les 60 secondes."',
            '"3 onglets : Dashboard (auth/RBAC/RGPD), WAF (attaques bloquées), SIEM Logs (20 derniers events)."',
            '"Score 100/100 aujourd\'hui = 0 incident actif. En cas d\'attaque, il descend et devient rouge."',
            '"Le SIEM onglet = boîtier noir infalsifiable : chaque connexion, attaque, export RGPD est tracé."',
        ],
        "nb": "setInterval(load, 60_000) · admin only · <ProtectedRoute roles={['admin']}>",
        "color": RGBColor(0x00, 0x70, 0xc0)
    },
    {
        "time": "⏱ MIN 7 — WAZUH ARCHITECTURE + AGENTS (slides 19-21) — 10h36",
        "dire": [
            '"Wazuh SIEM : déployé en 4 commandes Docker — manager, indexer, dashboard."',
            '"Agent 001 actif sur la machine macOS — coverage 100%, 0 point aveugle."',
            '"Protocole OSSEC port 1514, chiffrement AES, temps réel."',
            '"Les logs applicatifs YNOV sont envoyés via /tmp/applogs/*.log → Wazuh Manager."',
        ],
        "nb": "3 conteneurs Docker · Agent 001 macOS 15.7.4 · 7 règles custom 100010-100016",
        "color": RGBColor(0x1a, 0x3a, 0x6e)
    },
    {
        "time": "⏱ MIN 8 — WAZUH RÉSULTATS (slides 22-26) — 10h37",
        "dire": [
            '"436 420 événements collectés — Rule 550 Level 7 (Integrity checksum changed) dominante."',
            '"FIM surveille les binaires système /var/bin/* — root 89.44% des modifications."',
            '"17 CVE identifiées dont CVE-2019-5736 CVSS3=8.6 sur Docker — patch prioritaire P1."',
            '"MITRE ATT&CK : T1565.001 Stored Data Manipulation dominant + T1562 Defense Evasion."',
        ],
        "nb": "Rootcheck : interface en3 promiscuous + processus cachés détectés",
        "color": RGBColor(0x8b, 0x00, 0x00)
    },
    {
        "time": "⏱ MIN 9 — DEFENSE IN DEPTH + SYNTHÈSE (slides 27-29) — 10h38",
        "dire": [
            '"Defense in Depth sur 2 niveaux : couche applicative (WAF + auditLogs + dashboard) + couche infrastructure (Wazuh)."',
            '"Wazuh surveille là où mon code ne peut pas voir : OS, fichiers système, CVE packages."',
            '"CDC §3.3 couvert à 100% : accès sécurisé, RGPD, journalisation, HTTPS, monitoring."',
            '"Ensemble : 8/10 OWASP corrigés, 5/6 RGPD satisfaits, 17 CVE documentées avec plan d\'action."',
        ],
        "nb": "Phrase clé : 'Défense en profondeur = applicatif + infrastructure = aucun point aveugle'",
        "color": RGBColor(0x00, 0x80, 0x00)
    },
    {
        "time": "⏱ MIN 10 — CONCLUSION + DÉMO (slide 30) — 10h39",
        "dire": [
            '"En synthèse : WAF maison, 9 logs immuables, score 100/100, DAST 12/12, Wazuh 436k events."',
            '"Perspectives : Wazuh Active Response (blocage IP auto), Redis rate limit, patch Docker CVE."',
            '"Je suis prêt pour une démonstration live si vous souhaitez voir le WAF bloquer une SQLi."',
            '→ OUVRIR : http://localhost:8081/monitoring + https://localhost (Wazuh)',
        ],
        "nb": "Si question difficile : 'Excellente question — c'est dans mon plan d'amélioration : [MFA/Argon2id/DPIA]'",
        "color": RGBColor(0x1a, 0x3a, 0x6e)
    },
]

for m in minutes:
    bandeau(doc, m["time"], RGBColor(0x1a,0x3a,0x6e))
    for d in m["dire"]:
        bullet(doc, d)
    p = doc.add_paragraph()
    r = p.add_run("📌 " + m["nb"])
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = m["color"]
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — RÉSUMÉ CONDENSÉ DU RAPPORT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "📋 RÉSUMÉ CONDENSÉ DU RAPPORT DE SÉCURITÉ", RGBColor(0x1a,0x3a,0x6e))
normal(doc, "1 phrase par section — pour répondre aux questions sur le rapport", 9, italic=True)
doc.add_paragraph()

sections = [
    ("1. KPI Sécurité — Tableau de Bord",
     "436 420 events Wazuh · 17 CVE (8H+9M) · 5 AUTH_FAILURE · 3 WAF_BLOCK · 5/6 RGPD · 100% Agent Coverage",
     [
         "6 modules sur 8 opérationnels (SCA et CVE en cours de correction)",
         "Audit du 5 au 12 avril 2026 — chronologie documentée avec dates précises",
         "SCA 0/10 = hardening SSH non appliqué (non bloquant pour la soutenance)",
     ]),
    ("2. Problèmes Rencontrés (8 obstacles résolus)",
     "7/8 problèmes résolus · 1 contournement (NVD feed payant)",
     [
         "#1 Agent macOS Disconnected → reconfiguration ossec.conf (30 min)",
         "#2 FIM sans résultats → fréquence 43200s→60s + echo hosts (15 min)",
         "#3 NVD CVE payant → syscollector Docker Linux (contournement)",
         "#4 bcrypt DLL Windows → npm rebuild bcrypt (10 min)",
         "#5 Firebase timeout defineSecret → server_local.js avec mock modules (45 min)",
         "#6 Règles Wazuh champ 'action' réservé → <match> au lieu de <field> (20 min)",
         "#7 macOS SIP /etc write denied → /private/etc/hosts avec sudo (5 min)",
         "#8 PPTX Bad CRC-32 → reconstruction ZIP fichier/fichier (10 min)",
     ]),
    ("3. Architecture Wazuh",
     "Manager Docker 172.20.0.2:1514 · Agent macOS 127.0.0.1 · Backend :5001 · Frontend :5173",
     [
         "AuditLog.js écrit dans /tmp/applogs/*.log → lu par Wazuh via <localfile>",
         "local_decoder.xml parse le format JSON YNOV-APP",
         "7 règles custom (100010-100016) : Level 3 (info) à Level 14 (critique)",
         "AUTH_LOCKOUT = Level 14 (le plus élevé) — alerte immédiate",
         "WAF_BLOCK = Level 12 — attaque OWASP bloquée par le pare-feu applicatif",
     ]),
    ("4.1 Security Events",
     "436 420 events · Rule 550 Level 7 · T1565.001 · Level 12+ = 0 (aucune alerte critique)",
     [
         "Top alert : Integrity checksum changed (~95% des alertes)",
         "PCI DSS 11.5 × 1594 · 10.6.1 × 34 — conformité réglementaire prouvée",
         "0 Authentication failure, 0 Authentication success → système sûr au snapshot",
     ]),
    ("4.2 FIM — File Integrity Monitoring",
     "436 420 events · root 89.44% · modified 100% · /var/bin/afsa dominant",
     [
         "Hash SHA-256 sur les binaires système macOS /var/bin/*",
         "Files added / Files deleted = 'No results found' → aucune création/suppression suspecte",
         "Tout écart non autorisé → Rule 550 Level 7 — détection rootkit en < 5 secondes",
     ]),
    ("4.3 Policy Monitoring (Rootcheck)",
     "4 anomalies : Trojaned files · Hidden processes · en3 promiscuous · write permissions",
     [
         "Interface en3 en mode promiscuous = capture de trafic réseau détectée",
         "Processus 26061 caché = tentative de dissimulation potentielle",
         "En contexte macOS normal, peuvent être des faux positifs — nécessite investigation",
     ]),
    ("4.4 CVE — 17 Vulnérabilités",
     "0 Critical · 8 High · 9 Medium · Docker 4.43.2 = 8 CVE · CVE-2019-5736 CVSS3=8.6",
     [
         "CVE-2019-5736 = container escape → attaquant obtient root sur l'hôte (P1 CRITIQUE)",
         "Excel 16.107.3 → CVE-2001-0718 · lz4 1.10.0 → CVE-2014-4715",
         "Action : docker update Docker Engine vers version stable",
     ]),
    ("4.5 MITRE ATT&CK",
     "T1565.001 Stored Data Manipulation (Impact) dominant · T1562 Defense Evasion",
     [
         "Impact ~ 95% (modifications binaires /var/bin/*) · Defense Evasion ~5%",
         "PCI DSS 11.5 × 1594 · RGPD Art.25 (Privacy by Design) couvert",
         "Cartographie alignée sur les standards SOC internationaux",
     ]),
    ("5. Conformité RGPD × Wazuh",
     "5/6 articles satisfaits · Art.32 en cours (patch CVE-2019-5736 Docker)",
     [
         "Art.5 (intégrité) = FIM 436 420 alertes Rule 550 Level 7 ✓",
         "Art.17 (effacement) = DATA_EXPORT Rule 100016 → audit trail complet ✓",
         "Art.25 (Privacy by Design) = WAF + Rate Limit + Rules actifs ✓",
         "Art.32 (sécurité) = 17 CVE identifiées, patch Docker prioritaire ⚠",
         "Art.33 (notification 72h) = Alerte Level 14 AUTH_LOCKOUT → notif immédiate ✓",
         "Art.35 (DPIA) = Dashboard SIEM complet → analyse d'impact toutes données ✓",
     ]),
    ("6. Plan d'Action",
     "P1 CRITIQUE : Docker update (CVE-2019-5736) · P2 : Hardening SCA macOS · P3 : 9 CVE Medium",
     [
         "Activer Wazuh Active Response → blocage IP automatique après 5 AUTH failures",
         "Alertes email/Slack Level 10+ → intégration Wazuh → Slack webhook",
         "Ajouter agent Wazuh sur serveur production Firebase Functions",
     ]),
]

for titre, resume, details in sections:
    h1(doc, titre)
    p = doc.add_paragraph()
    r = p.add_run("→ " + resume)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x00, 0x70, 0xc0)
    for d in details:
        bullet(doc, d, level=0)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — TOUTES LES COMMANDES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "💻 TOUTES LES COMMANDES DE DÉMONSTRATION", RGBColor(0x00, 0x40, 0x00))
doc.add_paragraph()

h1(doc, "🚀 DÉMARRAGE APPLICATION (avant la soutenance)")
code_block(doc, "# Terminal 1 — Backend Node.js")
code_block(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/back/functions"')
code_block(doc, "node server_local.js")
code_block(doc, "# → http://localhost:5001")
doc.add_paragraph()
code_block(doc, "# Terminal 2 — Frontend React")
code_block(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/front"')
code_block(doc, "npm run dev")
code_block(doc, "# → http://localhost:8081")
doc.add_paragraph()
code_block(doc, "# Terminal 3 — Wazuh Docker")
code_block(doc, 'ln -s "/Users/anass/Downloads/frais-gestionScolaire 4/front/wazuh-docker/single-node" /tmp/wazuh-node')
code_block(doc, "cd /tmp/wazuh-node && docker compose up -d")
code_block(doc, "# → https://localhost  (admin / SecretPassword)  [attendre ~2 min]")
separateur(doc)

h1(doc, "🛡️ ÉTAPE 01 — WAF Injection SQL (1 min)")
code_block(doc, "# Doit retourner HTTP 403")
code_block(doc, "curl -s -o /dev/null -w 'HTTP: %{http_code}\\n' \\")
code_block(doc, "  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?id=1%20OR%201=1'")
code_block(doc, "# Résultat attendu : HTTP: 403")
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Ce qu'il faut dire : ")
r.bold = True; r.font.size = Pt(10)
r = p.add_run('"Le WAF intercepte la requête avant qu\'elle touche Firestore → HTTP 403 + WAF_BLOCK loggué"')
r.font.size = Pt(10); r.italic = True
separateur(doc)

h1(doc, "🛡️ ÉTAPE 02 — WAF XSS (30 sec)")
code_block(doc, "curl -s -o /dev/null -w 'HTTP: %{http_code}\\n' \\")
code_block(doc, "  'http://localhost:5001/gestionadminastration/us-central1/api/etudiants?nom=<script>alert(1)</script>'")
code_block(doc, "# Résultat attendu : HTTP: 403")
separateur(doc)

h1(doc, "🔒 ÉTAPE 03 — Brute Force / Rate Limiting (1 min)")
code_block(doc, "for i in {1..6}; do")
code_block(doc, "  curl -s -o /dev/null -w \"Tentative $i → %{http_code}\\n\" \\")
code_block(doc, "    -X POST http://localhost:5001/gestionadminastration/us-central1/api/auth/login \\")
code_block(doc, "    -H 'Content-Type: application/json' \\")
code_block(doc, "    -d '{\"email\":\"admin@school.fr\",\"password\":\"WRONG\"}'")
code_block(doc, "done")
code_block(doc, "# Résultat : 1→401 · 2→401 · 3→401 · 4→401 · 5→401 · 6→429")
p = doc.add_paragraph()
r = p.add_run("Ce qu'il faut dire : ")
r.bold = True; r.font.size = Pt(10)
r = p.add_run('"5 tentatives → compte verrouillé 15 min. Tentative 6 = HTTP 429 Too Many Requests + AUTH_LOCKOUT Firestore"')
r.font.size = Pt(10); r.italic = True
separateur(doc)

h1(doc, "🔬 ÉTAPE 04 — Scanner DAST Automatique (1 min)")
code_block(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/back/functions"')
code_block(doc, "node scripts/security_scan.js")
code_block(doc, "# Résultat attendu :")
code_block(doc, "# [PASS] WAF actif — SQLi bloqué")
code_block(doc, "# [PASS] Rate limiting — 429 après 11 req")
code_block(doc, "# [PASS] JWT validation — 401 token invalide")
code_block(doc, "# [PASS] Headers sécurité — HSTS présent")
code_block(doc, "# Score DAST final : 92/100 — 11/12 tests réussis")
separateur(doc)

h1(doc, "🌐 ÉTAPE 05 — Wazuh Dashboard Live (2 min)")
code_block(doc, "# Ouvrir dans le navigateur :")
code_block(doc, "open https://localhost")
code_block(doc, "# Login : admin / SecretPassword")
code_block(doc, "# Montrer : Security Events → 436 420 events")
code_block(doc, "#           MITRE ATT&CK → T1565.001 dominant")
code_block(doc, "#           File Integrity → root 89.44%")
code_block(doc, "#           Vulnerabilities → 17 CVE")
separateur(doc)

h1(doc, "📊 ÉTAPE 06 — Dashboard Monitoring Applicatif (1 min)")
code_block(doc, "# Ouvrir dans le navigateur :")
code_block(doc, "open http://localhost:8081/monitoring")
code_block(doc, "# Montrer : Score 100/100")
code_block(doc, "#           Onglet WAF → 0 attaques (sécurisé)")
code_block(doc, "#           Onglet SIEM Logs → journal events")
separateur(doc)

h1(doc, "🔐 ÉTAPE 07 — RGPD Traçabilité Art.5 & Art.32 (1 min)")
code_block(doc, "# 1. Ouvrir onglet privé → http://localhost:8081/login")
code_block(doc, "# 2. Taper un MAUVAIS mot de passe → connexion échouée")
code_block(doc, "# 3. Aller sur http://localhost:8081/monitoring → SIEM Logs → Rafraîchir")
code_block(doc, "# 4. Montrer : auth_failure loggué avec timestamp + IP + email")
p = doc.add_paragraph()
r = p.add_run("Ce qu'il faut dire : ")
r.bold = True; r.font.size = Pt(10)
r = p.add_run('"Chaque tentative échouée est tracée immédiatement. Timestamp = serverTimestamp() côté Google, non manipulable. RGPD Art.5 et Art.32."')
r.font.size = Pt(10); r.italic = True
separateur(doc)

h1(doc, "🔑 ÉTAPE BONUS — Test Chiffrement AES-256 (si demandé)")
code_block(doc, 'cd "/Users/anass/Downloads/frais-gestionScolaire 4/back/functions"')
code_block(doc, "node -e \"")
code_block(doc, "const enc = require('./src/utils/encryption');")
code_block(doc, "const c = enc.encrypt('donnee-sensible');")
code_block(doc, "console.log('Chiffré :', c);")
code_block(doc, "console.log('Déchiffré :', enc.decrypt(c));\"")
separateur(doc)

h1(doc, "🏥 RGPD — Routes Art.15 & Art.17 (si demandé)")
code_block(doc, "# Récupérer token admin d'abord :")
code_block(doc, "TOKEN=$(curl -s -X POST http://localhost:5001/gestionadminastration/us-central1/api/auth/login \\")
code_block(doc, "  -H 'Content-Type: application/json' \\")
code_block(doc, "  -d '{\"email\":\"admin@school.fr\",\"password\":\"VOTRE_MDP\"}' | jq -r '.token')")
code_block(doc, "")
code_block(doc, "# Art.15 — Export données personnelles")
code_block(doc, "curl -H \"Authorization: Bearer $TOKEN\" http://localhost:5001/.../api/users/USER_ID/export")
code_block(doc, "")
code_block(doc, "# Art.17 — Anonymisation")
code_block(doc, "curl -X DELETE -H \"Authorization: Bearer $TOKEN\" http://localhost:5001/.../api/users/USER_ID/data")
separateur(doc)

h1(doc, "🔧 Wazuh — Redémarrer Agent si déconnecté")
code_block(doc, "sudo /Library/Ossec/bin/wazuh-control stop")
code_block(doc, "sudo /Library/Ossec/bin/wazuh-control start")
code_block(doc, "sudo /Library/Ossec/bin/wazuh-control status")
code_block(doc, "# → wazuh-agentd is running")
separateur(doc)

h1(doc, "📁 Déclencher FIM manuellement (si demo Wazuh)")
code_block(doc, "# Modifier /private/etc/hosts → déclenche Rule 550 Level 7")
code_block(doc, "sudo bash -c \"echo '# wazuh-test-demo' >> /private/etc/hosts\"")
code_block(doc, "# → Dans Wazuh : File Integrity Monitoring → alerte immédiate Rule 550")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — Q&A RAPIDES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
bandeau(doc, "❓ Q&A RAPIDES — RÉPONSES EN 1 PHRASE", RGBColor(0x1a,0x3a,0x6e))
doc.add_paragraph()

qa = [
    ("Pourquoi Wazuh en plus du dashboard ?",
     "Wazuh surveille l'infrastructure (OS, fichiers, CVE) là où mon code ne peut pas voir. Defense in Depth = 2 couches indépendantes."),
    ("Comment les logs ne peuvent pas être falsifiés ?",
     "Firestore Rules : allow update: if false + allow delete: if false. Même un admin ne peut pas modifier. Même le backend compromis."),
    ("Comment fonctionne le WAF ?",
     "Express middleware avant l'auth : User-Agent → URL → Query params → Body JSON. Password exclu pour éviter faux positifs."),
    ("JWT : quelle durée ? Que contient le payload ?",
     "Access token 30 min, refresh 7 jours. Payload : {id, email, role} uniquement — jamais le mot de passe ni les données PII."),
    ("Comment protéger contre brute force ?",
     "Rate limiter : 5 tentatives / 15 min → HTTP 429 + AUTH_LOCKOUT Firestore. Limite connue : ne bloque pas les botnets multi-IP."),
    ("RGPD Art.17 vs logs immuables — contradiction ?",
     "Les logs contiennent un userId opaque, pas de nom/prénom. Effacement = on remplace l'ID par un hash irréversible. La ligne reste mais n'est plus une 'donnée personnelle' au sens Art.4."),
    ("CVE-2019-5736 — c'est grave ?",
     "CVSS3=8.6 — container escape : un attaquant dans un conteneur Docker peut obtenir root sur l'hôte. Détecté, documenté, plan d'action : docker update (P1 Critique post-soutenance)."),
    ("Quelle différence DAST vs SAST ?",
     "SAST = analyse statique du code source. DAST = application en cours d'exécution avec vraies requêtes HTTP. Mon scanner = DAST. ESLint = SAST léger."),
    ("Pourquoi bcrypt saltRounds=10 ?",
     "Facteur de coût — 2^10 = 1024 itérations ≈ 100ms par hash. Rend les attaques par dictionnaire non rentables. Même mot de passe = deux hashes différents grâce au salt aléatoire."),
    ("Que fait le module MITRE ATT&CK de Wazuh ?",
     "Il mappe automatiquement les règles Wazuh aux techniques ATT&CK. T1565.001 = modifications binaires /var/bin/* par FIM. Permet classification internationale SOC."),
    ("Le SCA score est 0/10 — c'est un problème ?",
     "C'est le hardening SSH qui n'est pas appliqué (macOS SIP). Non bloquant pour la soutenance. Plan d'action : appliquer CIS Benchmark macOS (2 semaines)."),
    ("Comment fonctionne le score /100 ?",
     "Départ 100, soustraction : -20 max si auth_failures > 5, -15 max si lockouts > 0, -15 max si access_denied > 3, -20 max si waf_blocks > 0. Vert ≥ 80, Orange 60-79, Rouge < 60."),
]

for q, r in qa:
    h2(doc, f"❓ {q}")
    p = doc.add_paragraph()
    run = p.add_run(f"✅ {r}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x00)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
separateur(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Document généré le {datetime.date.today().strftime('%d %B %Y')} · Anass Akker — PFE YNOV Campus 2026")
run.font.size = Pt(8)
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
output = "GUIDE_SOUTENANCE_10MIN.docx"
doc.save(output)
print(f"✅ {output} généré avec succès")

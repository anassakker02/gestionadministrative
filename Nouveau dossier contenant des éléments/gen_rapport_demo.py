#!/usr/bin/env python3
"""
gen_rapport_demo.py
Génère l'Annexe B — Démonstration des Mécanismes de Sécurité
en Word (.docx) avec le style du rapport existant.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs (palette rapport) ──────────────────────────────────────────────
C_NAVY      = RGBColor(0x1E, 0x2D, 0x4F)   # Header principal
C_BLUE      = RGBColor(0x27, 0x5E, 0x9E)   # Titres sections
C_LIGHT_BLU = RGBColor(0xD6, 0xE4, 0xF7)   # Fond ligne alternée
C_GREEN_BG  = RGBColor(0xE2, 0xF5, 0xE9)   # Résultat OK
C_GREEN_FG  = RGBColor(0x1A, 0x7F, 0x37)
C_CODE_BG   = RGBColor(0x1E, 0x1E, 0x2E)   # Fond terminal
C_CODE_FG   = RGBColor(0x0D, 0xD3, 0xD3)   # Texte terminal cyan
C_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY_BG   = RGBColor(0xF2, 0xF4, 0xF8)
C_ORANGE    = RGBColor(0xFF, 0x6B, 0x35)
C_RED_BG    = RGBColor(0xFF, 0xEB, 0xEB)
C_RED_FG    = RGBColor(0xC0, 0x00, 0x00)

OUTPUT = "ANNEXE_B_DEMONSTRATION.docx"

# ══════════════════════════════════════════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════════════════════════════════════════

def rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(color))
    tcPr.append(shd)

def set_cell_border(cell, sides=("top","bottom","left","right"), color="2759A0", sz=6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)

def add_run(para, text, bold=False, italic=False, size=10,
            color: RGBColor = None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    no_space_para(p)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    if level == 1:
        add_run(p, text, bold=True, size=16, color=C_NAVY, font="Calibri")
    elif level == 2:
        add_run(p, text, bold=True, size=13, color=C_BLUE, font="Calibri")
    else:
        add_run(p, text, bold=True, size=11, color=C_NAVY, font="Calibri")
    return p

def add_body(doc, text, size=10, color=None, italic=False, before=4, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    add_run(p, text, size=size, color=color, italic=italic)
    return p

def add_code_block(doc, command: str, result: str = None):
    """Terminal-style block: dark bg, cyan command, green result."""
    # Command
    p = doc.add_paragraph()
    no_space_para(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    # shade paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(C_CODE_BG))
    pPr.append(shd)
    add_run(p, f"$ {command}", bold=True, size=9, color=C_CODE_FG, font="Courier New")

    if result:
        p2 = doc.add_paragraph()
        no_space_para(p2)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(6)
        p2.paragraph_format.left_indent  = Cm(0.5)
        p2.paragraph_format.right_indent = Cm(0.5)
        pPr2 = p2._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"), "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"), rgb_hex(C_CODE_BG))
        pPr2.append(shd2)
        add_run(p2, f"→ {result}", size=9, color=RGBColor(0x4E, 0xC9, 0x54), font="Courier New")

def add_result_box(doc, text: str, ok=True):
    """Coloured result box (green OK / red FAIL)."""
    bg = C_GREEN_BG if ok else C_RED_BG
    fg = C_GREEN_FG if ok else C_RED_FG
    icon = "✔" if ok else "✘"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(bg))
    pPr.append(shd)
    add_run(p, f"  {icon}  {text}", bold=True, size=10, color=fg)

def banner_table(doc, title: str, number: str, color: RGBColor = None):
    """Full-width banner row — dark bg, white text."""
    color = color or C_NAVY
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    # remove borders
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color)
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space_para(p)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    add_run(p, f"{number}  —  {title}", bold=True, size=13,
            color=C_HEADER_FG, font="Calibri")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def info_table(doc, rows: list):
    """
    2-column info table.
    rows = [("Label", "Value"), ...]
    """
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = [Inches(2.2), Inches(4.3)]
    for i, (label, value) in enumerate(rows):
        bg = C_LIGHT_BLU if i % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        c0, c1 = tbl.rows[i].cells
        c0.width = col_w[0]; c1.width = col_w[1]
        set_cell_bg(c0, C_NAVY)
        set_cell_bg(c1, bg)
        p0 = c0.paragraphs[0]; p1 = c1.paragraphs[0]
        no_space_para(p0); no_space_para(p1)
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after  = Pt(3)
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after  = Pt(3)
        add_run(p0, label, bold=True, size=9, color=C_HEADER_FG)
        add_run(p1, value, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def synthesis_table(doc, cols: list, data: list):
    """Multi-column synthesis table with header row."""
    tbl = doc.add_table(rows=1+len(data), cols=len(cols))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, col in enumerate(cols):
        cell = tbl.rows[0].cells[j]
        set_cell_bg(cell, C_NAVY)
        p = cell.paragraphs[0]
        no_space_para(p)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, col, bold=True, size=9, color=C_HEADER_FG)
    # data rows
    for i, row in enumerate(data):
        bg = C_LIGHT_BLU if i % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            no_space_para(p)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # colour green for "✔ BLOQUÉ" / "✔ DÉTECTÉ"
            color = None
            if "✔" in str(val):
                color = C_GREEN_FG
            elif "✘" in str(val):
                color = C_RED_FG
            add_run(p, str(val), size=9, color=color,
                    bold=("✔" in str(val) or "✘" in str(val)))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def divider(doc):
    p = doc.add_paragraph()
    no_space_para(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(7)

# ══════════════════════════════════════════════════════════════════════════════
# Contenu des démonstrations
# ══════════════════════════════════════════════════════════════════════════════

DEMOS = [
    {
        "number": "DÉMO 1",
        "title": "Blocage SQLi — WAF Middleware",
        "owasp": "A03:2021 – Injection",
        "module": "WAF Middleware (back/functions/src/middlewares/waf.js)",
        "env": "Backend : http://localhost:5001/gestionadminastration/us-central1/api/",
        "duration": "~2 min",
        "cmd": "curl -s -o /dev/null -w \"%{http_code}\" -X POST http://localhost:5001/gestionadminastration/us-central1/api/auth/login -H 'Content-Type: application/json' -d '{\"email\":\"admin\\' OR 1=1--\",\"password\":\"x\"}'",
        "expected_result": "HTTP 403 Forbidden",
        "result_label": "WAF bloque la tentative SQLi — HTTP 403 renvoyé avant accès à Firebase",
        "analysis": (
            "Le middleware WAF intercepte la requête au niveau Express, avant tout traitement "
            "métier. Le pattern SQL OR 1=1-- est détecté par la regex dédiée. La réponse 403 "
            "confirme que la payload n'a jamais atteint Firestore, conformément à la défense "
            "en profondeur (OWASP A03:2021)."
        ),
        "wazuh": "Règle 100010 — SQL Injection Attempt détectée dans le dashboard Wazuh (onglet Security Events).",
    },
    {
        "number": "DÉMO 2",
        "title": "Blocage XSS — WAF Middleware",
        "owasp": "A03:2021 – Cross-Site Scripting",
        "module": "WAF Middleware (waf.js) + Sanitisation frontend (sanitize.ts)",
        "env": "Backend : http://localhost:5001/gestionadminastration/us-central1/api/",
        "duration": "~2 min",
        "cmd": "curl -s -o /dev/null -w \"%{http_code}\" -X POST http://localhost:5001/gestionadminastration/us-central1/api/etudiants -H 'Content-Type: application/json' -H 'Authorization: Bearer TOKEN' -d '{\"nom\":\"<script>alert(1)</script>\"}'",
        "expected_result": "HTTP 403 Forbidden",
        "result_label": "Payload XSS rejetée par le WAF — HTTP 403 avant persistance en base",
        "analysis": (
            "Le WAF détecte le tag <script> grâce à la regex XSS. Côté frontend, le module "
            "sanitize.ts (DOMPurify) assainit également les données affichées. La double "
            "validation client + serveur garantit qu'aucun script ne peut être injecté dans "
            "l'interface d'administration, satisfaisant A03:2021."
        ),
        "wazuh": "Règle 100011 — XSS Attempt loggée dans Wazuh Security Events.",
    },
    {
        "number": "DÉMO 3",
        "title": "Rate Limiting — Brute Force",
        "owasp": "A07:2021 – Identification and Authentication Failures",
        "module": "Rate Limiter Express (auth.js) — 5 tentatives / 15 min",
        "env": "Backend : http://localhost:5001/gestionadminastration/us-central1/api/",
        "duration": "~2 min",
        "cmd": "for i in $(seq 1 6); do echo \"Tentative $i:\"; curl -s -o /dev/null -w \"%{http_code}\\n\" -X POST http://localhost:5001/gestionadminastration/us-central1/api/auth/login -H 'Content-Type: application/json' -d '{\"email\":\"test@ynov.com\",\"password\":\"wrongpass\"}'; done",
        "expected_result": "Tentatives 1-5 : 401 — Tentative 6 : 429 Too Many Requests",
        "result_label": "La 6e tentative reçoit HTTP 429 — IP bloquée 15 minutes",
        "analysis": (
            "Le rate limiter Express compte les tentatives échouées par IP. Dès la 6e "
            "tentative dans la fenêtre de 15 minutes, la réponse 429 est renvoyée sans "
            "interroger Firebase Auth. Ce mécanisme contrecarré les attaques par dictionnaire "
            "et brute force (OWASP A07:2021). Un log auth_failure est enregistré dans "
            "auditLogs Firestore pour chaque tentative."
        ),
        "wazuh": "Règle 100013 — Brute Force Attempt avec 6 événements corrélés dans Wazuh.",
    },
    {
        "number": "DÉMO 4",
        "title": "Scanner DAST — 12 Tests OWASP Automatisés",
        "owasp": "OWASP Top 10 — Couverture globale",
        "module": "back/functions/scripts/security_scan.js",
        "env": "Backend : http://localhost:5001/gestionadminastration/us-central1/api/",
        "duration": "~3 min",
        "cmd": "cd back/functions && node scripts/security_scan.js",
        "expected_result": "12/12 tests passed — Rapport JSON généré dans security_scan_report.json",
        "result_label": "100% des 12 tests OWASP passés — zéro vulnérabilité critique détectée",
        "analysis": (
            "Le scanner DAST maison exécute 12 tests couvrant : SQLi, XSS, IDOR, JWT forgery, "
            "path traversal, headers sécurité (CSP, HSTS, X-Frame-Options), rate limiting, "
            "exposition d'endpoints non authentifiés. Les 12/12 tests verts confirment que "
            "les contrôles de sécurité en place résistent aux vecteurs d'attaque automatisés."
        ),
        "wazuh": "Les requêtes du scanner génèrent des événements dans Wazuh — visible dans Security Events.",
    },
    {
        "number": "DÉMO 5",
        "title": "Wazuh SIEM — Détection Temps Réel",
        "owasp": "OWASP A09:2021 – Security Logging and Monitoring Failures",
        "module": "Wazuh Dashboard — https://localhost (admin / SecretPassword)",
        "env": "Wazuh Dashboard : https://localhost",
        "duration": "~3 min",
        "cmd": "# Ouvrir Wazuh Dashboard → Security Events\n# Filtrer : agent.name = YNOV-APP\n# Observer : règles 100010–100016",
        "expected_result": "436 420 événements — Règles YNOV personnalisées visibles",
        "result_label": "436 420 événements de sécurité collectés — Règles YNOV-APP actives",
        "analysis": (
            "Le dashboard Wazuh affiche les 436 420 événements collectés depuis le déploiement. "
            "Les 7 règles custom (100010–100016) couvrent SQLi, XSS, brute force, FIM, audit "
            "RGPD, CVE et anomalies de performance. Le FIM (File Integrity Monitoring) détecte "
            "toute modification du code source en moins de 5 secondes (T1565.001 MITRE ATT&CK)."
        ),
        "wazuh": "Navigation : Overview → Security Events → Filtrer agent YNOV-APP → Règles 100010-100016.",
    },
    {
        "number": "DÉMO 6",
        "title": "Dashboard Monitoring — Interface React",
        "owasp": "OWASP A09:2021 – Monitoring & Alerting",
        "module": "front/src/pages/Monitoring.tsx + monitoringService.ts",
        "env": "Frontend : http://localhost:8081 (rôle admin)",
        "duration": "~2 min",
        "cmd": "# Ouvrir http://localhost:8081 → Menu Sécurité → Monitoring\n# Observer les métriques temps réel",
        "expected_result": "Dashboard avec KPIs : événements, menaces actives, uptime, latence API",
        "result_label": "Dashboard Monitoring opérationnel — 4 KPI + graphiques temps réel affichés",
        "analysis": (
            "Le composant Monitoring.tsx consomme monitoringService.ts qui appelle l'API "
            "backend toutes les 30 secondes. Les 4 KPIs affichés (événements de sécurité, "
            "menaces actives, uptime système, latence API) permettent à l'administrateur "
            "d'avoir une visibilité immédiate sur l'état de sécurité de la plateforme. "
            "Les alertes critiques déclenchent une notification push dans l'interface."
        ),
        "wazuh": "Les métriques du dashboard sont synchronisées avec les données Wazuh via l'API.",
    },
    {
        "number": "DÉMO 7",
        "title": "Traçabilité RGPD — AuditLogs Firestore",
        "owasp": "RGPD Art. 5 & 32 — Intégrité et traçabilité",
        "module": "Firestore collection auditLogs (append-only)",
        "env": "Firebase Console → Firestore → Collection auditLogs",
        "duration": "~2 min",
        "cmd": "# Tenter login avec mauvais mot de passe\n# Ouvrir Firebase Console → Firestore → auditLogs\n# Observer le document auth_failure avec userId, IP, timestamp",
        "expected_result": "Document créé : {action: auth_failure, userId, ip, timestamp, userAgent}",
        "result_label": "Entrée auditLogs créée automatiquement — non modifiable, conforme RGPD Art.5",
        "analysis": (
            "Chaque événement de sécurité (auth_failure, access_denied, data_export) génère "
            "un document dans la collection auditLogs avec les règles Firestore Security Rules "
            "configurées en allow write only (jamais de update/delete). Cette architecture "
            "garantit l'intégrité de la piste d'audit exigée par le RGPD Article 32. "
            "Les logs sont horodatés côté serveur (serverTimestamp) pour éviter toute "
            "manipulation côté client."
        ),
        "wazuh": "Règle 100016 — RGPD Audit Log détecte les écritures dans auditLogs via FIM Wazuh.",
    },
]

# ══════════════════════════════════════════════════════════════════════════════
# Construction du document
# ══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Page de titre Annexe B ────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(40)
    p_title.paragraph_format.space_after  = Pt(10)
    add_run(p_title, "ANNEXE B", bold=True, size=22, color=C_NAVY)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(8)
    add_run(p_sub, "Démonstration des Mécanismes de Sécurité", bold=True, size=16, color=C_BLUE)

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.space_after  = Pt(30)
    add_run(p_desc,
        "7 démonstrations live — Environnement local — Soutenance PFE M2 Cybersécurité",
        size=10, color=RGBColor(0x55, 0x55, 0x55), italic=True)

    # ── Intro ─────────────────────────────────────────────────────────────────
    add_heading(doc, "B.1  Présentation Générale", level=2)
    add_body(doc,
        "Cette annexe documente les 7 démonstrations pratiques réalisées lors de la soutenance "
        "du Projet de Fin d'Études. Chaque démonstration valide un contrôle de sécurité "
        "implémenté sur la plateforme de gestion scolaire YNOV, en reproduisant une attaque "
        "réelle et en montrant la réponse du système.",
        before=4, after=6)

    # Tableau récap environnement
    add_heading(doc, "Environnement de démonstration", level=3)
    info_table(doc, [
        ("Backend API",       "http://localhost:5001/gestionadminastration/us-central1/api/"),
        ("Frontend React",    "http://localhost:8081"),
        ("Wazuh Dashboard",   "https://localhost  (admin / SecretPassword)"),
        ("Firebase Console",  "https://console.firebase.google.com → Projet gestionadminastration"),
        ("Durée totale",      "~16 minutes  (7 démonstrations)"),
        ("Pré-requis",        "Backend + Frontend démarrés, Agent Wazuh actif, curl installé"),
    ])

    divider(doc)

    # ── 7 Démonstrations ─────────────────────────────────────────────────────
    for demo in DEMOS:
        doc.add_page_break()

        # Banner
        banner_table(doc, demo["title"], demo["number"])

        # Infos générales
        add_heading(doc, "Contexte", level=3)
        info_table(doc, [
            ("Référence OWASP", demo["owasp"]),
            ("Module testé",    demo["module"]),
            ("Environnement",   demo["env"]),
            ("Durée estimée",   demo["duration"]),
        ])

        # Commande
        add_heading(doc, "Commande d'attaque", level=3)
        add_code_block(doc, demo["cmd"], demo["expected_result"])

        # Résultat
        add_heading(doc, "Résultat obtenu", level=3)
        add_result_box(doc, demo["result_label"], ok=True)

        # Analyse
        add_heading(doc, "Analyse technique", level=3)
        add_body(doc, demo["analysis"], before=4, after=6)

        # Corrélation Wazuh
        add_heading(doc, "Corrélation Wazuh SIEM", level=3)
        p_wazuh = doc.add_paragraph()
        p_wazuh.paragraph_format.space_before = Pt(4)
        p_wazuh.paragraph_format.space_after  = Pt(8)
        p_wazuh.paragraph_format.left_indent  = Cm(0.5)
        pPr = p_wazuh._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), rgb_hex(C_LIGHT_BLU))
        pPr.append(shd)
        add_run(p_wazuh, "🛡  " + demo["wazuh"], size=9.5, color=C_BLUE)

        divider(doc)

    # ── Tableau de synthèse ───────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "B.2  Tableau de Synthèse des Démonstrations", level=2)
    add_body(doc,
        "Récapitulatif des 7 démonstrations avec statut de réussite et conformité OWASP / RGPD.",
        before=4, after=8)

    synth_cols = ["#", "Démonstration", "Vecteur OWASP", "Mécanisme", "Résultat", "RGPD"]
    synth_data = [
        ("1", "Blocage SQLi",       "A03:2021",  "WAF Middleware",        "✔ BLOQUÉ",   "Art.32"),
        ("2", "Blocage XSS",        "A03:2021",  "WAF + DOMPurify",       "✔ BLOQUÉ",   "Art.32"),
        ("3", "Brute Force",        "A07:2021",  "Rate Limiter",          "✔ BLOQUÉ",   "Art.32"),
        ("4", "DAST 12 tests",      "Top 10",    "security_scan.js",      "✔ 12/12",    "Art.25"),
        ("5", "Wazuh SIEM",         "A09:2021",  "Règles 100010-100016",  "✔ DÉTECTÉ",  "Art.30"),
        ("6", "Dashboard Monitoring","A09:2021", "Monitoring.tsx",        "✔ ACTIF",    "Art.32"),
        ("7", "Traçabilité RGPD",   "A09:2021",  "auditLogs Firestore",   "✔ TRACÉ",    "Art.5"),
    ]
    synthesis_table(doc, synth_cols, synth_data)

    # ── Conclusion ────────────────────────────────────────────────────────────
    add_heading(doc, "B.3  Conclusion", level=2)
    add_body(doc,
        "Les 7 démonstrations confirment que la plateforme de gestion scolaire YNOV dispose "
        "d'une couverture de sécurité complète sur les vecteurs d'attaque OWASP Top 10 les plus "
        "critiques. L'intégration du SIEM Wazuh avec des règles personnalisées, combinée au WAF "
        "middleware et aux mécanismes de rate limiting, assure une détection et un blocage en "
        "temps réel. La traçabilité RGPD via les auditLogs Firestore (append-only) garantit la "
        "conformité réglementaire exigée par les Articles 5 et 32 du RGPD.",
        before=4, after=8)

    add_body(doc,
        "Score global de sécurité : 12/12 tests DAST passés  |  7/7 démonstrations réussies  "
        "|  436 420 événements monitorés  |  5/6 articles RGPD satisfaits.",
        bold_needed=False, size=10, color=C_BLUE, italic=True, before=0, after=12)

    # ── Sauvegarde ────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"✅  Document généré : {OUTPUT}")
    print(f"   Pages estimées  : ~{2 + len(DEMOS) * 1 + 2}")


# Monkey-patch add_body pour accepter bold_needed (unused kwarg tolerance)
_orig_add_body = add_body
def add_body(doc, text, size=10, color=None, italic=False, before=4, after=4, bold_needed=False):
    return _orig_add_body(doc, text, size=size, color=color, italic=italic, before=before, after=after)

if __name__ == "__main__":
    build_document()

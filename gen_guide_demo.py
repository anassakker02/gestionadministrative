#!/usr/bin/env python3
"""
Génère GUIDE_DEMO_SECURITE.docx
Guide de démonstration live — PFE Cybersécurité YNOV 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = "/Users/anass/Downloads/frais-gestionScolaire 4"
DST  = os.path.join(BASE, "GUIDE_DEMO_SECURITE.docx")

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(text, level=1, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if not color else color
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5) if not color else color
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40) if not color else color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, color=None, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def code_block(text):
    """Bloc de code avec fond gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    # Fond gris clair via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E8F0FE")
    pPr.append(shd)
    return p

def result_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x10, 0x7C, 0x10)
    return p

def jury_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF8DC")
    pPr.append(shd)
    run = p.add_run(f'🎤  {text}')
    run.font.size   = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x7B, 0x44, 0x00)
    return p

def separator():
    p = doc.add_paragraph("─" * 90)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.runs[0]
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def etape_header(num, title, duration, color_hex):
    """En-tête coloré d'une étape."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.autofit = False
    widths = [Cm(2.5), Cm(11.5), Cm(2.5)]
    for i, w in enumerate(widths):
        tbl.columns[i].width = w
    cells = tbl.rows[0].cells
    for c in cells:
        set_cell_bg(c, color_hex)
    # Numéro
    r0 = cells[0].paragraphs[0].add_run(f"ÉTAPE {num}")
    r0.bold = True; r0.font.size = Pt(13); r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Titre
    r1 = cells[1].paragraphs[0].add_run(title)
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Durée
    r2 = cells[2].paragraphs[0].add_run(f"⏱  {duration}")
    r2.bold = True; r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run("GUIDE DE DÉMONSTRATION LIVE")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)

p2 = doc.add_paragraph()
run2 = p2.add_run("PFE Cybersécurité — Sécurisation Application Web Scolaire")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph()
run3 = p3.add_run("YNOV Campus 2026  ·  OWASP Top 10 + Wazuh SIEM + RGPD")
run3.font.size = Pt(12)
run3.italic = True
run3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Pré-requis ────────────────────────────────────────────────────────────────
heading("PRÉ-REQUIS AVANT LE PASSAGE", level=2, color=RGBColor(0xC0, 0x00, 0x00))
body("Ouvrir 3 terminaux et lancer :", bold=True)
code_block("# Terminal 1 — Backend\ncd \"/Users/anass/Downloads/frais-gestionScolaire 4/back/functions\"\nnode server_local.js\n# → http://localhost:5001")
code_block("# Terminal 2 — Frontend\ncd \"/Users/anass/Downloads/frais-gestionScolaire 4/front\"\nnpm run dev\n# → http://localhost:8081")
code_block("# Terminal 3 — Wazuh Docker (~2 min de démarrage)\ncd /tmp/wazuh-node\ndocker compose up -d\ndocker compose ps   # vérifier 3 conteneurs Up")
body("Wazuh Dashboard : https://localhost  →  admin / SecretPassword", bold=True, color=RGBColor(0x10, 0x7C, 0x10))
body("Total démonstration : ~8 minutes", bold=True)

separator()

# ═════════════════════════════════════════════════════════════════════════════
# ÉTAPE 1 — WAF SQLi
# ═════════════════════════════════════════════════════════════════════════════
etape_header("1", "WAF — Injection SQL bloquée", "1 min", "C0392B")

body("Dans Terminal 3, colle :", bold=True)
code_block('curl -s -o /dev/null -w "Réponse HTTP: %{http_code}\\n" \\\n  \'http://localhost:5001/api/etudiants?id=1%20OR%201=1\'')
body("Tu vois :", bold=True)
result_line("Réponse HTTP: 403")
doc.add_paragraph()
body("Tu dis au jury :", bold=True)
jury_quote('"J\'essaie une injection SQL dans l\'URL. Le WAF middleware intercepte la requête avant qu\'elle touche la base de données et retourne 403 Forbidden."')

separator()

# ═════════════════════════════════════════════════════════════════════════════
# ÉTAPE 2 — WAF XSS
# ═════════════════════════════════════════════════════════════════════════════
etape_header("2", "WAF — XSS bloqué", "30 sec", "E74C3C")

body("Dans Terminal 3, colle :", bold=True)
code_block('curl -s -o /dev/null -w "Réponse HTTP: %{http_code}\\n" \\\n  \'http://localhost:5001/api/etudiants?nom=<script>alert(1)</script>\'')
body("Tu vois :", bold=True)
result_line("Réponse HTTP: 403")
doc.add_paragraph()
body("Tu dis au jury :", bold=True)
jury_quote('"Même chose pour le XSS — script tag bloqué immédiatement. OWASP A03."')

separator()

# ═════════════════════════════════════════════════════════════════════════════
# ÉTAPE 3 — Brute Force
# ═════════════════════════════════════════════════════════════════════════════
etape_header("3", "Brute Force — Rate Limiting (OWASP A07)", "1 min", "E67E22")

body("Dans Terminal 3, colle :", bold=True)
code_block('for i in {1..6}; do\n  curl -s -o /dev/null -w "Tentative $i → %{http_code}\\n" \\\n  -X POST http://localhost:5001/api/auth/login \\\n  -H \'Content-Type: application/json\' \\\n  -d \'{"email":"admin@school.fr","password":"WRONG"}\'\ndone')
body("Tu vois :", bold=True)
result_line("Tentative 1 → 401")
result_line("Tentative 2 → 401")
result_line("Tentative 3 → 401")
result_line("Tentative 4 → 401")
result_line("Tentative 5 → 401")
result_line("Tentative 6 → 429  ← BLOQUÉ !")
doc.add_paragraph()
body("Tu dis au jury :", bold=True)
jury_quote('"5 tentatives échouées = compte verrouillé 15 minutes. Protection contre les attaques brute force — OWASP A07."')

separator()

# ═════════════════════════════════════════════════════════════════════════════
# ÉTAPE 4 — DAST
# ═════════════════════════════════════════════════════════════════════════════
etape_header("4", "Scan DAST automatique — OWASP Top 10", "1 min", "27AE60")

body("Dans Terminal 3, colle :", bold=True)
code_block('node "/Users/anass/Downloads/frais-gestionScolaire 4/back/functions/scripts/security_scan.js"')
body("Tu vois : liste de [PASS] en vert", bold=True)
result_line("[PASS] WAF actif — SQLi bloqué")
result_line("[PASS] Rate limiting — 429 après 5 req")
result_line("[PASS] JWT validation — 401 token invalide")
result_line("[PASS] Headers sécurité — HSTS présent")
p = doc.add_paragraph()
run = p.add_run("[WARN] 17 CVE détectées dans les dépendances npm")
run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
run.font.size = Pt(10)
run.font.bold = True
doc.add_paragraph()
body("Tu dis au jury :", bold=True)
jury_quote('"Ce script automatise tous les tests OWASP Top 10 — WAF, JWT, headers, rate limiting, CORS. Tout passe sauf 17 CVE dans les dépendances npm, détectées mais non critiques."')

separator()

# ═════════════════════════════════════════════════════════════════════════════
# ÉTAPE 5 — Wazuh SIEM
# ═════════════════════════════════════════════════════════════════════════════
etape_header("5", "Wazuh SIEM — Dashboard en direct", "2 min", "8E44AD")

body("Ouvrir dans le navigateur :", bold=True)
code_block("https://localhost   →   admin / SecretPassword")
doc.add_paragraph()

body("Script de navigation (dans l'ordre) :", bold=True)

# Tableau navigation Wazuh
tbl = doc.add_table(rows=5, cols=2)
tbl.style = "Table Grid"
tbl.autofit = False
tbl.columns[0].width = Cm(6)
tbl.columns[1].width = Cm(10.5)

headers = ["Clique sur", "Tu dis au jury"]
for j, h in enumerate(headers):
    cell = tbl.rows[0].cells[j]
    set_cell_bg(cell, "8E44AD")
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Security Events",
     '"436 420 événements collectés — chaque requête de l\'app est loguée ici."'),
    ("File Integrity",
     '"FIM surveille les fichiers système — 89% couverts. Si un attaquant modifie un fichier, alerte immédiate."'),
    ("MITRE ATT&CK",
     '"Corrélation automatique avec les techniques d\'attaque — ici T1565 Stored Data Manipulation détecté."'),
    ("Policy Monitoring",
     '"Vérification continue de la configuration OS — anomalies rootcheck signalées."'),
]

for i, (nav, say) in enumerate(rows_data):
    cells = tbl.rows[i+1].cells
    r0 = cells[0].paragraphs[0].add_run(nav)
    r0.bold = True; r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(0x5B, 0x21, 0x86)
    r1 = cells[1].paragraphs[0].add_run(say)
    r1.font.size = Pt(10); r1.font.italic = True
    r1.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)

doc.add_paragraph()

separator()

# ═════════════════════════════════════════════════════════════════════════════
# ÉTAPE 6 — Dashboard Monitoring
# ═════════════════════════════════════════════════════════════════════════════
etape_header("6", "Dashboard Monitoring intégré", "1 min", "16A085")

body("Dans le navigateur, ouvrir :", bold=True)
code_block("http://localhost:8081/monitoring")
body("Tu montres :", bold=True)
body("  •  Les alertes en temps réel")
body("  •  Les graphiques d'activité")
doc.add_paragraph()
body("Tu dis au jury :", bold=True)
jury_quote('"Ce dashboard est intégré dans l\'application — l\'admin voit les alertes de sécurité directement depuis l\'interface, sans aller sur Wazuh."')

separator()

# ═════════════════════════════════════════════════════════════════════════════
# ÉTAPE 7 — RGPD
# ═════════════════════════════════════════════════════════════════════════════
etape_header("7", "RGPD — Traçabilité Art.5 & Art.32", "1 min", "1F4E79")

body("Dans le navigateur, suivre ce flux :", bold=True)
code_block("1.  http://localhost:8081/users\n       → Voir la liste des utilisateurs\n\n2.  Cliquer sur l'icône ✏️  (Modifier) sur un utilisateur\n       → Changer le rôle ou le statut (actif/inactif)\n       → Cliquer Sauvegarder\n\n3.  http://localhost:8081/monitoring\n       → Voir l'événement tracé dans le monitoring")
doc.add_paragraph()
body("Tu vois dans le monitoring :", bold=True)
result_line("auditLog : timestamp + IP + identité de l'utilisateur")
result_line("DATA_EXPORT Art.15 disponible via l'API")
doc.add_paragraph()
body("Tu dis au jury :", bold=True)
jury_quote('"Chaque action est tracée dans Firestore avec timestamp, IP et identité de l\'utilisateur — conformité RGPD Art. 5 et Art. 32. L\'export des données Art. 15 est disponible via l\'API."')

separator()

# ═════════════════════════════════════════════════════════════════════════════
# RÉCAP FINAL
# ═════════════════════════════════════════════════════════════════════════════
heading("RÉCAPITULATIF — 8 minutes", level=2, color=RGBColor(0x1F, 0x49, 0x7D))

tbl2 = doc.add_table(rows=8, cols=3)
tbl2.style = "Table Grid"
tbl2.autofit = False
tbl2.columns[0].width = Cm(2)
tbl2.columns[1].width = Cm(9)
tbl2.columns[2].width = Cm(5.5)

hdrs = ["N°", "Démonstration", "Durée"]
for j, h in enumerate(hdrs):
    cell = tbl2.rows[0].cells[j]
    set_cell_bg(cell, "1F497D")
    run = cell.paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

recap = [
    ("01", "WAF — Injection SQL (SQLi) bloquée → 403",   "1 min"),
    ("02", "WAF — XSS bloqué → 403",                     "30 sec"),
    ("03", "Brute Force → Rate Limiting 429 (OWASP A07)","1 min"),
    ("04", "Scan DAST automatique OWASP Top 10",          "1 min"),
    ("05", "Wazuh SIEM — 436K events · FIM · MITRE",      "2 min"),
    ("06", "Dashboard /monitoring React intégré",          "1 min"),
    ("07", "RGPD traçabilité Art.5 & Art.32 Firestore",    "1 min"),
]

colors = ["C0392B","E74C3C","E67E22","27AE60","8E44AD","16A085","1F4E79"]
for i, (num, desc, dur) in enumerate(recap):
    cells = tbl2.rows[i+1].cells
    set_cell_bg(cells[0], colors[i])
    r0 = cells[0].paragraphs[0].add_run(num)
    r0.bold = True; r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r0.font.size = Pt(11)
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cells[1].paragraphs[0].add_run(desc)
    r1.font.size = Pt(10); r1.font.bold = True
    r2 = cells[2].paragraphs[0].add_run(dur)
    r2.font.size = Pt(10)
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
body("💡 CONSEIL : Ouvrir les 3 terminaux ET Wazuh AVANT ton passage. Wazuh démarre en ~2 min.",
     bold=True, color=RGBColor(0xC0, 0x00, 0x00), size=11)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
doc.save(DST)
print(f"✅ Fichier généré : {DST}")
